"""
Resume Matcher Web App
Upload resume + paste job URL → Get compatibility score + recommendations
"""

import streamlit as st
import traceback as _tb

# Page config must be first Streamlit call
st.set_page_config(
    page_title="ResumeSync - AI-Powered Job Matching",
    page_icon="🔄",
    layout="wide"
)

try:
    import streamlit.components.v1 as _components
    import os
    import re
    import html
    import json
    import hashlib
    import urllib.parse
    from datetime import datetime
    from typing import Optional
    from dotenv import load_dotenv
    from anthropic import Anthropic
    import requests
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import pypdf
    import io
    import base64
    import openpyxl
    from openpyxl import load_workbook
    try:
        from supabase import create_client as _supabase_create_client
        _SUPABASE_AVAILABLE = True
    except Exception:
        _SUPABASE_AVAILABLE = False

    # Load environment
    load_dotenv()

    # Configuration
    MODEL_NAME = "claude-sonnet-4-6"
    TRACKER_PATH = os.path.join(os.path.dirname(__file__), "job_applications.xlsx")
    COVER_LETTERS_PATH = os.path.join(os.path.dirname(__file__), "cover_letters")

    _IMPORT_ERROR = None
except Exception as _ie:
    _IMPORT_ERROR = _ie

# Custom CSS — VisualizePro · Forest + Sage design system
st.markdown("""
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Bricolage+Grotesque:opsz,wght@12..96,400;12..96,600;12..96,700;12..96,800&family=Lora:ital,wght@1,500;1,600&family=Space+Mono:wght@400;700&family=DM+Sans:opsz,wght@9..40,300;9..40,400;9..40,500&display=swap">
<style>

    /* ── Design tokens ── */
    :root {
        --bg:          #071512;
        --bg-2:        #0a1b16;
        --panel:       #0c2019;
        --panel-2:     #0f271e;
        --ink:         #ecf4ee;
        --ink2:        #9fb6a8;
        --ink3:        #6e8a7b;
        --mint:        #7ad79f;
        --mint-bright: #94e6b1;
        --mint-deep:   #4fae7a;
        --mint-dim:    rgba(122,215,159,0.10);
        --mint-dim2:   rgba(122,215,159,0.16);
        --mint-border: rgba(122,215,159,0.22);
        --amber:       #e0a14a;
        --line:        rgba(159,182,168,0.12);
        --line-2:      rgba(159,182,168,0.07);
        /* aliases used in component CSS */
        --gold:        #7ad79f;
        --gold-soft:   #94e6b1;
        --muted:       #6e8a7b;
        --line-strong: rgba(122,215,159,0.22);
    }

    /* ── App background ── */
    .stApp, .main {
        background: var(--bg) !important;
        background-image:
            radial-gradient(ellipse 80% 60% at 15% 20%, rgba(109,193,138,0.07), transparent 60%),
            radial-gradient(ellipse 60% 50% at 95% 90%, rgba(122,215,159,0.04), transparent 60%) !important;
        font-family: 'DM Sans', system-ui, sans-serif !important;
        color: var(--ink) !important;
    }

    /* ── Sidebar ── */
    section[data-testid="stSidebar"] {
        background: var(--bg-2) !important;
        border-right: 1px solid var(--line) !important;
    }
    section[data-testid="stSidebar"] .stMarkdown { color: var(--ink) !important; }
    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3 { color: var(--ink) !important; }

    /* ── Typography ── */
    h1, h2, h3 {
        font-family: 'DM Sans', sans-serif !important;
        color: var(--ink) !important;
        font-weight: 700 !important;
    }
    /* ── Analysis result headings (LLM markdown ## / ###) — keep compact ── */
    .stMarkdown h2 {
        font-size: 1.05rem !important;
        font-weight: 700 !important;
        margin: 1.4rem 0 0.35rem !important;
        padding-bottom: 0.3rem !important;
        border-bottom: 1px solid rgba(255,255,255,0.08) !important;
        letter-spacing: 0 !important;
    }
    .stMarkdown h3 {
        font-size: 0.92rem !important;
        font-weight: 600 !important;
        margin: 0.9rem 0 0.25rem !important;
    }
    .stMarkdown { font-family: 'DM Sans', sans-serif !important; color: var(--ink) !important; }
    p, li, label { color: var(--ink) !important; }

    /* ── Primary button — sage green fill ── */
    .stButton > button {
        background: var(--gold) !important;
        color: #000000 !important;
        border: none !important;
        border-radius: 4px !important;
        padding: 16px 32px !important;
        font-weight: 600 !important;
        font-size: 15px !important;
        font-family: 'DM Sans', sans-serif !important;
        transition: transform .15s, background .2s !important;
        min-height: 52px !important;
    }
    .stButton > button * { color: #000000 !important; }
    .stButton > button:hover {
        background: var(--gold-soft) !important;
        color: #000000 !important;
        transform: translateY(-1px) !important;
        box-shadow: none !important;
    }
    .stButton > button:hover * { color: #000000 !important; }

    /* ── Download button — ghost style ── */
    .stDownloadButton > button {
        background: transparent !important;
        color: var(--ink) !important;
        border: 1px solid var(--line-strong) !important;
        border-radius: 4px !important;
        padding: 12px 28px !important;
        font-weight: 600 !important;
        font-family: 'DM Sans', sans-serif !important;
        transition: background .2s, color .2s !important;
    }
    .stDownloadButton > button:hover {
        background: var(--gold) !important;
        color: #0a1f17 !important;
        border-color: var(--gold) !important;
    }

    /* ── Form elements ── */
    .stFileUploader, .stTextArea, .stTextInput { font-family: 'DM Sans', sans-serif !important; }

    .stTextInput input, .stTextArea textarea {
        font-family: 'DM Sans', sans-serif !important;
        background: var(--bg-2) !important;
        color: #ffffff !important;
        border-radius: 4px !important;
        border-color: var(--line-strong) !important;
    }
    .stSelectbox [data-baseweb="select"] > div:first-child {
        background: var(--bg-2) !important;
        border-color: var(--line-strong) !important;
    }
    .stSelectbox [data-baseweb="select"] [data-baseweb="value-container"] *,
    .stSelectbox [data-baseweb="select"] [role="combobox"],
    .stSelectbox [data-baseweb="select"] input,
    .stSelectbox [data-baseweb="single-select"],
    .stSelectbox [data-baseweb="single-select"] *,
    [data-baseweb="value-container"],
    [data-baseweb="value-container"] *,
    [data-baseweb="value-container"] > div,
    [data-baseweb="value-container"] > div > div {
        color: #ecf4ee !important;
        font-family: 'Bricolage Grotesque', system-ui, sans-serif !important;
        font-weight: 800 !important;
        font-size: 13px !important;
    }
    .stSelectbox [data-baseweb="select"] svg { fill: var(--muted) !important; }
    [data-baseweb="popover"] [data-baseweb="menu"] {
        background: var(--bg-2) !important;
    }
    [data-baseweb="popover"] [role="option"] {
        background: var(--bg-2) !important;
        color: #ffffff !important;
    }
    [data-baseweb="popover"] [role="option"]:hover,
    [data-baseweb="popover"] [aria-selected="true"] {
        background: rgba(109,193,138,0.15) !important;
        color: var(--gold) !important;
    }
    .stTextInput input::placeholder, .stTextArea textarea::placeholder {
        color: var(--muted) !important;
        opacity: 1 !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--gold) !important;
        box-shadow: 0 0 0 1px var(--gold) !important;
    }

    /* ── File uploader — dark theme, consistent with other inputs ── */
    [data-testid="stFileUploaderDropzone"] {
        background: var(--bg-2) !important;
        border: 1px solid var(--line-strong) !important;
        border-radius: 4px !important;
    }
    [data-testid="stFileUploaderDropzone"]:hover {
        border-color: var(--gold) !important;
    }
    /* Hide the duplicate upload SVG icon + its label — keep only browse button + file-type hint */
    [data-testid="stFileUploaderDropzone"] svg { display: none !important; }
    [data-testid="stFileUploaderDropzoneInstructions"] > div > div:first-child { display: none !important; }
    [data-testid="stFileUploaderDropzoneInstructions"],
    [data-testid="stFileUploaderDropzoneInstructions"] *,
    [data-testid="stFileUploaderDropzone"] span,
    [data-testid="stFileUploaderDropzone"] p,
    [data-testid="stFileUploaderDropzone"] small {
        color: var(--muted) !important;
        font-family: 'DM Sans', sans-serif !important;
    }
    [data-testid="stFileUploaderDropzone"] button {
        background: rgba(122,215,159,0.10) !important;
        border: 1px solid var(--line-strong) !important;
        border-radius: 4px !important;
        min-width: 120px !important;
    }
    [data-testid="stFileUploaderDropzone"] button * {
        display: none !important;
    }
    [data-testid="stFileUploaderDropzone"] button::after {
        content: "Browse files";
        display: inline !important;
        font-size: 0.875rem !important;
        font-family: 'DM Sans', sans-serif !important;
        color: var(--gold) !important;
    }
    [data-testid="stFileUploaderDropzone"] button:hover {
        background: var(--gold) !important;
    }
    [data-testid="stFileUploaderDropzone"] button:hover::after {
        color: #0a1f17 !important;
    }
    /* Uploaded file pill */
    [data-testid="uploadedFile"] {
        background: rgba(122,215,159,0.07) !important;
        border: 1px solid var(--line) !important;
        border-radius: 4px !important;
    }
    [data-testid="uploadedFile"] *,
    [data-testid="uploadedFileData"] * { color: var(--ink) !important; }
    button[title="Remove file"] { color: var(--muted) !important; }

    /* ── Feedback boxes ── */
    .stSuccess, .stSuccess * {
        background-color: rgba(122,215,159,0.07) !important;
        color: #000000 !important;
        border-left: 4px solid var(--gold) !important;
    }
    .stInfo, .stInfo * {
        background-color: rgba(122,215,159,0.04) !important;
        border-left: 4px solid var(--line-strong) !important;
        color: #000000 !important;
    }
    .stError { font-family: 'DM Sans', sans-serif !important; }
    .stStatus { font-family: 'DM Sans', sans-serif !important; }

    /* ── Spinner ── */
    .stSpinner > div { border-top-color: var(--gold) !important; }

    /* ── Divider ── */
    hr { border-color: var(--line) !important; opacity: 1 !important; }

    /* ── Block container ── */
    .block-container, .main .block-container { padding-top: 0.75rem !important; }

    /* ── Remove Streamlit's top header bar entirely ── */
    [data-testid="stHeader"] { display: none !important; }
    #MainMenu { visibility: hidden !important; }
    footer { visibility: hidden !important; }

    /* ── Hide scrollbars ── */
    ::-webkit-scrollbar { display: none !important; }
    * { scrollbar-width: none !important; -ms-overflow-style: none !important; }

    /* ── Radio ── */
    .stRadio label { color: var(--muted) !important; }

    /* ── Checkbox ── */
    .stCheckbox label { color: var(--ink) !important; font-family: 'DM Sans', sans-serif !important; }

    /* ── Dataframe ── */
    .stDataFrame { border: 1px solid var(--line) !important; border-radius: 8px !important; }

    /* ── Cover letter white box — force dark ink ── */
    .cl-box, .cl-box * { color: #1b1b1b !important; font-family: 'DM Sans', sans-serif !important; }
</style>
""", unsafe_allow_html=True)


# ============= HELPERS =============

def _read_secret(key: str) -> str:
    """Read a secret: env var → st.secrets → local secrets.toml."""
    # 1. Environment variable
    val = os.getenv(key, "").strip()
    if val:
        return val
    # 2. Streamlit Cloud secrets manager
    try:
        val = str(st.secrets.get(key, "")).strip()
        if val:
            return val
    except Exception:
        pass
    # 3. Local .streamlit/secrets.toml (dev)
    secrets_path = os.path.join(os.path.dirname(__file__), ".streamlit", "secrets.toml")
    if os.path.exists(secrets_path):
        text = open(secrets_path, encoding="utf-8").read()
        m = re.search(rf'{key}\s*=\s*"([^"]+)"', text)
        if m:
            return m.group(1).strip()
    return ""


def get_client():
    api_key = _read_secret("ANTHROPIC_API_KEY")
    if not api_key or api_key == "your-api-key-here":
        return None
    return Anthropic(api_key=api_key)


@st.cache_resource
def get_supabase():
    if not _SUPABASE_AVAILABLE:
        return None
    url = _read_secret("SUPABASE_URL")
    key = _read_secret("SUPABASE_KEY")
    if url and key:
        try:
            return _supabase_create_client(url, key)
        except Exception:
            return None
    return None


# ── Auth & subscription helpers ──────────────────────────────────────────────

FREE_CREDITS = 3

PACK_INFO = {
    "pack5":  {"name": "Starter Pack", "credits": 5,  "price": "A$9",  "amount": "9.00",  "color": "#6fb1e0"},
    "pack20": {"name": "Value Pack",   "credits": 20, "price": "A$15", "amount": "15.00", "color": "#7ad79f"},
    "pack30": {"name": "Pro Pack",     "credits": 30, "price": "A$20", "amount": "20.00", "color": "#a78bfa"},
    "pack50": {"name": "Max Pack",     "credits": 50, "price": "A$30", "amount": "30.00", "color": "#f59e0b"},
}

PAY_ADVANCED_URL = "https://pad.live/pa53992/pay"
APP_URL = "https://resumesync-adc8kjcujdwh8jnnhcgcpe.streamlit.app"


def _fresh_supabase():
    """Non-cached Supabase client used for auth calls."""
    if not _SUPABASE_AVAILABLE:
        return None
    url = _read_secret("SUPABASE_URL")
    key = _read_secret("SUPABASE_KEY")
    if url and key:
        try:
            return _supabase_create_client(url, key)
        except Exception:
            pass
    return None


def _auth_user_dict(resp) -> Optional[dict]:
    """Extract user dict from a Supabase auth response."""
    if resp and resp.user and resp.session:
        return {
            "user_id":       resp.user.id,
            "email":         resp.user.email or "",
            "access_token":  resp.session.access_token,
            "refresh_token": resp.session.refresh_token,
        }
    return None


def sign_in(email: str, password: str) -> tuple:
    """Sign in with email + password. Returns (user_dict | None, error_str)."""
    sb = _fresh_supabase()
    if not sb:
        return None, "Supabase not configured."
    try:
        resp = sb.auth.sign_in_with_password({"email": email, "password": password})
        user = _auth_user_dict(resp)
        return user, ("" if user else "Invalid email or password.")
    except Exception as e:
        msg = str(e)
        if "Invalid login credentials" in msg:
            return None, "Invalid email or password."
        return None, msg


def sign_up(email: str, password: str) -> tuple:
    """Create a pre-confirmed account using the service role key, then sign in."""
    if not _SUPABASE_AVAILABLE:
        return None, "Supabase not configured."
    url  = _read_secret("SUPABASE_URL")
    skey = _read_secret("SUPABASE_SERVICE_ROLE_KEY")
    if not url or not skey:
        # Fall back to regular signup if no service role key
        return _sign_up_basic(email, password)
    try:
        admin_sb = _supabase_create_client(url, skey)
        # Create user as already confirmed — no email needed
        admin_sb.auth.admin.create_user({
            "email": email,
            "password": password,
            "email_confirm": True,
        })
    except Exception as e:
        msg = str(e)
        if "already been registered" in msg.lower() or "already registered" in msg.lower() or "already exists" in msg.lower():
            return None, "An account with that email already exists. Please sign in."
        return None, msg
    # Now sign in with the new credentials
    return sign_in(email, password)


def _sign_up_basic(email: str, password: str) -> tuple:
    """Fallback signup without service role key (requires email confirmation off in Supabase)."""
    sb = _fresh_supabase()
    if not sb:
        return None, "Supabase not configured."
    try:
        resp = sb.auth.sign_up({"email": email, "password": password})
        user = _auth_user_dict(resp)
        if user:
            return user, ""
        return None, "Account created — please check your inbox to confirm your email, then sign in."
    except Exception as e:
        msg = str(e)
        if "already registered" in msg.lower() or "already been registered" in msg.lower():
            return None, "An account with that email already exists. Please sign in."
        return None, msg


def get_or_create_profile(user_id: str, email: str) -> dict:
    """Load user profile from Supabase, creating it on first login."""
    sb = get_supabase()
    if not sb:
        return {"id": user_id, "email": email, "credits": FREE_CREDITS}
    try:
        res = sb.table("profiles").select("*").eq("id", user_id).execute()
        if res.data:
            profile = res.data[0]
            if profile.get("credits") is None:
                sb.table("profiles").update({"credits": FREE_CREDITS}).eq("id", user_id).execute()
                profile["credits"] = FREE_CREDITS
            return profile
        new_profile = {"id": user_id, "email": email, "credits": FREE_CREDITS}
        sb.table("profiles").insert(new_profile).execute()
        return new_profile
    except Exception:
        return {"id": user_id, "email": email, "credits": FREE_CREDITS}


def can_run_analysis(profile: dict) -> tuple:
    """Returns (allowed: bool, reason: str)."""
    credits = int(profile.get("credits") or 0)
    if credits > 0:
        return True, ""
    return False, "You have no analyses left. Top up with a credit pack to continue."


def decrement_credits(user_id: str):
    """Deduct 1 credit after a successful analysis."""
    sb = get_supabase()
    if not sb or not user_id:
        return
    try:
        res = sb.table("profiles").select("credits").eq("id", user_id).execute()
        if res.data:
            new_credits = max(0, int(res.data[0].get("credits") or 0) - 1)
            sb.table("profiles").update({"credits": new_credits}).eq("id", user_id).execute()
            if "user_profile" in st.session_state:
                st.session_state["user_profile"]["credits"] = new_credits
    except Exception:
        pass


def add_credits(user_id: str, amount: int) -> bool:
    """Add credits to a user's account after a pack purchase."""
    sb = _fresh_supabase()
    if not sb:
        return False
    try:
        res = sb.table("profiles").select("credits").eq("id", user_id).execute()
        current = int(res.data[0].get("credits") or 0) if res.data else 0
        sb.table("profiles").update({"credits": current + amount}).eq("id", user_id).execute()
        return True
    except Exception:
        return False


def extract_text_from_pdf(file):
    try:
        pdf_reader = pypdf.PdfReader(file)
        return "".join(page.extract_text() for page in pdf_reader.pages)
    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def extract_text_from_docx(file):
    try:
        file_bytes = file.getvalue() if hasattr(file, 'getvalue') else file.read()
        doc = Document(io.BytesIO(file_bytes))
        texts = []
        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        # Table cells — many resume templates use tables for layout
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            texts.append(para.text)
        return "\n".join(texts)
    except Exception as e:
        return f"Error reading Word document: {str(e)}"


def analyze_docx_ats_structure(file_bytes: bytes) -> dict:
    """Inspect a DOCX's actual structure to simulate what a strict ATS parser
    (Taleo/Workday-style: paragraphs only, no tables, no headers/footers, no text boxes)
    would extract versus a lenient one (Greenhouse/Lever-style: paragraphs + tables)."""
    doc = Document(io.BytesIO(file_bytes))

    strict_lines = [p.text for p in doc.paragraphs if p.text.strip()]
    strict_text = "\n".join(strict_lines)

    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        table_lines.append(para.text)
    table_text = "\n".join(table_lines)

    header_footer_lines = []
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is not None and not part.is_linked_to_previous:
                for para in part.paragraphs:
                    if para.text.strip():
                        header_footer_lines.append(para.text)
    header_footer_text = "\n".join(header_footer_lines)

    lenient_text = strict_text + (("\n" + table_text) if table_text else "")

    has_multicolumn = False
    for section in doc.sections:
        cols_el = section._sectPr.find(qn('w:cols'))
        if cols_el is not None:
            num = cols_el.get(qn('w:num'))
            if num and int(num) > 1:
                has_multicolumn = True
                break

    has_textboxes = doc.element.body.find('.//' + qn('w:txbxContent')) is not None

    has_skills_section = any(
        len(line.strip()) < 40 and re.search(r'skill|competenc', line, re.IGNORECASE)
        for line in strict_lines
    )

    contact_pattern = re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+|(\+?\d[\d\-\s()]{7,}\d)')
    has_contact_in_header_footer = bool(contact_pattern.search(header_footer_text))

    total_chars = len(strict_text) + len(table_text) + len(header_footer_text)
    at_risk_chars = len(table_text) + len(header_footer_text)
    at_risk_pct = round(100 * at_risk_chars / total_chars) if total_chars else 0

    return {
        'strict_text': strict_text,
        'lenient_text': lenient_text,
        'table_text': table_text,
        'header_footer_text': header_footer_text,
        'has_tables': len(doc.tables) > 0,
        'table_count': len(doc.tables),
        'has_multicolumn': has_multicolumn,
        'has_textboxes': has_textboxes,
        'has_skills_section': has_skills_section,
        'has_contact_in_header_footer': has_contact_in_header_footer,
        'at_risk_pct': at_risk_pct,
    }


def build_ats_safe_docx(file_bytes: bytes, struct: dict) -> tuple:
    """Non-destructively duplicate content that lives in tables/headers/footers as
    plain paragraphs, so a strict ATS parser (which skips tables and headers/footers)
    still captures it. Original tables/headers/footers are left untouched — this only
    adds visible lines, it never deletes or rewrites existing content."""
    doc = Document(io.BytesIO(file_bytes))
    fixes = []

    if struct.get('has_contact_in_header_footer') and struct.get('header_footer_text'):
        contact_line = struct['header_footer_text'].replace('\n', '   |   ')
        new_para = doc.add_paragraph()
        run = new_para.add_run(contact_line)
        run.bold = True
        body = doc.element.body
        body.remove(new_para._p)
        body.insert(0, new_para._p)
        _preview = contact_line[:70] + ('…' if len(contact_line) > 70 else '')
        fixes.append(f'Added your header/footer contact info as a visible line at the top: "{_preview}"')

    if doc.tables:
        for i, table in enumerate(doc.tables, start=1):
            cell_texts = []
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in cell_texts:
                        cell_texts.append(t)
            if not cell_texts:
                continue
            flattened = ' • '.join(cell_texts)
            new_para = doc.add_paragraph()
            new_para.add_run(flattened)
            table._tbl.addnext(new_para._p)
            _preview = flattened[:70] + ('…' if len(flattened) > 70 else '')
            fixes.append(f'Added a plain-text line under table {i} duplicating its content: "{_preview}"')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), fixes


def extract_text_from_txt(file):
    try:
        return file.read().decode('utf-8')
    except Exception as e:
        return f"Error reading text file: {str(e)}"


def extract_resume_text(uploaded_file):
    if uploaded_file is None:
        return "No file uploaded"
    file_type = uploaded_file.name.split('.')[-1].lower()
    if file_type == 'pdf':
        return extract_text_from_pdf(uploaded_file)
    elif file_type == 'docx':
        return extract_text_from_docx(uploaded_file)
    elif file_type == 'doc':
        return "Error: Old .doc format is not supported. Please open your resume in Word, save it as .docx, then re-upload."
    elif file_type == 'txt':
        return extract_text_from_txt(uploaded_file)
    return "Unsupported file type. Please upload PDF, DOCX, or TXT."



ATS_PATTERNS = [
    ('Greenhouse',      r'boards\.greenhouse\.io|greenhouse\.io/embed',
     '#3ab87a', '🌿',
     "Greenhouse parses PDFs reasonably well but still prefers .docx. "
     "Scoring is heavily keyword-frequency based — use the job posting's exact terms. "
     "Single-column layout is safest. Standard section headings (Experience, Skills, Education) are required.",
     ["Use exact keywords from the job description — frequency matters",
      "Submit .docx over PDF when given the option",
      "Include both acronym and full form for every key term (e.g. 'SQL (Structured Query Language)')",
      "Avoid tables and text boxes — single-column only"]),

    ('Lever',           r'jobs\.lever\.co',
     '#4d90fe', '⚙️',
     "Lever is one of the more modern ATS systems and handles PDFs well. "
     "It emphasises keyword matching but also passes resumes to humans quickly — formatting matters more here than in older systems.",
     ["PDF is fine, but .docx is still safer",
      "Lever surfaces resumes to recruiters early — clean formatting and strong summary matter",
      "Include a Skills section — Lever's parsing highlights it separately",
      "Quantified achievements stand out since humans review quickly"]),

    ('Workday',         r'myworkdayjobs\.com|wd\d+\.myworkdayjobs',
     '#e87722', '🔶',
     "Workday is used by large enterprises and is notoriously strict. "
     "It often re-formats your resume entirely and pulls data into structured fields — "
     "so your formatting matters less, but your keyword matches matter enormously.",
     ["Workday strips most formatting — focus on keywords, not layout",
      "Exact keyword matches are critical — Workday's scoring is very literal",
      "Fill in every field manually even if it duplicates your resume",
      "Keep the resume file itself simple — Workday may ignore complex layouts entirely"]),

    ('Taleo',           r'taleo\.net|tbe\.taleo',
     '#cc0000', '🔴',
     "Taleo (used by large corporates, banks, government) is the strictest and oldest major ATS. "
     "It is notoriously bad at parsing PDFs, tables, columns, and anything non-standard.",
     ["Submit .docx — Taleo's PDF parsing is unreliable",
      "Absolutely no tables, columns, text boxes, or headers/footers",
      "Standard section names only: Experience, Education, Skills, Summary",
      "Keyword matching is purely literal — no semantic understanding whatsoever",
      "Dates must be formatted consistently (e.g. Jan 2022 – Mar 2024)"]),

    ('Ashby',           r'jobs\.ashbyhq\.com',
     '#7c3aed', '🟣',
     "Ashby is a modern ATS used by tech startups. It handles formatting well and passes resumes to humans quickly.",
     ["Clean formatting is appreciated — humans see your resume early",
      "Skills section is scanned separately — include one",
      "Keyword matching still matters but Ashby is more forgiving than Taleo/Workday"]),

    ('SmartRecruiters', r'smartrecruiters\.com',
     '#00b4d8', '🔵',
     "SmartRecruiters handles modern formats well and is relatively lenient on formatting.",
     ["PDF and .docx both work well",
      "Keyword matching is important — include both acronyms and full terms",
      "Skills section is highlighted in the recruiter view"]),

    ('iCIMS',           r'icims\.com',
     '#005eb8', '🏢',
     "iCIMS is widely used in large enterprises and healthcare. It parses reasonably well but prefers simple formatting.",
     ["Simple single-column layout strongly preferred",
      "Submit .docx when possible",
      "Keyword frequency matters — repeat key terms naturally across sections"]),

    ('Workable',        r'apply\.workable\.com',
     '#ff6b35', '🟠',
     "Workable is used by SMEs and handles modern formats well. Resumes reach humans relatively quickly.",
     ["Both PDF and .docx work fine",
      "Include a Skills section — Workable highlights it in recruiter view",
      "Quantified results help since humans review early"]),

    ('BambooHR',        r'bamboohr\.com',
     '#73c41d', '🌱',
     "BambooHR is used by SMEs and is relatively lenient. Resumes often reach humans quickly.",
     ["Standard formatting is fine — no major restrictions",
      "Keyword matching still applies — mirror the job description language",
      "Keep it clean and concise — hiring managers usually review directly"]),
]


def detect_ats(url: str) -> dict | None:
    """Detect which ATS is being used from a job or application URL. Returns ATS info dict or None."""
    if not url or url == 'Manual Input':
        return None
    for name, pattern, color, icon, summary, tips in ATS_PATTERNS:
        if re.search(pattern, url, re.IGNORECASE):
            return {'name': name, 'color': color, 'icon': icon,
                    'summary': summary, 'tips': tips}
    return None


def detect_ats_from_text(text: str) -> dict | None:
    """Scan pasted job description text for embedded ATS URLs."""
    urls = re.findall(r'https?://[^\s\)"\'<>]+', text)
    for url in urls:
        result = detect_ats(url)
        if result:
            return result
    return None


def scrape_job_url(url: str) -> dict:
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'DNT': '1',
            'Connection': 'keep-alive',
        }
        response = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        lines = [l.strip() for l in soup.get_text(separator='\n', strip=True).splitlines() if l.strip()]
        return {'success': True, 'content': '\n'.join(lines)[:8000], 'url': url}
    except Exception as e:
        return {'success': False, 'error': str(e)}


# ============= ANALYSIS =============

def analyze_resume_vs_job(resume_text: str, job_content: str, job_url: str, client):
    """Tiered keyword extraction, three-way match (exact/implied/missing), ATS formatting flags."""
    prompt = f"""You are an expert ATS analyst and resume strategist. Produce a rigorous, honest gap analysis — not flattery. Be specific and conservative: when in doubt, mark something as implied or missing rather than matched.

**MY RESUME:**
{resume_text}

**JOB POSTING (from {job_url}):**
{job_content}

---

## COMPATIBILITY SCORE: [X]%

Weight your score as follows — hard requirements matter most:
- Hard requirements (must-haves): 50%
- Core skills and tools: 30%
- Preferred / nice-to-have: 15%
- Soft skills and culture fit: 5%

Provide one sentence explaining the primary driver of the score.

## JOB DETAILS
- **Job Title:**
- **Company:**
- **Location:**
- **Employment Type:**

## KEYWORD EXTRACTION & MATCH

Extract all significant requirements from the job posting into four tiers, then assign each a match status:
- **EXACT** — the term or a recognised synonym/acronym appears clearly in the resume
- **IMPLIED** — the resume describes the underlying experience but without matching terminology (note what bridges it)
- **MISSING** — no evidence in the resume

Be conservative: partial overlap (e.g. "manage" vs "management") = IMPLIED. "Cloud platforms" is NOT "AWS". "Scripting" is NOT "Python".

### Hard Requirements *(degree, certification, years of experience, "required" / "must have")*
| Requirement | Status | Notes |
|-------------|--------|-------|
| [item] | EXACT / IMPLIED / MISSING | [note] |

### Core Skills & Tools *(technologies, software, methodologies explicitly named)*
| Skill / Tool | Status | Notes |
|--------------|--------|-------|
| [item] | EXACT / IMPLIED / MISSING | [note] |

### Preferred Qualifications *("preferred", "plus", "bonus", "nice to have")*
| Qualification | Status | Notes |
|---------------|--------|-------|
| [item] | EXACT / IMPLIED / MISSING | [note] |

### Soft Skills & Competencies
| Competency | Status | Notes |
|------------|--------|-------|
| [item] | EXACT / IMPLIED / MISSING | [note] |

## CLOSING THE GAPS

For every IMPLIED or MISSING item, provide either:
- A specific, honest reword the candidate could use **only if the underlying experience genuinely applies** — do not invent experience
- OR flag it as **"Verify with candidate"** if it cannot be determined from the resume

## RESUME SECTION REVIEW

**Professional Summary:** Is it tailored? Does it open with the right strengths for this specific role?
**Work Experience:** Score bullets against: Action Verb + Exact Skill + Quantified Impact + Business Outcome. Name specific weak bullets.
**Skills Section:** Does it mirror the job's exact terminology? What's missing or misworded?
**Education / Certifications:** Any required credentials absent?

## SPECIFIC IMPROVEMENT RECOMMENDATIONS

Provide 6–8 concrete rewrites of existing resume text. Each must:
- Quote the EXACT current text from the resume (copy it verbatim including bullet characters)
- Replace weak verbs, vague language, or synonym mismatches — use the job ad's exact terminology where the candidate clearly has the skill
- Apply the formula: Action Verb + Exact Skill + Quantified Impact + Business Outcome
- NEVER add skills, tools, or experience not already in the resume

Format each as:
**Before:** [exact current text, copied verbatim]
**After:** [improved version — same facts, better expression]
**Why:** [one sentence]

## ATS FORMATTING FLAGS

Flag any structural issues that cause ATS parsing failures:
- Tables, text boxes, or multi-column layouts (ATS often garbles text in these)
- Headers/footers containing key information (most ATS ignore header/footer content)
- Non-standard section titles (e.g. "My Journey" instead of "Experience", "My Toolkit" instead of "Skills")
- Graphics or images in place of text
- Unparseable date formats
- **Missing Skills section** — a dedicated "Skills" or "Core Competencies" section is scanned separately by most ATS; flag if absent
- **Resume length** — flag if resume appears over 2 pages (too long gets skimmed) or under 1 full page (looks thin)

If none found, state: *No formatting issues detected.*

## IMPACT SCORE

Count bullet points in the Experience section. Categorise each as:
- **Results-focused:** contains a metric, outcome, or business impact ("reduced by 30%", "increased revenue", "launched", "delivered", "saved X hours")
- **Duty-focused:** describes a responsibility without measurable outcome ("responsible for", "managed", "assisted with", "involved in")

Report as: **X of Y bullets show measurable impact (Z%)** then list the top 3 duty-focused bullets that most need converting to results-focused, with a suggested rewrite.

## ATS KEYWORD OPTIMISATION
- **Must-add (hard filter risk):** exact terms the ATS almost certainly filters on
- **Should-add (ranking boost):** phrases that increase ranking

⚠️ KEYWORD EXACTNESS WARNING: ATS systems match literal strings. Flag every case where the resume uses a vague or paraphrased term instead of the job ad's exact wording.
- "SQL" ≠ "database querying" · "Python" ≠ "scripting" · "AWS" ≠ "cloud platforms" · "Tableau" ≠ "visualisation tools"
- **Acronym coverage:** for any key term that has both an acronym and a full form (e.g. "SEO / Search Engine Optimisation", "KPI / Key Performance Indicator", "ETL / Extract Transform Load"), flag if the resume includes only one form — ATS may index either, so both should appear at least once
- Flag missing semantic phrases where relevant: "end-to-end ownership", "data-driven decision making", "cross-functional collaboration", "stakeholder management", "scalable solutions", "production-ready systems", "business insights"
- Warn if the resume shows keyword stuffing — modern ATS platforms detect and penalise it

## COVER LETTER TALKING POINTS
Top 3 specific talking points for this role and company, focused on honestly bridging the most important implied or missing gaps.

Be direct. If the match is weak, say so and explain why.
"""

    try:
        message = client.messages.create(
            model=MODEL_NAME,
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}]
        )
        return {'success': True, 'analysis': message.content[0].text}
    except Exception as e:
        return {'success': False, 'error': str(e)}


def _extract_applicant_name(resume_text: str) -> str:
    """Pull the applicant's name from the top of the resume (usually the first non-empty line)."""
    for line in resume_text.splitlines():
        line = line.strip()
        if not line:
            continue
        # Skip lines that look like contact details, section headers, or dates
        if re.search(r'[@|/\\]|\d{4}|resume|curriculum|vitae|cv\b', line, re.IGNORECASE):
            continue
        # Likely a name: 2-4 words, mostly alpha, no punctuation beyond hyphens
        words = line.split()
        if 2 <= len(words) <= 5 and all(re.match(r"[A-Za-z\-'\.]+$", w) for w in words):
            return line
    return ""


def _extract_applicant_phone(resume_text: str) -> str:
    """Pull the first phone number found in the resume."""
    m = re.search(
        r'(\+?\d[\d\s\-().]{7,}\d)',
        resume_text
    )
    return m.group(1).strip() if m else ""


def _ensure_closing(letter: str, name: str, phone: str) -> str:
    """
    Guarantee the cover letter ends with a proper closing block.
    If one already exists (any variant of Regards/Sincerely/Yours), replace it
    with our canonical version. Otherwise append it.
    """
    closing_pat = re.compile(
        r'(kind regards|warm regards|regards|sincerely|yours sincerely|yours faithfully|'
        r'best regards|with regards|thank you)[,.]?\s*\n',
        re.IGNORECASE
    )
    # Build the closing block
    closing_lines = ["Kind regards,", ""]
    if name:
        closing_lines.append(name)
    if phone:
        closing_lines.append(phone)
    closing_block = "\n".join(closing_lines)

    # If a closing already exists, strip everything from it onward and replace
    m = closing_pat.search(letter)
    if m:
        return letter[:m.start()].rstrip() + "\n\n" + closing_block
    # No closing found — append it
    return letter.rstrip() + "\n\n" + closing_block


def generate_cover_letter(resume_text: str, job_content: str, job_url: str, analysis_text: str, client,
                          tone: str = "Professional", length: str = "Standard (300–350 words)",
                          incorporate_recs: bool = True, prior_letter: str = None,
                          change_instructions: str = None, user_guidance: str = None):
    word_map = {"Brief (≈200 words)": "approximately 200", "Standard (300–350 words)": "300–350", "Detailed (≈450 words)": "approximately 450"}
    word_target = word_map.get(length, "300–350")
    recs_line = (
        "4. Subtly weave in 1–2 of the specific improvement recommendations from the match analysis — show you understand what the role demands."
        if incorporate_recs else
        "4. Address any apparent gaps using transferable skills and relevant context."
    )

    applicant_name  = _extract_applicant_name(resume_text)
    applicant_phone = _extract_applicant_phone(resume_text)
    closing_name    = applicant_name  or "[Your Name]"
    closing_phone   = applicant_phone or ""
    closing_instruction = (
        f"IMPORTANT — the final lines of the letter MUST be exactly:\n\n"
        f"Kind regards,\n\n{closing_name}"
        + (f"\n{closing_phone}" if closing_phone else "")
    )

    if prior_letter and change_instructions:
        prompt = f"""You are an expert career coach and professional cover letter writer.

You previously wrote the cover letter below. The user has requested specific changes. Apply them faithfully while keeping everything else intact.

**CURRENT COVER LETTER:**
{prior_letter}

**REQUESTED CHANGES:**
{change_instructions}

**CONTEXT — MY RESUME:**
{resume_text}

**CONTEXT — JOB POSTING (from {job_url}):**
{job_content}

**Keep these settings:**
- Tone: {tone}
- Length: {word_target} words
- Do NOT include placeholder text or generic phrases.
- {closing_instruction}

Return only the revised cover letter, starting directly with the salutation.
"""
    else:
        prompt = f"""You are an expert career coach and professional cover letter writer.

Based on this resume analysis and job posting, write a compelling cover letter.

**MY RESUME:**
{resume_text}

**JOB POSTING (from {job_url}):**
{job_content}

**MATCH ANALYSIS:**
{analysis_text}

**Instructions:**
1. Tone: {tone} — maintain this voice consistently throughout
2. Length: {word_target} words (body only — not counting the closing lines)
3. Strong opening hook that immediately shows value — no "I am writing to apply" clichés
{recs_line}
5. Highlight 2–3 key achievements from the resume that directly match job requirements
6. Clear, confident call to action in the closing paragraph
7. Tailored to this specific job and company — reference the role and company by name
{f"8. Additional guidance from the user: {user_guidance}" if user_guidance and user_guidance.strip() else ""}

Do NOT include placeholder text or generic phrases.
Start directly with the salutation. Make it ready to copy-paste.
{closing_instruction}
"""
    try:
        message = client.messages.create(
            model=MODEL_NAME,
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )
        letter = message.content[0].text
        # Guarantee the closing is present and correct regardless of what Claude produced
        letter = _ensure_closing(letter, applicant_name, applicant_phone)
        return {'success': True, 'cover_letter': letter}
    except Exception as e:
        return {'success': False, 'error': str(e)}


# ============= EXCEL TRACKER =============

def parse_analysis_fields(analysis_text: str) -> dict:
    """Extract job title, company, location and match % from analysis text."""
    fields = {'job_title': '', 'company': '', 'location': '', 'match_pct': ''}

    score_match = re.search(r'COMPATIBILITY SCORE[^0-9]*(\d+)%', analysis_text, re.IGNORECASE)
    if score_match:
        fields['match_pct'] = score_match.group(1) + '%'

    title_match = re.search(r'\*{0,2}Job Title[:\*\s]+(.+)', analysis_text, re.IGNORECASE)
    if title_match:
        fields['job_title'] = title_match.group(1).strip().strip('*')

    company_match = re.search(r'\*{0,2}Company[:\*\s]+(.+)', analysis_text, re.IGNORECASE)
    if company_match:
        fields['company'] = company_match.group(1).strip().strip('*')

    location_match = re.search(r'\*{0,2}Location[:\*\s]+(.+)', analysis_text, re.IGNORECASE)
    if location_match:
        fields['location'] = location_match.group(1).strip().strip('*')

    return fields


def extract_recommendations_summary(analysis_text: str) -> str:
    """Pull the SPECIFIC IMPROVEMENT RECOMMENDATIONS section out of the analysis."""
    m = re.search(
        r'##\s*SPECIFIC IMPROVEMENT RECOMMENDATIONS\s*\n(.*?)(?=\n##|\Z)',
        analysis_text, re.DOTALL | re.IGNORECASE
    )
    if m:
        return m.group(1).strip()[:1500]
    return ""


def update_application_status(record_id, new_status: str):
    """Update the status of a single application row in Supabase."""
    sb = get_supabase()
    if sb and record_id:
        sb.table("applications").update({"status": new_status}).eq("id", record_id).execute()
        load_tracker_data.clear()


def save_to_tracker(job_title: str, company: str, location: str,
                    resume_filename: str, match_pct: str, job_url: str,
                    cover_letter: str = '', cover_letter_path: str = '',
                    notes: str = '', updated_resume_file: str = '',
                    user_id: str = None):
    """Insert a job application row into Supabase."""
    sb = get_supabase()
    row = {
        "date":                datetime.now().strftime('%Y-%m-%d'),
        "job_title":           job_title,
        "company":             company,
        "location":            location,
        "resume_file":         resume_filename,
        "updated_resume_file": updated_resume_file,
        "match_pct":           match_pct,
        "job_url":             job_url if job_url != 'Manual Input' else '',
        "status":              "Applied",
        "notes":               notes,
        "cover_letter":        cover_letter_path if cover_letter_path else cover_letter,
        "user_id":             user_id,
    }
    if not sb:
        raise RuntimeError("Supabase is not configured — cannot save application.")
    result = sb.table("applications").insert(row).execute()
    if hasattr(result, 'error') and result.error:
        raise RuntimeError(f"Supabase insert failed: {result.error}")


@st.cache_data(ttl=30)
def load_tracker_data(user_id: str = None):
    """Load applications from Supabase for the given user, newest first."""
    sb = get_supabase()
    if not sb:
        return []
    try:
        query = sb.table("applications").select("*").order("created_at", desc=True)
        if user_id:
            query = query.eq("user_id", user_id)
        rows = query.execute()
        return rows.data or []
    except Exception:
        return []


def generate_tracker_excel(tracker_data: list) -> bytes:
    """Build an in-memory Excel workbook from tracker_data (list of dicts)."""
    headers  = ['Date', 'Job Title', 'Company', 'Location', 'Original Resume', 'Updated Resume', 'Match %', 'Job URL', 'Status', 'Cover Letter']
    col_keys = ['date', 'job_title', 'company', 'location', 'resume_file', 'updated_resume_file', 'match_pct', 'job_url', 'status', 'cover_letter']

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Applications"

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font      = openpyxl.styles.Font(bold=True, color="FFFFFF")
        cell.fill      = openpyxl.styles.PatternFill(fill_type="solid", fgColor="230344")
        cell.alignment = openpyxl.styles.Alignment(horizontal="center")

    widths = [14, 30, 25, 20, 35, 35, 10, 45, 15, 30]
    for col, width in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = width

    for i, item in enumerate(tracker_data, 2):
        for col, key in enumerate(col_keys, 1):
            ws.cell(row=i, column=col, value=item.get(key) or '')
        if i % 2 == 0:
            fill = openpyxl.styles.PatternFill(fill_type="solid", fgColor="F3F0F8")
            for col in range(1, len(headers) + 1):
                ws.cell(row=i, column=col).fill = fill

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============= RESUME UPDATER =============

def generate_guidance_updates(resume_text: str, user_guidance: str, client):
    """
    Apply the candidate's OWN stated facts/metrics/tone instructions as locked find→replace pairs.
    These are treated as ground truth — not suggestions — and run before any AI improvements.
    """
    prompt = f"""You are a resume editor. The candidate has provided SPECIFIC facts or instructions below.
These are truths they are vouching for — your ONLY job is to weave them into the most relevant bullets.

CANDIDATE INSTRUCTIONS (must be applied exactly):
{user_guidance.strip()}

RESUME TEXT:
{resume_text}

Rules:
- Find the single most relevant existing bullet point for each instruction and incorporate it naturally
- Do NOT add new bullets; modify an existing one
- Do NOT fabricate anything beyond what the candidate stated
- "find" must be copied verbatim from the resume (including bullet prefix characters)
- Keep the same format/prefix as the original bullet
- Return ONLY a JSON array — no other text

```json
[
  {{
    "find": "exact text verbatim from resume",
    "replace": "improved version incorporating the candidate's stated fact",
    "description": "Applied candidate guidance: [their instruction]",
    "type": "user_guidance"
  }}
]
```"""

    try:
        resp = client.messages.create(
            model=MODEL_NAME,
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = resp.content[0].text.strip()
        # strip markdown fence if present
        if '```' in raw:
            raw = re.sub(r'```(?:json)?\s*', '', raw).strip().rstrip('`').strip()
        updates = json.loads(raw)
        for u in updates:
            u['type'] = 'user_guidance'
        return {'success': True, 'updates': updates}
    except Exception as e:
        return {'success': False, 'updates': [], 'error': str(e)}


def generate_resume_updates(resume_text: str, analysis_text: str, client):
    """
    Ask Claude to return targeted find→replace pairs based on recommendations.
    Strictly forbidden from fabricating skills or experience not in the original resume.
    User guidance is applied separately as a locked Step 0 in generate_guidance_updates().
    """

    prompt = f"""You are a professional resume editor with a strict honesty policy.

**CURRENT RESUME TEXT:**
{resume_text}

**ANALYSIS & RECOMMENDATIONS:**
{analysis_text}

Your task is to improve how existing experience is communicated — NOT to add experience that doesn't exist.

⚠️ HONESTY RULES (non-negotiable):
- NEVER add skills, tools, certifications, job titles, or responsibilities not already present in the resume
- NEVER invent metrics, numbers, or achievements that are not in either the resume OR the candidate guidance above
- NEVER add experience with software, frameworks, or industries not mentioned in the resume
- You MAY: use stronger action verbs, improve sentence structure, make achievements clearer, reframe existing content more powerfully, fix vague language
- PRIORITISE these action verbs wherever they naturally fit existing content: Built, Designed, Developed, Automated, Optimised, Implemented, Led, Delivered, Reduced, Increased, Migrated, Integrated — replace weak or passive verbs with these where accurate
- ✅ CANDIDATE GUIDANCE IS TRUTH: any metrics, numbers, team sizes, or outcomes the candidate has provided in the guidance section (above) are real facts they are vouching for. You MUST weave these into the relevant bullet points — find the closest matching bullet and incorporate the figure. Do not skip them.
- Where a bullet point is missing a quantified result and one could logically exist, note this in the description field (e.g. "Consider adding a metric here — what % improvement or time saving resulted?") so the user knows where to strengthen further
- CRITICAL — technical keyword matching: if the job ad uses specific tool/technology names (e.g. "SQL", "Python", "Tableau", "AWS"), the resume must use those EXACT words — not synonyms or paraphrases. "SQL" not "database querying". "Python" not "scripting". "AWS" not "cloud platforms". "Tableau" not "visualisation tools". Where the candidate clearly has the skill but has used a vague term, replace it with the exact keyword from the job ad
- SEMANTIC PHRASES — where naturally supported by existing experience, weave in AI-scoring phrases: "end-to-end ownership", "data-driven decision making", "cross-functional collaboration", "stakeholder management", "scalable solutions", "production-ready systems", "business insights" — only where they accurately describe what the candidate did
- AVOID keyword stuffing, repeating the same keyword multiple times, or adding keyword lists — modern ATS systems detect and downgrade this
- THE MAGIC FORMULA for a strong bullet: Action Verb + Exact Skill/Tool + Quantified Impact + Business Outcome. Example: "Built automated Tableau dashboards using SQL and Python, reducing manual reporting time by 40% and improving stakeholder decision-making speed." Apply this structure to existing bullets wherever possible
- If additional guidance is provided, follow it only where it aligns with existing resume content
- If a recommendation requires adding something the candidate clearly doesn't have, SKIP that change entirely

Return ONLY a JSON array. Each item must have:
- "find": exact text from the resume (copy verbatim, including bullet characters)
- "replace": the improved version (same factual content, better expression)
- "description": one short sentence explaining what changed and why (e.g. "Replaced weak verb 'managed' with 'led'" or "Made outcome explicit from existing context")

```json
[
  {{
    "find": "exact text copied verbatim from the resume",
    "replace": "improved version — same facts, better expression",
    "description": "Short explanation of the change"
  }}
]
```

Rules:
- "find" MUST appear exactly in the resume — copy it verbatim including bullet prefix characters
- Keep the same bullet prefix/format as the original
- Make 5–10 targeted changes — not a full rewrite
- Do NOT change: job titles, company names, dates, education details, section headers

Return only the JSON array, no other text."""

    def _sanitize_json(s: str) -> str:  # noqa: E301
        """Escape unescaped control characters inside JSON string values."""
        result = []
        in_string = False
        i = 0
        _escapes = {'\n': '\\n', '\r': '\\r', '\t': '\\t', '\b': '\\b', '\f': '\\f'}
        while i < len(s):
            c = s[i]
            if c == '\\' and in_string:
                result.append(c)
                i += 1
                if i < len(s):
                    result.append(s[i])
            elif c == '"':
                result.append(c)
                in_string = not in_string
            elif in_string and ord(c) < 0x20:
                result.append(_escapes.get(c, ''))
            else:
                result.append(c)
            i += 1
        return ''.join(result)

    try:
        message = client.messages.create(
            model=MODEL_NAME,
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}]
        )
        response_text = message.content[0].text
        json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
        json_str = json_match.group(1) if json_match else response_text.strip()
        try:
            updates = json.loads(json_str)
        except json.JSONDecodeError:
            updates = json.loads(_sanitize_json(json_str))
        return {'success': True, 'updates': updates}
    except Exception as e:
        return {'success': False, 'error': str(e)}


def generate_implied_to_exact_updates(resume_text: str, analysis_text: str, client):
    """
    Parse IMPLIED rows from the analysis keyword tables and generate find→replace
    pairs that swap vague resume language for the job posting's exact terminology.
    """
    # Extract all IMPLIED rows from the markdown tables
    implied_items = []
    for row in re.finditer(
        r'\|\s*([^|]+?)\s*\|\s*IMPLIED\s*\|\s*([^|]*?)\s*\|',
        analysis_text, re.IGNORECASE
    ):
        term = row.group(1).strip()
        note = row.group(2).strip()
        if term and term not in ('Requirement', 'Skill / Tool', 'Qualification', 'Competency'):
            implied_items.append({"term": term, "note": note})

    if not implied_items:
        return {'success': True, 'updates': [], 'message': 'No IMPLIED matches found in the analysis.'}

    implied_list = "\n".join(
        f"- Exact term needed: \"{i['term']}\" | Resume currently implies it via: {i['note']}"
        for i in implied_items
    )

    prompt = f"""You are a precise resume editor. Your only job is to convert IMPLIED keyword matches into EXACT matches.

The analysis has identified these terms as IMPLIED — the candidate has the experience but the resume uses vague or paraphrased language instead of the job posting's exact terminology:

{implied_list}

RESUME TEXT:
{resume_text}

For each IMPLIED item, find the specific phrase in the resume that implies this skill and replace it with the job posting's exact term — keeping everything else identical.

Rules:
- "find" MUST be copied verbatim from the resume (including bullet prefix characters)
- Only change the terminology — do NOT rewrite the whole bullet
- Do NOT add any skill or experience that isn't already in the resume
- If you cannot find a clear phrase to change for an item, skip it
- Each "find" must be unique — no duplicates

Return ONLY a JSON array:
```json
[
  {{
    "find": "exact phrase from resume, copied verbatim",
    "replace": "same phrase with the exact job posting terminology substituted in",
    "description": "IMPLIED → EXACT: '[vague term]' → '[exact term]'"
  }}
]
```

Return only the JSON array, no other text."""

    try:
        message = client.messages.create(
            model=MODEL_NAME,
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )
        response_text = message.content[0].text.strip()
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', response_text)
        json_str = json_match.group(1) if json_match else response_text
        updates = json.loads(json_str)
        return {'success': True, 'updates': updates, 'implied_count': len(implied_items)}
    except Exception as e:
        return {'success': False, 'error': str(e)}


def generate_removal_suggestions(resume_text: str, client) -> dict:
    """
    Identify unnecessary content to remove from the resume.
    Returns find→replace pairs where replace='' means delete the line.
    """
    prompt = f"""You are a ruthless resume editor. Identify content in the resume below that should be REMOVED to make it tighter and stronger.

RESUME:
{resume_text}

Find lines or bullet points that fall into these categories:

1. DUTY-BASED BULLETS — bullets that only describe a responsibility with no outcome or metric.
   Examples of phrases to flag: "Responsible for", "Assisted with", "Helped", "Involved in", "Supported", "Worked on", "Participated in", "Contributed to"

2. DUPLICATES — the same skill or responsibility mentioned more than once across different roles. Keep it in the most recent role only; flag older duplicates for removal.

3. FILLER LINES — "References available on request", hobbies/interests section content, generic objective statements with no specific value.

4. OUTDATED ROLE BULLETS — if a role is 10+ years old and has more than 2 bullets, flag the weakest ones (duty-based, no metrics, outdated tech).

5. VERBOSE PADDING — standalone sentences that restate something already said more concisely elsewhere.

For each item you find, return:
- "find": the EXACT text as it appears in the resume (copy it verbatim including bullet prefix characters)
- "replace": "" (empty string — meaning delete this line)
- "description": one short sentence starting with "REMOVE:" explaining why (e.g. "REMOVE: duty-based bullet with no outcome in 2012 role")
- "type": "remove"

Return ONLY a JSON array — no other text:
```json
[
  {{
    "find": "exact text from resume copied verbatim",
    "replace": "",
    "description": "REMOVE: reason",
    "type": "remove"
  }}
]
```

Rules:
- "find" MUST be copied verbatim from the resume — it must match exactly
- Only flag lines you are confident should be removed — be selective, not exhaustive
- Do NOT flag job titles, company names, dates, or qualifications
- Maximum 8 removal suggestions — prioritise the highest-impact cuts"""

    try:
        message = client.messages.create(
            model=MODEL_NAME,
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )
        response_text = message.content[0].text.strip()
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', response_text)
        json_str = json_match.group(1) if json_match else response_text
        updates = json.loads(json_str)
        # Ensure all removal items have type and empty replace
        for u in updates:
            u['type']    = 'remove'
            u['replace'] = ''
        return {'success': True, 'updates': updates}
    except Exception as e:
        return {'success': False, 'error': str(e)}

def estimate_pages(word_count: int) -> float:
    """Rough page estimate: ~400 resume words per page (accounting for whitespace/headers)."""
    return round(word_count / 400, 1)


def trim_resume(resume_text: str, client, target_pages: int = 2) -> dict:
    """
    Ask Claude to identify what to cut/condense, returning find→replace pairs
    that can be applied to the original DOCX to preserve formatting.
    replace='' means delete the line entirely; non-empty replace means condense.
    """
    current_words = len(resume_text.split())
    max_words     = target_pages * 550
    must_cut      = max(0, current_words - max_words)

    prompt = f"""You are a ruthless resume editor. Identify what to CUT or CONDENSE from the resume below to bring it to {target_pages} pages.

CURRENT LENGTH: {current_words} words  |  TARGET MAX: {max_words} words  |  MUST CUT: {must_cut} words

RESUME:
{resume_text}

Work through in priority order:
1. DELETE entirely: Hobbies/Interests, "References available on request", generic objective statements
2. CONDENSE old roles (10+ years ago): reduce to title/company/dates + ONE best bullet — return the multi-bullet block as find with a single-bullet replacement
3. DELETE duty-based bullets with no outcome ("Responsible for", "Assisted with", "Helped", "Involved in", "Supported")
4. DELETE duplicate bullets that appear in multiple roles (keep only the most recent instance)
5. CONDENSE verbose bullets: tighten to one line, remove filler words

Return ONLY a JSON array — no other text:
```json
[
  {{
    "find": "exact text from resume copied verbatim (including bullet symbols if present)",
    "replace": "",
    "description": "REMOVE: reason"
  }},
  {{
    "find": "exact multi-line block to replace",
    "replace": "condensed single line",
    "description": "CONDENSE: reason"
  }}
]
```

Rules:
- "find" MUST be verbatim from the resume — it must match exactly including spacing and punctuation
- NEVER delete job titles, company names, date ranges, qualifications, or results-based bullets (those with %, $, numbers)
- Aim for {must_cut}+ words removed across all changes
- Maximum 12 items — prioritise highest-word-count cuts first"""

    try:
        message = client.messages.create(
            model=MODEL_NAME,
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = message.content[0].text.strip()
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', raw)
        json_str = json_match.group(1) if json_match else raw
        pairs = json.loads(json_str)
        # Tag all as trim type
        for p in pairs:
            p['type'] = 'trim'
        return {'success': True, 'pairs': pairs}
    except Exception as e:
        return {'success': False, 'error': str(e)}


def build_plain_docx(text: str, filename_hint: str = "resume") -> bytes:
    """Convert plain trimmed resume text into a compact, tight DOCX (no Word default bloat)."""
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement

    def _set_line_spacing(paragraph, lines=1.0):
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = 0   # exact
        pf.line_spacing      = Pt(13)

    def _zero_space(paragraph, before=0, after=2):
        pf = paragraph.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after  = Pt(after)

    doc = Document()

    # Narrow margins
    for sec in doc.sections:
        sec.top_margin    = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin   = Inches(0.75)
        sec.right_margin  = Inches(0.75)

    # Kill default Normal style spacing
    normal = doc.styles['Normal']
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)
    normal.font.size = Pt(10.5)

    blank_count = 0   # collapse multiple blank lines into one
    for line in text.splitlines():
        stripped = line.strip()

        if not stripped:
            blank_count += 1
            if blank_count == 1:          # allow single blank separator
                p = doc.add_paragraph("")
                _zero_space(p, before=0, after=0)
            continue
        blank_count = 0

        # Section heading: ALL CAPS line, 4+ chars
        if stripped.isupper() and len(stripped) >= 4 and not stripped.startswith(('•', '-', '·', '*', '–')):
            p = doc.add_paragraph()
            _zero_space(p, before=6, after=2)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
        # Bullet line
        elif stripped[0] in ('•', '-', '·', '*', '–'):
            p = doc.add_paragraph()
            _zero_space(p, before=0, after=1)
            _set_line_spacing(p)
            p.paragraph_format.left_indent   = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            run = p.add_run("• " + stripped[1:].strip())
            run.font.size = Pt(10.5)
        # Normal text
        else:
            p = doc.add_paragraph()
            _zero_space(p, before=0, after=2)
            _set_line_spacing(p)
            run = p.add_run(stripped)
            run.font.size = Pt(10.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def render_analysis(analysis_text: str):
    """Render the analysis with styled keyword tables for the EXACT/IMPLIED/MISSING section."""

    STATUS_STYLE = {
        'exact':   ('✅ EXACT',   '#7ad79f', 'rgba(122,215,159,0.12)'),
        'implied': ('🔶 IMPLIED', '#e0a14a', 'rgba(224,161,74,0.12)'),
        'missing': ('❌ MISSING', '#ef4444', 'rgba(239,68,68,0.10)'),
    }

    def _status_badge(cell: str) -> str:
        c = cell.strip().upper()
        for key, (label, color, bg) in STATUS_STYLE.items():
            if key in c.lower():
                return (
                    f'<span style="font-family:\'Space Mono\',monospace;font-size:10px;'
                    f'font-weight:700;letter-spacing:0.06em;color:{color};'
                    f'background:{bg};border-radius:5px;padding:2px 8px;white-space:nowrap;">'
                    f'{label}</span>'
                )
        return cell.strip()

    def _parse_md_table(block: str) -> list[list[str]]:
        rows = []
        for line in block.splitlines():
            line = line.strip()
            if not line.startswith('|') or re.match(r'^\|[-| ]+\|$', line):
                continue
            cells = [c.strip() for c in line.strip('|').split('|')]
            rows.append(cells)
        return rows

    def _render_subsection_table(heading: str, table_rows: list[list[str]]) -> str:
        if not table_rows:
            return ''
        header = table_rows[0]
        body   = table_rows[1:]
        th_cells = ''.join(
            f'<th style="padding:4px 10px;text-align:left;font-family:\'Space Mono\',monospace;'
            f'font-size:9px;letter-spacing:0.10em;text-transform:uppercase;color:#6e8a7b;'
            f'border-bottom:1px solid rgba(159,182,168,0.15);white-space:nowrap;">{h}</th>'
            for h in header
        )
        tr_rows = ''
        for i, row in enumerate(body):
            bg = 'rgba(255,255,255,0.02)' if i % 2 == 0 else 'transparent'
            tds = ''
            for j, cell in enumerate(row):
                if j == 1:  # Status column
                    content = _status_badge(cell)
                else:
                    content = f'<span style="font-family:\'DM Sans\',sans-serif;font-size:12px;color:#ecf4ee;">{cell}</span>'
                tds += (
                    f'<td style="padding:4px 10px;vertical-align:top;'
                    f'border-bottom:1px solid rgba(159,182,168,0.07);background:{bg};">'
                    f'{content}</td>'
                )
            tr_rows += f'<tr>{tds}</tr>'
        return (
            f'<div style="margin:0.4rem 0 0.2rem;">'
            f'<p style="font-family:\'Space Mono\',monospace;font-size:9px;letter-spacing:0.14em;'
            f'text-transform:uppercase;color:#7ad79f;margin:0 0 0.3rem;">{heading}</p>'
            f'<div style="overflow-x:auto;border-radius:8px;border:1px solid rgba(159,182,168,0.12);">'
            f'<table style="width:100%;border-collapse:collapse;">'
            f'<thead><tr>{th_cells}</tr></thead>'
            f'<tbody>{tr_rows}</tbody>'
            f'</table></div></div>'
        )

    # Split the analysis into sections
    section_pat = re.compile(r'(^##\s+.+$)', re.MULTILINE)
    parts = section_pat.split(analysis_text)

    i = 0
    while i < len(parts):
        chunk = parts[i]

        # Section heading
        if re.match(r'^##\s+', chunk):
            heading_text = chunk.strip()
            # Check if the NEXT chunk contains the keyword tables
            content = parts[i + 1] if i + 1 < len(parts) else ''
            if 'KEYWORD EXTRACTION' in heading_text.upper():
                heading_label = heading_text.lstrip('#').strip()
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:8px;margin:0.8rem 0 0.3rem;">'
                    f'<span style="font-family:\'Space Mono\',monospace;font-size:9.5px;letter-spacing:0.16em;'
                    f'text-transform:uppercase;color:#7ad79f;white-space:nowrap;">{heading_label}</span>'
                    f'<span style="flex:1;height:1px;background:rgba(159,182,168,0.15);"></span>'
                    f'</div>',
                    unsafe_allow_html=True
                )
                # Parse subsections (### headings) within this content block
                sub_pat = re.compile(r'(^###\s+.+$)', re.MULTILINE)
                sub_parts = sub_pat.split(content)
                j = 0
                # Any leading text before first ###
                if sub_parts[0].strip():
                    st.markdown(sub_parts[0])
                j = 1
                while j < len(sub_parts):
                    sub_heading = sub_parts[j].strip('# ').strip()
                    sub_content = sub_parts[j + 1] if j + 1 < len(sub_parts) else ''
                    rows = _parse_md_table(sub_content)
                    if rows:
                        html = _render_subsection_table(sub_heading, rows)
                        st.markdown(html, unsafe_allow_html=True)
                        # Render any non-table text after the table in this subsection
                        non_table = re.sub(r'\|.*\|', '', sub_content).strip()
                        if non_table:
                            st.markdown(non_table)
                    else:
                        st.markdown(f'### {sub_heading}\n{sub_content}')
                    j += 2
                i += 2
                continue
            else:
                heading_label = heading_text.lstrip('#').strip()
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:8px;margin:0.8rem 0 0.25rem;">'
                    f'<span style="font-family:\'Space Mono\',monospace;font-size:9.5px;letter-spacing:0.16em;'
                    f'text-transform:uppercase;color:#7ad79f;white-space:nowrap;">{heading_label}</span>'
                    f'<span style="flex:1;height:1px;background:rgba(159,182,168,0.15);"></span>'
                    f'</div>',
                    unsafe_allow_html=True
                )
                if content.strip():
                    st.markdown(content)
                i += 2
                continue

        # Plain text before first heading
        if chunk.strip():
            st.markdown(chunk)
        i += 1


def parse_keyword_tables(analysis_text: str) -> dict:
    """
    Parse the four EXACT/IMPLIED/MISSING keyword tables from the analysis into
    structured data: {'hard': [...], 'core': [...], 'preferred': [...], 'soft': [...]}
    Each item: {'term': str, 'status': 'exact'|'implied'|'missing'}
    """
    CATEGORIES = [
        ('hard',      re.compile(r'###\s*Hard\s+Req',    re.I)),
        ('core',      re.compile(r'###\s*Core\s+Skill',  re.I)),
        ('preferred', re.compile(r'###\s*Preferred',     re.I)),
        ('soft',      re.compile(r'###\s*Soft\s+Skill',  re.I)),
    ]
    SKIP_HEADERS = {'requirement', 'skill / tool', 'skill/tool', 'qualification', 'competency', 'skill', 'competencies'}

    result = {k: [] for k in ('hard', 'core', 'preferred', 'soft')}

    sub_pat = re.compile(r'^###\s+.+$', re.MULTILINE)
    headings = sub_pat.findall(analysis_text)
    sections = sub_pat.split(analysis_text)

    for heading, content in zip(headings, sections[1:]):
        cat = next((k for k, p in CATEGORIES if p.search(heading)), None)
        if cat is None:
            continue
        for line in content.splitlines():
            line = line.strip()
            if not line.startswith('|') or re.match(r'^\|[-| ]+\|$', line):
                continue
            cells = [c.strip() for c in line.strip('|').split('|')]
            if len(cells) < 2 or cells[0].lower() in SKIP_HEADERS:
                continue
            term = cells[0]
            raw_status = cells[1].lower() if len(cells) > 1 else ''
            if 'exact' in raw_status:
                status = 'exact'
            elif 'implied' in raw_status:
                status = 'implied'
            elif 'missing' in raw_status:
                status = 'missing'
            else:
                continue
            result[cat].append({'term': term, 'status': status})

    return result


def compute_score_analytically(table_data: dict, n_promote: int = 0) -> int | None:
    """
    Compute weighted compatibility score from keyword table data.
    EXACT=1.0, IMPLIED=0.5, MISSING=0.0
    n_promote: treat this many IMPLIED items as EXACT (applied by boost).
    Promoted items are taken from the most-weighted categories first.
    """
    WEIGHTS = {'hard': 0.50, 'core': 0.30, 'preferred': 0.15, 'soft': 0.05}
    POINTS  = {'exact': 1.0, 'implied': 0.5, 'missing': 0.0}
    CAT_ORDER = ['hard', 'core', 'preferred', 'soft']

    # Build a mutable copy with promotion applied
    promoted_left = n_promote
    promoted = {cat: [] for cat in CAT_ORDER}
    for cat in CAT_ORDER:
        for item in table_data.get(cat, []):
            if item['status'] == 'implied' and promoted_left > 0:
                promoted[cat].append('exact')
                promoted_left -= 1
            else:
                promoted[cat].append(item['status'])

    total_score = 0.0
    has_data = False
    for cat in CAT_ORDER:
        items = promoted[cat]
        if not items:
            continue
        has_data = True
        cat_score = sum(POINTS.get(s, 0.0) for s in items) / len(items)
        total_score += WEIGHTS[cat] * cat_score

    return round(total_score * 100) if has_data else None


def re_score_resume(updated_text: str, job_content: str, client,
                    orig_score: int = None, changes: list = None) -> int | None:
    """LLM re-score for general resume improvements. Returns integer % or None."""
    context = ""
    if orig_score is not None:
        context = f"The previous compatibility score was {orig_score}%. Calibrate relative to that.\n\n"
    if changes:
        descs = [c.get('description', '') for c in changes[:8] if c.get('description')]
        if descs:
            context += "Changes made to the resume:\n" + "\n".join(f"- {d}" for d in descs) + "\n\n"
    try:
        message = client.messages.create(
            model=MODEL_NAME,
            max_tokens=10,
            messages=[{"role": "user", "content":
                f"{context}"
                f"Re-score this updated resume against the job posting using this exact weighting:\n"
                f"- Hard requirements (degree, certs, years exp, 'required'/'must have'): 50%\n"
                f"- Core skills and tools explicitly named: 30%\n"
                f"- Preferred/nice-to-have qualifications: 15%\n"
                f"- Soft skills and culture fit: 5%\n"
                f"EXACT keyword matches score higher than IMPLIED. "
                f"Return ONLY a single integer 0-100. No other text.\n\n"
                f"RESUME:\n{updated_text[:4000]}\n\nJOB POSTING:\n{job_content[:4000]}"}]
        )
        val = re.search(r'\d+', message.content[0].text.strip())
        return min(int(val.group()), 100) if val else None
    except Exception:
        return None


def apply_updates_to_docx(file_bytes: bytes, updates: list, original_filename: str) -> bytes:
    """
    Apply find→replace pairs to the original DOCX, preserving all formatting.
    Searches both body paragraphs and table cells (many templates use tables for layout).
    Returns the updated DOCX as bytes.
    """
    doc = Document(io.BytesIO(file_bytes))
    applied = 0

    def _all_paragraphs(doc):
        yield from doc.paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    def _replace_in_para(para, find_text, replace_text):
        # Best case: find_text lives entirely within one run — replace only that run,
        # leaving every other run's formatting completely untouched.
        for run in para.runs:
            if find_text in run.text:
                run.text = run.text.replace(find_text, replace_text)
                return True

        # Fallback: text spans multiple runs — concatenate, replace, rebuild.
        if find_text not in para.text:
            return False

        new_text = para.text.replace(find_text, replace_text)

        # Pick the last non-empty content run as the formatting reference so we
        # inherit colour/size from the actual text, not from a structural run[0].
        ref_run = next(
            (r for r in reversed(para.runs) if r.text.strip()),
            para.runs[0] if para.runs else None,
        )

        if para.runs:
            tgt = para.runs[0]
            tgt.text = new_text
            if ref_run and ref_run is not tgt:
                try:
                    if ref_run.font.color.rgb is not None:
                        tgt.font.color.rgb = ref_run.font.color.rgb
                except Exception:
                    pass
                if ref_run.font.size:
                    tgt.font.size = ref_run.font.size
                if ref_run.font.name:
                    tgt.font.name = ref_run.font.name
                tgt.bold = ref_run.bold
                tgt.italic = ref_run.italic
            for run in para.runs[1:]:
                run.text = ''
        return True

    for update in updates:
        find_text    = update.get('find', '').strip()
        replace_text = update.get('replace', '').strip()
        if not find_text:
            continue

        for para in _all_paragraphs(doc):
            if find_text in para.text:
                if replace_text == '':
                    # Deletion — remove the entire paragraph element
                    try:
                        p_elem = para._element
                        p_elem.getparent().remove(p_elem)
                        applied += 1
                    except Exception:
                        pass
                else:
                    if _replace_in_para(para, find_text, replace_text):
                        applied += 1
                break

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue(), applied


# ============= DOCX BUILDERS =============

def _add_md_run(para, text: str):
    """Add runs to a paragraph, handling **bold** markers."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part:
            para.add_run(part)


def _add_md_table(doc, table_lines: list):
    """Render pipe-delimited markdown table rows as a Word table."""
    rows = []
    for line in table_lines:
        if re.match(r'^\|[\s\-|:]+\|$', line.strip()):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if cells:
            rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ''
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
            cell = table.rows[i].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(clean)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True


def create_cover_letter_docx(cover_letter_text: str) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(1.1)
    sec.bottom_margin = Inches(1.1)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for line in cover_letter_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0 if not line.strip() else 6)
        _add_md_run(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def create_analysis_docx(analysis_text: str, job_url: str, resume_filename: str, job_content: str = '') -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1.15)
    sec.right_margin  = Inches(1.15)

    # Title block
    title = doc.add_heading('ResumeSync — Compatibility Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta = doc.add_paragraph()
    meta.add_run('Resume: ').bold = True
    meta.add_run(resume_filename)
    meta2 = doc.add_paragraph()
    meta2.add_run('Job: ').bold = True
    meta2.add_run(job_url)
    meta3 = doc.add_paragraph()
    meta3.add_run('Date: ').bold = True
    meta3.add_run(datetime.now().strftime('%Y-%m-%d'))
    doc.add_paragraph()

    # Job description section
    if job_content and job_content.strip():
        doc.add_heading('Job Description', level=1)
        for para in job_content.strip().split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
            else:
                doc.add_paragraph()
        doc.add_page_break()

    doc.add_heading('Analysis', level=1)
    doc.add_paragraph()

    lines = analysis_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)

        # Table block — collect consecutive pipe lines
        elif line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            _add_md_table(doc, table_lines)
            doc.add_paragraph()
            continue

        # Bullet list
        elif re.match(r'^[-*]\s', line):
            p = doc.add_paragraph(style='List Bullet')
            _add_md_run(p, line[2:].strip())

        # Numbered list
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            _add_md_run(p, re.sub(r'^\d+\.\s', '', line))

        # Empty line
        elif not line.strip():
            doc.add_paragraph()

        # Normal paragraph (with inline bold)
        else:
            p = doc.add_paragraph()
            _add_md_run(p, line)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============= LOGO =============

@st.cache_data
def get_vp_logo_b64():
    from PIL import Image
    import numpy as np
    logo_path = os.path.join(os.path.dirname(__file__), "assets", "logo.png")
    if not os.path.exists(logo_path):
        return None
    img = Image.open(logo_path).convert("RGB")
    # Resize to sidebar-friendly width
    w, h = img.size
    new_w = 200
    img = img.resize((new_w, int(h * new_w / w)), Image.LANCZOS)
    data = np.array(img)
    r, g, b = data[:, :, 0], data[:, :, 1], data[:, :, 2]
    # Purple background mask: #230344 ≈ R<80, G<30, B>30
    mask = (r < 80) & (g < 30) & (b > 30)
    # Replace background with sidebar green #0a1b16
    data[mask, 0] = 8
    data[mask, 1] = 32
    data[mask, 2] = 25
    # Gold/bronze icon mask: warm tone with high R, moderate G, low B
    r2, g2, b2 = data[:, :, 0], data[:, :, 1], data[:, :, 2]
    gold_mask = (r2 > 100) & (g2 > 60) & (b2 < 120) & (r2 > b2) & (~mask)
    # Replace gold with accent green #7ad79f
    data[gold_mask, 0] = 109
    data[gold_mask, 1] = 193
    data[gold_mask, 2] = 138
    result = Image.fromarray(data.astype(np.uint8))
    buf = io.BytesIO()
    result.save(buf, format="PNG")
    buf.seek(0)
    return base64.b64encode(buf.getvalue()).decode()


# ============= LANDING PAGE =============

def show_landing():
    """Marketing landing page — full-page iframe with sidebar trigger for navigation."""

    # Hide all Streamlit chrome and set dark background to match design
    st.markdown("""
<style>
section[data-testid="stSidebar"]{display:none!important;}
[data-testid="collapsedControl"]{display:none!important;}
header[data-testid="stHeader"]{display:none!important;}
#MainMenu{display:none!important;}
footer[data-testid="stFooter"]{display:none!important;}
.stApp,[data-testid="stApp"]{background:#071512!important;overflow:hidden!important;}
[data-testid="stAppViewContainer"],.stAppViewContainer{background:#071512!important;height:100vh!important;overflow:hidden!important;}
section.main,.main{background:#071512!important;padding:0!important;height:100vh!important;overflow:hidden!important;}
.block-container,[data-testid="stMainBlockContainer"]{background:#071512!important;padding:0!important;max-width:100%!important;margin:0!important;height:100vh!important;overflow:hidden!important;}
[data-testid="stVerticalBlock"]{height:100vh!important;gap:0!important;overflow:hidden!important;}
iframe[title="st.iframe"]{height:100vh!important;width:100vw!important;border:none!important;display:block!important;}
</style>""", unsafe_allow_html=True)

    # Hidden login trigger in sidebar (hidden by display:none CSS above).
    # Iframe JS finds this button by text and clicks it to trigger a Python-side rerun.
    with st.sidebar:
        if st.button("__rs_login__", key="rs_login_trigger"):
            st.session_state.show_login = True
            st.rerun()

    # Full landing page in isolated iframe — no Streamlit container constraints.
    _landing_doc = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>ResumeSync — know your match before you apply</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Bricolage+Grotesque:opsz,wght@12..96,400;12..96,600;12..96,700;12..96,800&family=Lora:ital,wght@1,500;1,600&family=Space+Mono:wght@400;700&family=DM+Sans:opsz,wght@9..40,300;9..40,400;9..40,500&display=swap" rel="stylesheet">
<style>
:root{
--bg:#071512;--bg-2:#0a1b16;--panel:#0c2019;--panel-2:#0f271e;
--ink:#ecf4ee;--ink2:#9fb6a8;--ink3:#6e8a7b;
--mint:#7ad79f;--mint-bright:#94e6b1;--mint-deep:#4fae7a;
--mint-dim:rgba(122,215,159,0.10);--mint-dim2:rgba(122,215,159,0.16);
--mint-border:rgba(122,215,159,0.22);--amber:#e0a14a;
--line:rgba(159,182,168,0.12);--line-2:rgba(159,182,168,0.07);
--shadow-lg:0 24px 60px rgba(0,0,0,0.45),0 4px 14px rgba(0,0,0,0.35);
--mono:'Space Mono',ui-monospace,monospace;
--display:'Bricolage Grotesque',system-ui,sans-serif;
--serif:'Lora',Georgia,serif;--sans:'DM Sans',system-ui,sans-serif;
}
*{box-sizing:border-box;margin:0;padding:0;}
html{scroll-behavior:smooth;}
body{font-family:var(--sans);background:var(--bg);color:var(--ink);min-height:100vh;overflow-x:hidden;-webkit-font-smoothing:antialiased;position:relative;}
body::before{content:'';position:fixed;inset:0;z-index:0;pointer-events:none;background:radial-gradient(900px 600px at 78% -8%,rgba(122,215,159,0.10),transparent 60%),radial-gradient(700px 700px at 6% 4%,rgba(122,215,159,0.05),transparent 55%);}
body>*{position:relative;z-index:1;}
.eyebrow{font-family:var(--mono);font-size:12px;font-weight:400;letter-spacing:0.22em;text-transform:uppercase;color:var(--mint);display:inline-flex;align-items:center;gap:14px;}
.eyebrow::before,.eyebrow::after{content:'';width:26px;height:1px;background:var(--mint);opacity:0.5;}
.eyebrow.one::after{display:none;}
.nav{position:fixed;top:0;left:0;right:0;z-index:9999;display:flex;align-items:center;justify-content:space-between;padding:18px 56px;background:rgba(7,21,18,0.92);backdrop-filter:blur(14px);border-bottom:1px solid var(--line-2);}
body{padding-top:74px;}
.logo{display:flex;align-items:center;gap:12px;text-decoration:none;}
.logo-mark{width:38px;height:38px;border-radius:10px;background:linear-gradient(150deg,var(--mint),var(--mint-deep));display:grid;place-items:center;font-family:var(--display);font-size:18px;font-weight:800;color:#06140f;box-shadow:0 4px 14px rgba(122,215,159,0.30);}
.logo-txt{display:flex;flex-direction:column;line-height:1;}
.logo-name{font-family:var(--display);font-size:19px;font-weight:700;letter-spacing:-0.02em;color:var(--ink);}
.logo-name span{color:var(--mint);}
.logo-by{font-family:var(--mono);font-size:9.5px;letter-spacing:0.14em;text-transform:uppercase;color:var(--ink3);margin-top:3px;}
.nav-links{display:flex;align-items:center;gap:30px;}
.nav-link{font-size:14px;color:var(--ink2);text-decoration:none;font-weight:500;transition:color .15s;}
.nav-link:hover{color:var(--ink);}
.nav-cta{font-size:13.5px;font-weight:600;padding:10px 20px;border-radius:9px;background:var(--mint);color:#06140f;text-decoration:none;transition:background .15s,transform .15s;box-shadow:0 4px 14px rgba(122,215,159,0.22);}
.nav-cta:hover{background:var(--mint-bright);transform:translateY(-1px);}
.hero{padding:78px 56px 96px;max-width:1280px;margin:0 auto;display:grid;grid-template-columns:1fr 472px;gap:60px;align-items:center;}
h1{font-family:var(--display);font-weight:800;font-size:clamp(38px,4.1vw,60px);line-height:1.0;letter-spacing:-0.03em;color:var(--ink);margin:22px 0 24px;}
h1 .italic{font-family:var(--serif);font-weight:600;font-style:italic;color:var(--mint);letter-spacing:-0.01em;}
.hero-lede{font-size:18px;line-height:1.66;color:var(--ink2);max-width:50ch;margin-bottom:36px;}
.hero-lede strong{color:var(--ink);font-weight:500;}
.hero-btns{display:flex;align-items:center;gap:14px;flex-wrap:wrap;margin-bottom:34px;}
.btn-primary{font-size:15px;font-weight:600;padding:15px 28px;border-radius:11px;background:var(--mint);color:#06140f;text-decoration:none;box-shadow:0 10px 30px rgba(122,215,159,0.26);transition:background .15s,transform .15s;display:inline-block;}
.btn-primary:hover{background:var(--mint-bright);transform:translateY(-1px);}
.btn-ghost{font-size:15px;font-weight:500;padding:15px 26px;border-radius:11px;border:1px solid var(--line);color:var(--ink);text-decoration:none;background:rgba(255,255,255,0.015);transition:border-color .15s,background .15s;display:inline-block;}
.btn-ghost:hover{border-color:var(--mint-border);background:var(--mint-dim);}
.hero-social{display:flex;align-items:center;gap:18px;flex-wrap:wrap;}
.social-item{display:flex;align-items:center;gap:8px;font-size:13px;color:var(--ink3);}
.social-item b{color:var(--ink2);font-weight:600;}
.social-item .tick{color:var(--mint);}
.social-div{width:1px;height:14px;background:var(--line);}
.hero-visual{position:relative;}
.photo-frame{position:relative;border-radius:20px;overflow:hidden;border:1px solid var(--line);box-shadow:var(--shadow-lg);aspect-ratio:5/6;background:#0b1e18 center/cover no-repeat;}
.photo-frame img{position:absolute;inset:0;width:100%;height:100%;object-fit:cover;display:block;}
.photo-frame::after{content:'';position:absolute;inset:0;pointer-events:none;background:linear-gradient(to top,rgba(6,17,13,0.92) 0%,rgba(6,17,13,0.55) 22%,rgba(6,17,13,0) 46%);z-index:2;}
.photo-badge{position:absolute;top:18px;left:18px;z-index:3;font-family:var(--mono);font-size:10.5px;font-weight:700;letter-spacing:0.14em;text-transform:uppercase;padding:7px 12px;border-radius:7px;background:rgba(6,17,13,0.78);color:var(--mint);border:1px solid var(--mint-border);backdrop-filter:blur(4px);display:inline-flex;align-items:center;gap:8px;}
.pulse{width:6px;height:6px;border-radius:50%;background:var(--mint);animation:pulse 2.2s infinite;}
@keyframes pulse{0%{box-shadow:0 0 0 0 rgba(122,215,159,0.5);}70%{box-shadow:0 0 0 8px rgba(122,215,159,0);}100%{box-shadow:0 0 0 0 rgba(122,215,159,0);}}
.photo-name{position:absolute;left:22px;bottom:20px;z-index:3;}
.photo-name .nm{font-family:var(--display);font-size:21px;font-weight:700;color:#fff;letter-spacing:-0.01em;}
.photo-name .role{font-family:var(--mono);font-size:10.5px;letter-spacing:0.12em;text-transform:uppercase;color:var(--mint);margin-top:6px;}
.float{position:absolute;z-index:5;background:rgba(10,27,21,0.82);border:1px solid var(--mint-border);border-radius:16px;box-shadow:0 12px 30px rgba(0,0,0,0.30);backdrop-filter:blur(10px);}
.float-score{top:-44px;right:-34px;padding:18px 20px 16px;width:172px;text-align:center;}
.fs-label{font-family:var(--mono);font-size:9.5px;letter-spacing:0.16em;text-transform:uppercase;color:var(--ink3);margin-bottom:10px;}
.ring-wrap{position:relative;width:108px;height:108px;margin:0 auto;}
.ring-wrap svg{transform:rotate(-90deg);}
.ring-num{position:absolute;inset:0;display:grid;place-items:center;font-family:var(--display);font-weight:800;font-size:30px;color:var(--mint);}
.ring-num small{font-size:14px;color:var(--ink2);margin-left:1px;}
.fs-fit{font-family:var(--mono);font-size:10px;letter-spacing:0.14em;text-transform:uppercase;color:var(--mint);margin-top:10px;}
.float-skills{bottom:64px;right:-44px;padding:16px 18px;width:228px;}
.fsk-label{font-family:var(--mono);font-size:9.5px;letter-spacing:0.16em;text-transform:uppercase;color:var(--ink3);margin-bottom:12px;}
.skill-row{display:grid;grid-template-columns:64px 1fr 26px;align-items:center;gap:10px;margin-bottom:9px;}
.skill-row:last-child{margin-bottom:0;}
.skill-name{font-size:11.5px;color:var(--ink2);font-weight:500;}
.skill-track{height:5px;border-radius:99px;background:rgba(159,182,168,0.14);overflow:hidden;}
.skill-fill{height:100%;border-radius:99px;background:linear-gradient(90deg,var(--mint-deep),var(--mint));}
.skill-val{font-family:var(--mono);font-size:11px;color:var(--mint);text-align:right;}
.section{max-width:1180px;margin:0 auto;padding:96px 56px;}
.band{background:var(--bg-2);border-top:1px solid var(--line-2);border-bottom:1px solid var(--line-2);}
.sec-head{text-align:center;margin-bottom:62px;}
h2{font-family:var(--display);font-weight:800;font-size:clamp(34px,3.8vw,52px);letter-spacing:-0.025em;line-height:1.04;margin:16px 0 0;}
h2 .italic{font-family:var(--serif);font-weight:600;font-style:italic;color:var(--mint);}
.sec-sub{color:var(--ink2);font-size:16.5px;max-width:54ch;margin:18px auto 0;line-height:1.6;}
.sec-head .eyebrow{justify-content:center;}
.features-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:18px;}
.feature-card{background:var(--panel);border:1px solid var(--line);border-radius:18px;padding:28px;transition:transform .2s,border-color .2s,background .2s;position:relative;overflow:hidden;}
.feature-card:hover{transform:translateY(-4px);border-color:var(--mint-border);background:var(--panel-2);}
.feature-num{font-family:var(--mono);font-size:11px;color:var(--ink3);letter-spacing:0.1em;margin-bottom:18px;}
.feature-card h3{font-family:var(--display);font-size:18px;font-weight:700;color:var(--ink);margin-bottom:10px;letter-spacing:-0.01em;}
.feature-card p{font-size:13.5px;line-height:1.62;color:var(--ink2);}
.feature-tag{display:inline-block;margin-top:16px;font-family:var(--mono);font-size:10px;letter-spacing:0.08em;text-transform:uppercase;padding:4px 10px;border-radius:99px;background:var(--mint-dim);color:var(--mint);border:1px solid var(--mint-border);}
.steps{display:grid;grid-template-columns:repeat(4,1fr);gap:14px;}
.step{background:var(--panel);border:1px solid var(--line);border-radius:18px;padding:26px 22px;position:relative;transition:border-color .2s,background .2s;}
.step:hover{border-color:var(--mint-border);background:var(--panel-2);}
.step-num{font-family:var(--display);font-size:15px;font-weight:800;color:#06140f;width:34px;height:34px;border-radius:9px;background:var(--mint);display:grid;place-items:center;margin-bottom:18px;}
.step h3{font-family:var(--display);font-size:16px;font-weight:700;color:var(--ink);margin-bottom:8px;}
.step p{font-size:13px;line-height:1.58;color:var(--ink2);}
.pricing-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:18px;max-width:960px;margin:0 auto;align-items:start;}
.price-card{background:var(--panel);border:1px solid var(--line);border-radius:20px;padding:30px 26px;position:relative;}
.price-card.featured{background:linear-gradient(170deg,#103025,#0a1f18);border-color:var(--mint-border);box-shadow:0 24px 60px rgba(0,0,0,0.4),0 0 0 1px rgba(122,215,159,0.10) inset;transform:translateY(-10px);}
.price-badge{position:absolute;top:-12px;left:50%;transform:translateX(-50%);font-family:var(--mono);font-size:9.5px;font-weight:700;letter-spacing:0.14em;text-transform:uppercase;background:var(--mint);color:#06140f;padding:5px 14px;border-radius:99px;white-space:nowrap;box-shadow:0 6px 16px rgba(122,215,159,0.3);}
.price-plan{font-family:var(--display);font-size:18px;font-weight:700;margin-bottom:6px;}
.price-desc{font-size:12.5px;color:var(--ink3);line-height:1.5;margin-bottom:22px;min-height:38px;}
.price-num{font-family:var(--display);font-size:50px;font-weight:800;line-height:1;display:flex;align-items:flex-start;gap:3px;}
.price-cur{font-size:20px;margin-top:9px;color:var(--ink2);font-weight:600;}
.price-period{font-family:var(--mono);font-size:11px;color:var(--ink3);letter-spacing:0.04em;margin:8px 0 24px;}
.price-cta{display:block;text-align:center;width:100%;padding:13px 0;border-radius:10px;font-size:14px;font-weight:600;text-decoration:none;cursor:pointer;margin-bottom:24px;transition:background .15s,transform .15s,border-color .15s;}
.cta-fill{background:var(--mint);color:#06140f;}
.cta-fill:hover{background:var(--mint-bright);transform:translateY(-1px);}
.cta-out{background:transparent;border:1px solid var(--line);color:var(--ink);}
.cta-out:hover{border-color:var(--mint-border);background:var(--mint-dim);}
.price-features{list-style:none;display:flex;flex-direction:column;gap:11px;}
.pf{display:flex;align-items:flex-start;gap:10px;font-size:13px;line-height:1.4;color:var(--ink2);}
.pf b{color:var(--ink);font-weight:600;}
.ck{flex:0 0 17px;width:17px;height:17px;border-radius:50%;display:grid;place-items:center;font-size:9px;margin-top:1px;}
.ck.y{background:var(--mint-dim2);color:var(--mint);}
.ck.n{background:rgba(159,182,168,0.08);color:var(--ink3);}
.cta-strip{max-width:1180px;margin:0 auto;padding:0 56px 96px;}
.cta-inner{background:linear-gradient(150deg,#11362a,#0a1f18);border:1px solid var(--mint-border);border-radius:24px;padding:56px;text-align:center;position:relative;overflow:hidden;}
.cta-inner::before{content:'';position:absolute;inset:0;background:radial-gradient(600px 300px at 50% -30%,rgba(122,215,159,0.16),transparent 70%);pointer-events:none;}
.cta-inner h2{position:relative;}
.cta-inner p{position:relative;color:var(--ink2);font-size:16px;margin:16px auto 30px;max-width:46ch;line-height:1.6;}
.cta-inner .hero-btns{justify-content:center;margin-bottom:0;position:relative;}
.rs-footer{border-top:1px solid var(--line-2);padding:38px 56px;max-width:1280px;margin:0 auto;display:flex;align-items:center;justify-content:space-between;gap:20px;flex-wrap:wrap;}
.footer-logo{display:flex;align-items:center;gap:10px;font-family:var(--display);font-weight:700;font-size:15px;color:var(--ink);}
.footer-mark{width:28px;height:28px;border-radius:7px;background:linear-gradient(150deg,var(--mint),var(--mint-deep));display:grid;place-items:center;font-size:12px;font-weight:800;color:#06140f;}
.footer-note{font-family:var(--mono);font-size:11px;color:var(--ink3);letter-spacing:0.03em;}
.footer-note a{color:var(--mint);text-decoration:none;}
@media(max-width:1080px){.hero{grid-template-columns:1fr;gap:64px;padding-bottom:80px;}.hero-visual{max-width:460px;}.float-skills{right:-20px;}.float-score{right:-14px;}.features-grid{grid-template-columns:1fr 1fr;}.steps{grid-template-columns:1fr 1fr;}.pricing-grid{grid-template-columns:1fr;max-width:440px;}.price-card.featured{transform:none;}}
@media(max-width:640px){.nav{padding:14px 20px;}.nav-links{display:none;}.hero{padding:44px 20px 56px;}.section{padding:64px 20px;}.cta-strip{padding:0 20px 64px;}.cta-inner{padding:38px 24px;}.features-grid,.steps{grid-template-columns:1fr;}.float-score{top:-18px;right:-8px;transform:scale(0.88)}.float-skills{right:-8px;bottom:20px;transform:scale(0.9)}.rs-footer{padding:28px 20px;flex-direction:column;text-align:center;}}
</style>
</head>
<body>
<nav class="nav">
<a class="logo" href="#">
<div class="logo-mark">R</div>
<div class="logo-txt">
<div class="logo-name">Resume<span>Sync</span></div>
<div class="logo-by">by VisualizePro</div>
</div>
</a>
<div class="nav-links">
<a class="nav-link" href="#features">Features</a>
<a class="nav-link" href="#how">How it works</a>
<a class="nav-link" href="#pricing">Pricing</a>
</div>
<a class="nav-cta cta-login" id="nav-cta-btn" href="#">Log in &#8594;</a>
</nav>
<section class="hero">
<div class="hero-copy">
<div class="eyebrow">know before you apply</div>
<h1>Stop applying blindly.<br><span class="italic">Know your match</span><br>before you apply.</h1>
<p class="hero-lede">Upload your resume, paste any job posting &#8212; and in seconds know your exact <strong>match score</strong>, which skills are missing, which keywords get you <strong>past ATS</strong>, and walk away with a tailored cover letter and resume ready to send.</p>
<div class="hero-btns">
<a class="btn-primary cta-login" href="#">Start free &#8212; no card needed</a>
<a class="btn-ghost" href="#how">See how it works &#8594;</a>
</div>
<div class="hero-social">
<div class="social-item"><span class="tick">&#10003;</span><b>3</b> free applications</div>
<div class="social-div"></div>
<div class="social-item"><span class="tick">&#10003;</span><b>AI</b> resume rewrites</div>
<div class="social-div"></div>
<div class="social-item"><span class="tick">&#10003;</span>Cover letter <b>included</b></div>
</div>
</div>
<div class="hero-visual">
<div class="photo-frame">
<img src="https://images.unsplash.com/photo-1573496359142-b8d87734a5a2?w=900&q=80&auto=format&fit=crop" alt="Professional" loading="lazy"/>
<div class="photo-badge"><span class="pulse"></span>Now hiring &middot; You</div>
<div class="photo-name">
<div class="nm">Madeli, BI Specialist</div>
<div class="role">Currently matching to 12 open roles</div>
</div>
</div>
<div class="float float-score">
<div class="fs-label">Match Score</div>
<div class="ring-wrap">
<svg width="108" height="108" viewBox="0 0 108 108">
<circle cx="54" cy="54" r="48" fill="none" stroke="rgba(159,182,168,0.16)" stroke-width="8"/>
<circle cx="54" cy="54" r="48" fill="none" stroke="var(--mint)" stroke-width="8" stroke-linecap="round" stroke-dasharray="301.6" stroke-dashoffset="24.1"/>
</svg>
<div class="ring-num">92<small>%</small></div>
</div>
<div class="fs-fit">Strong fit</div>
</div>
<div class="float float-skills">
<div class="fsk-label">Skill alignment</div>
<div class="skill-row"><span class="skill-name">Tableau</span><span class="skill-track"><span class="skill-fill" style="width:96%"></span></span><span class="skill-val">96</span></div>
<div class="skill-row"><span class="skill-name">Power BI</span><span class="skill-track"><span class="skill-fill" style="width:82%"></span></span><span class="skill-val">82</span></div>
<div class="skill-row"><span class="skill-name">SQL</span><span class="skill-track"><span class="skill-fill" style="width:90%"></span></span><span class="skill-val">90</span></div>
<div class="skill-row"><span class="skill-name">Snowflake</span><span class="skill-track"><span class="skill-fill" style="width:74%"></span></span><span class="skill-val">74</span></div>
</div>
</div>
</section>
<div class="band">
<section class="section" id="features">
<div class="sec-head">
<div class="eyebrow one">what resumesync does</div>
<h2>Every tool you need to<br><span class="italic">land the job.</span></h2>
<p class="sec-sub">From pasting a job posting to sending your application &#8212; one workflow handles the score, the rewrite, the cover letter, and the tracking.</p>
</div>
<div class="features-grid">
<div class="feature-card"><div class="feature-num">01</div><h3>Match Score</h3><p>Paste any job description and get an instant percentage showing how well your resume fits &#8212; scored on skills, keywords, experience, and tone.</p><span class="feature-tag">Instant results</span></div>
<div class="feature-card"><div class="feature-num">02</div><h3>Skill &amp; keyword gaps</h3><p>See exactly what&#8217;s missing &#8212; which keywords to add, which experience to emphasise, which sections to rewrite. Specific, actionable, never vague.</p><span class="feature-tag">AI-powered</span></div>
<div class="feature-card"><div class="feature-num">03</div><h3>Resume rewriter</h3><p>One click tailors your resume to the job &#8212; sharpening language and lifting your ATS score without inventing anything or losing your voice.</p><span class="feature-tag">Claude AI</span></div>
<div class="feature-card"><div class="feature-num">04</div><h3>Cover letter generator</h3><p>A tailored cover letter built from the posting and your updated resume. Sounds human, not robotic &#8212; and you can refine it with your own prompts.</p><span class="feature-tag">Tailored per role</span></div>
<div class="feature-card"><div class="feature-num">05</div><h3>Prompt &amp; refine</h3><p>Not happy with the output? Tell it what to change. &#8220;Make it more senior.&#8221; &#8220;Emphasise leadership.&#8221; You always direct the final result.</p><span class="feature-tag">You direct the AI</span></div>
<div class="feature-card"><div class="feature-num">06</div><h3>Application tracker</h3><p>Every application in one place &#8212; role, company, match score, resume version, cover letter, and status. Your whole job search, organised.</p><span class="feature-tag">Never lose track</span></div>
</div>
</section>
</div>
<section class="section" id="how">
<div class="sec-head">
<div class="eyebrow one">how it works</div>
<h2>Four steps to a<br><span class="italic">stronger application.</span></h2>
<p class="sec-sub">From job posting to ready-to-send application in under ten minutes.</p>
</div>
<div class="steps">
<div class="step"><div class="step-num">1</div><h3>Paste the job</h3><p>Drop in a Seek or LinkedIn URL, or paste the description. We scrape the rest.</p></div>
<div class="step"><div class="step-num">2</div><h3>Get your score</h3><p>See your match percentage instantly, with a clear list of gaps to close.</p></div>
<div class="step"><div class="step-num">3</div><h3>Rewrite &amp; refine</h3><p>Let AI rewrite your resume and cover letter, then tweak with your own prompts.</p></div>
<div class="step"><div class="step-num">4</div><h3>Track &amp; apply</h3><p>Save it all to your tracker and send. Every document stored, every status logged.</p></div>
</div>
</section>
<div class="band">
<section class="section" id="pricing">
<div class="sec-head">
<div class="eyebrow one">simple pricing</div>
<h2>Buy what you need.<br><span class="italic">No subscription.</span></h2>
<p class="sec-sub">Start with 3 free analyses. Top up whenever you like &#8212; credits never expire.</p>
</div>
<div class="pricing-grid">
<div class="price-card">
<div class="price-plan">Free</div>
<div class="price-desc">Try it on your first three applications. No credit card needed.</div>
<div class="price-num"><span class="price-cur">A$</span>0</div>
<div class="price-period">3 analyses included</div>
<a class="price-cta cta-out cta-login" data-plan="free" href="#">Get started &#8594;</a>
<ul class="price-features">
<li class="pf"><span class="ck y">&#10003;</span><span>3 analyses included</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Match score + gap report</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Resume rewriter</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Cover letter generator</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Application tracker</span></li>
</ul>
</div>
<div class="price-card featured">
<div class="price-badge">Most popular</div>
<div class="price-plan">Value Pack</div>
<div class="price-desc">Great value for an active job search. Credits never expire.</div>
<div class="price-num"><span class="price-cur">A$</span>15</div>
<div class="price-period">20 analyses &middot; one-off</div>
<a class="price-cta cta-fill cta-login" data-plan="pack20" href="#">Buy 20 analyses &#8594;</a>
<ul class="price-features">
<li class="pf"><span class="ck y">&#10003;</span><span>20 analyses</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Match score + gap report</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Resume rewriter</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Cover letter generator</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Application tracker</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Credits never expire</span></li>
</ul>
</div>
<div class="price-card">
<div class="price-plan">Starter Pack</div>
<div class="price-desc">Enough for a focused sprint of applications.</div>
<div class="price-num"><span class="price-cur">A$</span>9</div>
<div class="price-period">5 analyses &middot; one-off</div>
<a class="price-cta cta-out cta-login" data-plan="pack5" href="#">Buy 5 analyses &#8594;</a>
<ul class="price-features">
<li class="pf"><span class="ck y">&#10003;</span><span>5 analyses</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Match score + gap report</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Resume rewriter</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Cover letter generator</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Application tracker</span></li>
<li class="pf"><span class="ck y">&#10003;</span><span>Credits never expire</span></li>
</ul>
</div>
</div>
<div style="text-align:center;margin-top:1.5rem;">
<p style="font-family:'DM Sans',sans-serif;font-size:14px;color:#6e8a7b;margin:0 0 0.75rem;">
  Need more? <strong style="color:#9fb6a8;">Pro Pack — 30 analyses for A$20</strong> &middot; <strong style="color:#9fb6a8;">Max Pack — 50 analyses for A$30</strong>
</p>
<a class="price-cta cta-out cta-login" data-plan="pack30" href="#" style="display:inline-block;margin:0 6px 6px;">Pro Pack A$20 &#8594;</a>
<a class="price-cta cta-out cta-login" data-plan="pack50" href="#" style="display:inline-block;margin:0 6px 6px;">Max Pack A$30 &#8594;</a>
</div>
</section>
</div>
<div class="cta-strip" id="cta">
<div class="cta-inner">
<div class="eyebrow one" style="justify-content:center;">start in 60 seconds</div>
<h2 style="margin-top:14px;">Know your match.<br><span class="italic">Then apply with confidence.</span></h2>
<p>Your first three applications are free. Upload a resume, paste a job, and see your score &#8212; no card, no commitment.</p>
<div class="hero-btns">
<a class="btn-primary cta-login" href="#">Start free &#8212; no card needed</a>
<a class="btn-ghost" href="#features">Explore features &#8594;</a>
</div>
</div>
</div>
<div class="rs-footer">
<div class="footer-logo"><div class="footer-mark">R</div>ResumeSync</div>
<div class="footer-note">Built by <a href="https://visualizepro.com.au" target="_blank">VisualizePro</a> &middot; Sydney, Australia &middot; Powered by Claude AI</div>
</div>
<script>
(function() {
  // Resize iframe to match parent viewport so no double-scrollbar confusion
  function resizeToViewport() {
    try {
      var frames = window.parent.document.querySelectorAll('iframe');
      for (var i = 0; i < frames.length; i++) {
        if (frames[i].contentWindow === window) {
          frames[i].style.height = (window.parent.innerHeight || 768) + 'px';
          break;
        }
      }
    } catch(e) {}
  }
  resizeToViewport();
  window.addEventListener('resize', resizeToViewport);

  // Navigate parent page to login, optionally with a plan
  function triggerLogin(plan) {
    var planParam = plan ? '&plan=' + encodeURIComponent(plan) : '';
    // Method 1: direct parent location (works for same-origin user-gesture in most browsers)
    try {
      window.parent.location.href = window.parent.location.pathname + '?login=1' + planParam;
      return;
    } catch(e1) {}
    // Method 2: click the hidden trigger button in the parent Streamlit DOM
    try {
      var btns = window.parent.document.querySelectorAll('button');
      for (var i = 0; i < btns.length; i++) {
        if (btns[i].textContent.includes('__rs_login__')) {
          btns[i].click();
          return;
        }
      }
    } catch(e2) {}
    // Method 3: navigate same tab as last resort
    try {
      window.location.href = '?login=1' + planParam;
    } catch(e3) {}
  }

  // Single event-delegation listener — handles all clicks in the document
  document.addEventListener('click', function(e) {
    // Walk up from clicked element to find the <a> tag
    var el = e.target;
    while (el && el !== document.body) {
      if (el.tagName === 'A') {
        var href = el.getAttribute('href') || '';
        if (el.classList.contains('cta-login')) {
          e.preventDefault();
          var plan = el.getAttribute('data-plan') || '';
          triggerLogin(plan);
          return;
        }
        if (href.startsWith('#') && href.length > 1) {
          e.preventDefault();
          var target = document.getElementById(href.slice(1));
          if (target) target.scrollIntoView({ behavior: 'smooth', block: 'start' });
          return;
        }
        break;
      }
      el = el.parentElement;
    }
  });
})();
</script>
</body>
</html>"""

    _components.html(_landing_doc, height=850, scrolling=True)



# ============= LOGIN PAGE =============

def show_login_page():
    """Email + password login / signup screen."""
    _, col, _ = st.columns([1, 1.6, 1])
    with col:
        # Toggle between sign-in and sign-up
        mode = st.session_state.get("login_mode", "signin")

        st.markdown(f"""
        <div style="padding:3rem 0 1.5rem;">
          <div style="display:flex;align-items:center;gap:12px;margin-bottom:2.5rem;">
            <div style="width:44px;height:44px;border-radius:12px;
                        background:linear-gradient(150deg,#7ad79f,#4fae7a);
                        display:grid;place-items:center;
                        font-family:'Bricolage Grotesque',sans-serif;
                        font-weight:800;font-size:22px;color:#06140f;
                        box-shadow:0 4px 16px rgba(122,215,159,0.30);">R</div>
            <div>
              <div style="font-family:'Bricolage Grotesque',sans-serif;font-weight:700;
                          font-size:20px;color:#ecf4ee;letter-spacing:-0.02em;">ResumeSync</div>
              <div style="font-family:'Space Mono',monospace;font-size:9px;
                          letter-spacing:0.18em;text-transform:uppercase;color:#6e8a7b;">by VisualizePro</div>
            </div>
          </div>
          <h1 style="font-family:'Bricolage Grotesque',sans-serif;font-weight:800;
                     font-size:30px;color:#ecf4ee;margin:0 0 0.4rem;letter-spacing:-0.02em;">
            {"Create account" if mode == "signup" else "Sign in"}
          </h1>
          <p style="font-family:'DM Sans',sans-serif;font-size:15px;color:#9fb6a8;margin:0 0 2rem;">
            {"Enter your email and choose a password." if mode == "signup" else "Welcome back — enter your email and password."}
          </p>
        </div>
        """, unsafe_allow_html=True)

        email_input = st.text_input(
            "Email address",
            placeholder="you@example.com",
            label_visibility="collapsed",
            key="login_email_input",
        )
        password_input = st.text_input(
            "Password",
            placeholder="Password (min 6 characters)",
            type="password",
            label_visibility="collapsed",
            key="login_password_input",
        )
        _components.html("""<script>
(function() {
  function patch() {
    var inputs = window.parent.document.querySelectorAll('input[type="password"]');
    inputs.forEach(function(el) { el.setAttribute('autocomplete', 'off'); });
  }
  patch();
  var obs = new MutationObserver(patch);
  obs.observe(window.parent.document.body, { childList: true, subtree: true });
})();
</script>""", height=0)

        if mode == "signup":
            if st.button("Create account →", type="primary", use_container_width=True, key="signup_btn"):
                email_val    = (email_input or "").strip().lower()
                password_val = (password_input or "").strip()
                if not email_val or "@" not in email_val:
                    st.error("Please enter a valid email address.")
                elif len(password_val) < 6:
                    st.error("Password must be at least 6 characters.")
                else:
                    info, err = sign_up(email_val, password_val)
                    if info:
                        st.session_state["auth_user_id"]  = info["user_id"]
                        st.session_state["auth_email"]    = info["email"]
                        st.session_state["auth_token"]    = info["access_token"]
                        st.session_state["auth_refresh"]  = info["refresh_token"]
                        st.session_state["user_profile"]  = get_or_create_profile(info["user_id"], info["email"])
                        st.session_state.pop("login_mode", None)
                        st.rerun()
                    else:
                        st.error(err)
            st.markdown("<div style='height:0.5rem;'></div>", unsafe_allow_html=True)
            if st.button("Already have an account? Sign in", key="switch_to_signin"):
                st.session_state["login_mode"] = "signin"
                st.rerun()
        else:
            if st.button("Sign in →", type="primary", use_container_width=True, key="signin_btn"):
                email_val    = (email_input or "").strip().lower()
                password_val = (password_input or "").strip()
                if not email_val or "@" not in email_val:
                    st.error("Please enter a valid email address.")
                elif not password_val:
                    st.error("Please enter your password.")
                else:
                    info, err = sign_in(email_val, password_val)
                    if info:
                        st.session_state["auth_user_id"]  = info["user_id"]
                        st.session_state["auth_email"]    = info["email"]
                        st.session_state["auth_token"]    = info["access_token"]
                        st.session_state["auth_refresh"]  = info["refresh_token"]
                        st.session_state["user_profile"]  = get_or_create_profile(info["user_id"], info["email"])
                        st.session_state.pop("login_mode", None)
                        st.rerun()
                    else:
                        st.error(err)
            st.markdown("<div style='height:0.5rem;'></div>", unsafe_allow_html=True)
            if st.button("No account yet? Create one free", key="switch_to_signup"):
                st.session_state["login_mode"] = "signup"
                st.rerun()

        st.markdown("""
        <div style="margin-top:2.5rem;padding-top:1.5rem;border-top:1px solid rgba(159,182,168,0.12);">
          <p style="font-family:'DM Sans',sans-serif;font-size:13px;color:#6e8a7b;margin:0 0 1.2rem;">
            <b style="color:#9fb6a8;">Free:</b> 3 analyses included ·
            <b style="color:#9fb6a8;">Packs from A$9</b> — no subscription, never expires
          </p>
          <a href="https://madelivniekerk.github.io/Resumesync"
             style="font-family:'DM Sans',sans-serif;font-size:13px;color:#6e8a7b;
                    text-decoration:none;display:inline-flex;align-items:center;gap:4px;
                    transition:color 0.2s;"
             onmouseover="this.style.color='#9fb6a8'"
             onmouseout="this.style.color='#6e8a7b'">
            ← Back to main page
          </a>
        </div>
        """, unsafe_allow_html=True)


# ============= PAYMENT PAGES =============

def show_payment_redirect(pack: str):
    """Shown after login when user selected a pack — redirect them to Pay Advanced checkout."""
    info = PACK_INFO.get(pack, {})
    if not info:
        st.rerun()
        return

    auth_email = st.session_state.get("auth_email", "")
    _ref = f"{auth_email}|{pack}"[:50]
    pay_params = urllib.parse.urlencode({
        "paymentamount":      info["amount"],
        "paymentdescription": f"ResumeSync {info['name']} ({info['credits']} analyses)",
        "paymentref":         _ref,
        "email":              auth_email,
    })
    pay_url = PAY_ADVANCED_URL + "?" + pay_params

    st.markdown("""
<style>
[data-testid="stHeader"]{display:none!important;}
#MainMenu{visibility:hidden!important;}
footer{visibility:hidden!important;}
</style>""", unsafe_allow_html=True)

    _, col, _ = st.columns([1, 1.6, 1])
    with col:
        st.markdown(f"""
<div style="padding:3rem 0 1rem;">
  <div style="display:flex;align-items:center;gap:12px;margin-bottom:2.5rem;">
    <div style="width:44px;height:44px;border-radius:12px;
                background:linear-gradient(150deg,#7ad79f,#4fae7a);
                display:grid;place-items:center;
                font-family:'Bricolage Grotesque',sans-serif;
                font-weight:800;font-size:22px;color:#06140f;
                box-shadow:0 4px 16px rgba(122,215,159,0.30);">R</div>
    <div>
      <div style="font-family:'Bricolage Grotesque',sans-serif;font-weight:700;
                  font-size:20px;color:#ecf4ee;letter-spacing:-0.02em;">ResumeSync</div>
      <div style="font-family:'Space Mono',monospace;font-size:9px;
                  letter-spacing:0.18em;text-transform:uppercase;color:#6e8a7b;">by VisualizePro</div>
    </div>
  </div>
  <div style="background:#0c2019;border:1px solid rgba(159,182,168,0.12);
              border-radius:14px;padding:1.8rem;margin-bottom:1.5rem;">
    <div style="font-family:'Space Mono',monospace;font-size:10px;letter-spacing:0.14em;
                text-transform:uppercase;color:{info['color']};margin-bottom:0.5rem;">Selected pack</div>
    <div style="font-family:'Bricolage Grotesque',sans-serif;font-weight:800;
                font-size:26px;color:#ecf4ee;margin-bottom:0.25rem;">{info['name']}</div>
    <div style="font-family:'DM Sans',sans-serif;font-size:22px;font-weight:700;
                color:{info['color']};margin-bottom:0.5rem;">{info['price']} one-off</div>
    <div style="font-family:'DM Sans',sans-serif;font-size:14px;color:#9fb6a8;">{info['credits']} analyses · never expires · no subscription</div>
  </div>
  <p style="font-family:'DM Sans',sans-serif;font-size:14px;color:#9fb6a8;margin:0 0 1.5rem;">
    You're signed in as <strong style="color:#ecf4ee;">{auth_email}</strong>.
    Click below to complete your payment — you'll be redirected back automatically.
  </p>
</div>""", unsafe_allow_html=True)

        st.markdown(f"""
<a href="{pay_url}" target="_self" style="
  display:block;width:100%;padding:0.65rem 1rem;
  background:linear-gradient(135deg,#7ad79f,#4fae7a);
  color:#06140f;font-family:'DM Sans',sans-serif;font-weight:700;
  font-size:15px;text-align:center;border-radius:9px;
  text-decoration:none;box-sizing:border-box;
  box-shadow:0 4px 14px rgba(122,215,159,0.30);">
  Pay {info['price']} — get {info['credits']} analyses →
</a>""", unsafe_allow_html=True)
        st.markdown("<div style='height:0.75rem;'></div>", unsafe_allow_html=True)
        if st.button("Skip — use my free credits for now", key="skip_payment_btn", use_container_width=True):
            st.rerun()


def show_payment_pending():
    """Shown after returning from Pay Advanced while the webhook processes."""
    st.markdown("""
<style>
[data-testid="stHeader"]{display:none!important;}
#MainMenu{visibility:hidden!important;}
footer{visibility:hidden!important;}
</style>""", unsafe_allow_html=True)

    _, col, _ = st.columns([1, 1.6, 1])
    with col:
        st.markdown("""
<div style="padding:3rem 0 1rem;">
  <div style="display:flex;align-items:center;gap:12px;margin-bottom:2.5rem;">
    <div style="width:44px;height:44px;border-radius:12px;
                background:linear-gradient(150deg,#7ad79f,#4fae7a);
                display:grid;place-items:center;
                font-family:'Bricolage Grotesque',sans-serif;
                font-weight:800;font-size:22px;color:#06140f;
                box-shadow:0 4px 16px rgba(122,215,159,0.30);">R</div>
    <div>
      <div style="font-family:'Bricolage Grotesque',sans-serif;font-weight:700;
                  font-size:20px;color:#ecf4ee;letter-spacing:-0.02em;">ResumeSync</div>
      <div style="font-family:'Space Mono',monospace;font-size:9px;
                  letter-spacing:0.18em;text-transform:uppercase;color:#6e8a7b;">by VisualizePro</div>
    </div>
  </div>
  <div style="background:#0c2019;border:1px solid rgba(122,215,159,0.20);
              border-radius:14px;padding:1.8rem;margin-bottom:1.5rem;text-align:center;">
    <div style="font-size:36px;margin-bottom:0.75rem;">✓</div>
    <div style="font-family:'Bricolage Grotesque',sans-serif;font-weight:800;
                font-size:22px;color:#ecf4ee;margin-bottom:0.5rem;">Payment received!</div>
    <p style="font-family:'DM Sans',sans-serif;font-size:14px;color:#9fb6a8;margin:0;">
      Your credits are being added. It may take a moment to activate —
      click below to continue and your balance will reflect shortly.
    </p>
  </div>
</div>""", unsafe_allow_html=True)

        if st.button("Continue to app →", type="primary", use_container_width=True, key="payment_continue_btn"):
            # Refresh the profile from Supabase to pick up the webhook-updated tier
            uid = st.session_state.get("auth_user_id")
            email = st.session_state.get("auth_email", "")
            if uid:
                st.session_state["user_profile"] = get_or_create_profile(uid, email)
            st.rerun()


# ============= TRACKER PAGE =============

def show_tracker():
    """Dedicated Applications Tracker page."""

    # Handle status changes submitted via the HTML select + URL params
    _p = st.query_params
    _rid = _p.get('rid', '')
    _sc  = _p.get('sc', '')
    if _rid and _sc in ['Applied', 'Interview', 'Offer', 'Declined', 'Rejected', 'Draft']:
        update_application_status(_rid, _sc)
        st.query_params.clear()
        st.rerun()

    # Hide Streamlit header
    st.markdown("""
    <style>
    [data-testid="stHeader"]{display:none!important;}
    #MainMenu{visibility:hidden!important;}
    footer{visibility:hidden!important;}
    </style>
    """, unsafe_allow_html=True)

    _tracker_uid = st.session_state.get('auth_user_id')
    tracker_data = load_tracker_data(user_id=_tracker_uid)

    # ── Sidebar ──────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown("""
        <div style="padding:1.5rem 1rem 0;">
          <div style="display:flex;align-items:center;gap:11px;margin-bottom:1.6rem;">
            <div style="width:34px;height:34px;border-radius:9px;
                        background:linear-gradient(150deg,#7ad79f,#4fae7a);
                        display:grid;place-items:center;
                        font-family:'Bricolage Grotesque',system-ui,sans-serif;
                        font-weight:800;font-size:16px;color:#06140f;
                        box-shadow:0 4px 12px rgba(122,215,159,0.28);flex-shrink:0;">R</div>
            <div>
              <div style="font-family:'Bricolage Grotesque',system-ui,sans-serif;font-weight:700;
                          font-size:17px;letter-spacing:-0.02em;color:#ecf4ee;line-height:1.1;">ResumeSync</div>
              <div style="font-family:'Space Mono',monospace;font-size:9px;letter-spacing:0.18em;
                          text-transform:uppercase;color:#6e8a7b;margin-top:3px;">by VisualizePro</div>
            </div>
          </div>

          <div style="border-top:1px solid rgba(159,182,168,0.12);margin-bottom:1.2rem;"></div>

          <p style="font-family:'Space Mono',monospace;font-size:9px;letter-spacing:0.16em;
                    text-transform:uppercase;color:#7ad79f;margin:0 0 0.8rem;">Workspace</p>
        </div>
        """, unsafe_allow_html=True)

        if st.button("✦ New Analysis", key="tracker_nav_new", use_container_width=True):
            st.session_state['page'] = 'app'
            st.rerun()

        # Applications — active/current page
        st.markdown("""
        <div style="background:rgba(122,215,159,0.10);border:1px solid rgba(122,215,159,0.22);
                    border-radius:6px;padding:10px 14px;margin:4px 0;
                    font-family:'DM Sans',sans-serif;font-size:0.875rem;font-weight:600;color:#7ad79f;">
          📋 Applications
        </div>
        """, unsafe_allow_html=True)

        st.markdown("""
        <div style='margin-top:1.5rem;border-top:1px solid rgba(159,182,168,0.12);padding-top:1rem;'>
          <a href='https://madelivniekerk.github.io/Resumesync/' target='_self'
             style='display:block;text-align:center;font-family:DM Sans,sans-serif;font-size:13px;
                    color:#6e8a7b;text-decoration:none;padding:6px 0;'>
            ← Back to home
          </a>
        </div>""", unsafe_allow_html=True)

    # ── Stats computation ─────────────────────────────────────────────────────
    total = len(tracker_data)
    applied_count = sum(1 for r in tracker_data if r.get('status') == 'Applied')
    interview_count = sum(1 for r in tracker_data if r.get('status') == 'Interview')
    pct_vals = []
    for r in tracker_data:
        raw = str(r.get('match_pct', '') or '').replace('%', '').strip()
        if raw.isdigit():
            pct_vals.append(int(raw))
    avg_match = (sum(pct_vals) // len(pct_vals)) if pct_vals else 0

    # ── Page header ───────────────────────────────────────────────────────────
    hdr_l, hdr_r = st.columns([3, 1])
    with hdr_l:
        st.markdown("""
        <div style="padding:28px 0 8px;">
          <h1 style="font-family:'Bricolage Grotesque',system-ui,sans-serif;font-weight:700;
                     font-size:clamp(26px,3vw,38px);letter-spacing:-0.025em;color:#ecf4ee;margin:0 0 8px;">
            Applications
          </h1>
          <p style="font-family:'DM Sans',sans-serif;font-size:14px;color:#9fb6a8;margin:0;">
            Every role you've analysed — resume, cover letter, score and status in one place.
          </p>
        </div>
        """, unsafe_allow_html=True)
    with hdr_r:
        st.markdown("<div style='padding-top:32px;'>", unsafe_allow_html=True)
        if st.button("New analysis →", key="tracker_hdr_new", type="primary", use_container_width=True):
            st.session_state['page'] = 'app'
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    # ── Stats row ─────────────────────────────────────────────────────────────
    s1, s2, s3, s4 = st.columns(4)
    stat_cards = [
        (s1, str(total), "Tracked", "#7ad79f"),
        (s2, str(applied_count), "Applied", "#6fb1e0"),
        (s3, str(interview_count), "Interviews", "#e0a14a"),
        (s4, (str(avg_match) + "%" if pct_vals else "—"), "Avg Match", "#7ad79f"),
    ]
    for col, val, label, color in stat_cards:
        with col:
            st.markdown(
                '<div style="background:#0c2019;border:1px solid rgba(159,182,168,0.12);'
                'border-radius:14px;padding:20px 18px;text-align:center;margin-bottom:16px;">'
                '<div style="font-family:\'Bricolage Grotesque\',system-ui,sans-serif;font-weight:800;'
                'font-size:2rem;color:' + color + ';line-height:1;">' + val + '</div>'
                '<div style="font-family:\'Space Mono\',monospace;font-size:9px;letter-spacing:0.16em;'
                'text-transform:uppercase;color:#6e8a7b;margin-top:6px;">' + label + '</div>'
                '</div>',
                unsafe_allow_html=True
            )

    # ── Empty state ───────────────────────────────────────────────────────────
    if total == 0:
        st.markdown("""
        <div style="background:#0c2019;border:1px solid rgba(159,182,168,0.12);border-radius:16px;
                    padding:48px 24px;text-align:center;margin-top:8px;">
          <div style="font-size:2.4rem;margin-bottom:14px;">📋</div>
          <div style="font-family:'Bricolage Grotesque',system-ui,sans-serif;font-weight:700;
                      font-size:20px;color:#ecf4ee;margin-bottom:8px;">No applications yet</div>
          <div style="font-family:'DM Sans',sans-serif;font-size:14px;color:#9fb6a8;margin-bottom:24px;">
            Run your first resume analysis to start tracking your job applications.
          </div>
        </div>
        """, unsafe_allow_html=True)
        _, mid_col, _ = st.columns([1, 1, 1])
        with mid_col:
            if st.button("Run your first analysis →", key="tracker_empty_cta", type="primary", use_container_width=True):
                st.session_state.page = 'app'
                st.rerun()
    else:
        STATUS_OPTIONS = ['Applied', 'Interview', 'Offer', 'Declined', 'Rejected', 'Draft']
        STATUS_CONFIG = {
            'Applied':   {'color': '#6fb1e0', 'bg': 'rgba(111,177,224,0.12)', 'border': 'rgba(111,177,224,0.30)'},
            'Interview': {'color': '#e0a14a', 'bg': 'rgba(224,161,74,0.14)',  'border': 'rgba(224,161,74,0.34)'},
            'Offer':     {'color': '#7ad79f', 'bg': 'rgba(122,215,159,0.12)', 'border': 'rgba(122,215,159,0.30)'},
            'Declined':  {'color': '#e07a5f', 'bg': 'rgba(224,122,95,0.12)', 'border': 'rgba(224,122,95,0.30)'},
            'Rejected':  {'color': '#ef4444', 'bg': 'rgba(239,68,68,0.10)',  'border': 'rgba(239,68,68,0.30)'},
            'Draft':     {'color': '#6e8a7b', 'bg': 'rgba(110,138,123,0.12)','border': 'rgba(110,138,123,0.28)'},
        }
        logo_colors = ['#1d3a31', '#1a2f3e', '#2d2518', '#2a1f35', '#1a2e2a', '#2e1f1f']

        # ── Compact row styling ────────────────────────────────────────────────
        st.markdown("""<style>
        [data-testid="stHorizontalBlock"] .element-container {
            margin-bottom: 0 !important;
            padding-bottom: 0 !important;
        }
        [data-testid="stHorizontalBlock"] [data-testid="stVerticalBlockBorderWrapper"],
        [data-testid="stHorizontalBlock"] [data-testid="stVerticalBlock"] {
            gap: 0 !important;
        }
        /* Remove padding around component iframes used for HTML selects */
        [data-testid="stHorizontalBlock"] iframe {
            display: block !important;
        }
        </style>""", unsafe_allow_html=True)

        # ── Table card header ─────────────────────────────────────────────────
        st.markdown("""
        <div style="background:#0c2019;border:1px solid rgba(159,182,168,0.12);
                    border-radius:14px 14px 0 0;padding:16px 22px 14px;margin-top:4px;
                    border-bottom:1px solid rgba(159,182,168,0.09);">
          <span style="font-family:'Bricolage Grotesque',sans-serif;font-weight:700;
                       font-size:16px;color:#ecf4ee;">All applications</span>
        </div>
        """, unsafe_allow_html=True)

        # ── Column headers ────────────────────────────────────────────────────
        hc0, hc1, hc2, hc3 = st.columns([3.2, 0.9, 1.3, 2.4])
        hs = ("font-family:'Space Mono',monospace;font-size:9.5px;letter-spacing:0.14em;"
              "text-transform:uppercase;color:#6e8a7b;padding:10px 0 6px;")
        with hc0: st.markdown(f'<div style="{hs}">Role</div>', unsafe_allow_html=True)
        with hc1: st.markdown(f'<div style="{hs}text-align:center;">Match</div>', unsafe_allow_html=True)
        with hc2: st.markdown(f'<div style="{hs}">Applied</div>', unsafe_allow_html=True)
        with hc3: st.markdown(f'<div style="{hs}">Status</div>', unsafe_allow_html=True)

        # ── Rows ──────────────────────────────────────────────────────────────
        for idx, row in enumerate(tracker_data):
            record_id = row.get('id')
            company   = row.get('company', '') or ''
            job_title = row.get('job_title', '') or ''
            match_pct = str(row.get('match_pct', '') or '')
            status    = row.get('status', 'Draft') or 'Draft'
            date_val  = str(row.get('date', '') or '')
            initial   = company[0].upper() if company else '?'
            logo_bg   = logo_colors[idx % len(logo_colors)]
            cfg       = STATUS_CONFIG.get(status, STATUS_CONFIG['Draft'])

            st.markdown(
                '<div style="height:1px;background:rgba(159,182,168,0.08);"></div>',
                unsafe_allow_html=True
            )

            col_role, col_match, col_date, col_status = st.columns([3.2, 0.9, 1.3, 2.4])

            with col_role:
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:12px;padding:5px 0 4px;">'
                    f'<div style="width:36px;height:36px;border-radius:9px;flex-shrink:0;'
                    f'background:{logo_bg};border:1px solid rgba(255,255,255,0.07);'
                    f'display:grid;place-items:center;font-family:\'Bricolage Grotesque\',sans-serif;'
                    f'font-weight:800;font-size:14px;color:#ecf4ee;">{initial}</div>'
                    f'<div style="flex:1;min-width:0;">'
                    f'<div style="font-family:\'DM Sans\',sans-serif;font-weight:600;font-size:13.5px;'
                    f'color:#ecf4ee;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{job_title}</div>'
                    f'<div style="font-family:\'DM Sans\',sans-serif;font-size:11.5px;color:#6e8a7b;margin-top:2px;">{company}</div>'
                    f'</div></div>',
                    unsafe_allow_html=True
                )

            with col_match:
                st.markdown(
                    f'<div style="display:flex;align-items:center;justify-content:center;min-height:28px;">'
                    f'<span style="font-family:\'Bricolage Grotesque\',sans-serif;font-weight:800;'
                    f'font-size:18px;color:#ecf4ee;">{match_pct}</span>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            with col_status:
                # Native HTML <select> — full color control, no Base Web conflicts
                opts_html = ''.join(
                    f'<option value="{o}" {"selected" if o == status else ""}>{o}</option>'
                    for o in STATUS_OPTIONS
                )
                _components.html(f"""
                <html><head><style>
                html,body{{margin:0;padding:0;background:transparent;overflow:hidden;}}
                select{{
                    font-family:'Bricolage Grotesque',system-ui,sans-serif;
                    font-weight:800;font-size:12px;
                    color:{cfg['color']};
                    background:{cfg['bg']};
                    border:1px solid {cfg['border']};
                    border-radius:20px;
                    padding:4px 8px 4px 12px;
                    cursor:pointer;outline:none;
                    width:100%;max-width:150px;
                    margin-top:5px;
                }}
                </style></head><body>
                <select onchange="try{{
                    var u=new URL(window.parent.location.href);
                    u.searchParams.set('sc',this.value);
                    u.searchParams.set('rid','{record_id}');
                    window.parent.location.href=u.href;
                }}catch(e){{}}">
                    {opts_html}
                </select>
                </body></html>
                """, height=38)

            with col_date:
                st.markdown(
                    f'<div style="display:flex;align-items:center;min-height:28px;">'
                    f'<span style="font-family:\'Space Mono\',monospace;font-size:11px;'
                    f'color:#6e8a7b;white-space:nowrap;">{date_val}</span>'
                    f'</div>',
                    unsafe_allow_html=True
                )

        st.markdown("<div style='height:16px;'></div>", unsafe_allow_html=True)

        # Download tracker Excel from tracker page too
        st.download_button(
            label="📥 Download Tracker (Excel)",
            data=generate_tracker_excel(tracker_data),
            file_name="job_applications.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="tracker_page_download",
            use_container_width=False
        )


# ============= STREAMLIT UI =============

def main():
    # ── Handle ?login=1[&plan=X] query params ────────────────────────────────
    if st.query_params.get('login'):
        plan = st.query_params.get('plan', '')
        st.query_params.clear()
        st.session_state.show_login = True
        if plan in ('free', 'pack5', 'pack20', 'pack30', 'pack50'):
            st.session_state.pending_plan = plan
        st.rerun()

    # ── Handle ?payment_return=1 (returning from Pay Advanced) ──────────────
    if st.query_params.get('payment_return'):
        st.query_params.clear()
        st.session_state.show_payment_pending = True
        # If session was lost during redirect, send to login first (not landing)
        if 'auth_user_id' not in st.session_state:
            st.session_state.show_login = True
        st.rerun()

    # ── Auth gate ─────────────────────────────────────────────────────────────
    if 'auth_user_id' not in st.session_state:
        if st.session_state.get('show_login'):
            show_login_page()
        else:
            show_landing()
        return

    # ── Post-login payment screens ────────────────────────────────────────────
    if st.session_state.pop('show_payment_pending', None):
        show_payment_pending()
        return

    pending_plan = st.session_state.get('pending_plan', '')
    if pending_plan in PACK_INFO:
        st.session_state.pop('pending_plan', None)
        show_payment_redirect(pending_plan)
        return
    st.session_state.pop('pending_plan', None)

    auth_user_id = st.session_state['auth_user_id']
    auth_email   = st.session_state.get('auth_email', '')
    profile = get_or_create_profile(auth_user_id, auth_email)
    st.session_state['user_profile'] = profile

    # ── Route to tracker or app workspace ─────────────────────────────────────
    page = st.session_state.get('page', 'app')
    if page == 'tracker':
        show_tracker()
        return

    # ── Sidebar ─────────────────────────────────────────────────────────────
    with st.sidebar:
        tracker_data = load_tracker_data(user_id=auth_user_id)
        count = len(tracker_data)
        st.markdown(f"""
        <div style="padding: 1.5rem 1rem 0;">

          <!-- Brand mark -->
          <div style="display:flex;align-items:center;gap:11px;margin-bottom:1.4rem;">
            <div style="width:34px;height:34px;border-radius:9px;background:linear-gradient(150deg,#7ad79f,#4fae7a);display:grid;place-items:center;font-family:'Bricolage Grotesque',system-ui,sans-serif;font-weight:800;font-size:16px;color:#06140f;box-shadow:0 4px 12px rgba(122,215,159,0.28);flex-shrink:0;">R</div>
            <div>
              <div style="font-family:'Bricolage Grotesque',system-ui,sans-serif;font-weight:700;font-size:17px;letter-spacing:-0.02em;color:#ecf4ee;line-height:1.1;">ResumeSync</div>
              <div style="font-family:'Space Mono',monospace;font-size:9px;letter-spacing:0.18em;text-transform:uppercase;color:#6e8a7b;margin-top:3px;">by VisualizePro</div>
            </div>
          </div>

          <!-- Divider -->
          <div style="border-top:1px solid rgba(159,182,168,0.12); margin-bottom:1.2rem;"></div>

          <!-- User / credits block -->
          <div style="background:#0c2019;border:1px solid rgba(159,182,168,0.10);
                      border-radius:10px;padding:12px 14px;margin-bottom:1.2rem;">
            <div style="font-family:'DM Sans',sans-serif;font-size:12px;
                        color:#9fb6a8;white-space:nowrap;overflow:hidden;
                        text-overflow:ellipsis;margin-bottom:6px;">{auth_email}</div>
            <div style="display:flex;align-items:center;justify-content:space-between;gap:8px;">
              <div style="display:flex;align-items:center;gap:8px;">
                <span style="font-family:'Bricolage Grotesque',sans-serif;font-weight:800;
                             font-size:22px;color:#7ad79f;line-height:1;">
                  {int(profile.get('credits') or 0)}
                </span>
                <span style="font-family:'DM Sans',sans-serif;font-size:11px;color:#6e8a7b;">
                  {"analysis" if int(profile.get('credits') or 0) == 1 else "analyses"} remaining
                </span>
              </div>
            </div>
          </div>

          <!-- What you get -->
          <div style="border-top:1px solid rgba(159,182,168,0.12); padding-top:1rem; margin-bottom:1.2rem;">
            <p style="font-family:'Space Mono',monospace; font-size:9px; letter-spacing:0.16em; color:#7ad79f; text-transform:uppercase; margin:0 0 0.8rem;">What you get</p>
            <div style="display:flex; flex-direction:column; gap:0.6rem;">
              <div style="display:flex; gap:10px; align-items:flex-start;">
                <span style="color:#7ad79f; font-size:14px; flex-shrink:0; margin-top:1px;">📊</span>
                <span style="font-family:'DM Sans',sans-serif; font-size:0.8rem; color:#ecf4ee; line-height:1.4;">Match % score <span style="color:#9fb6a8;">— see your fit before you apply</span></span>
              </div>
              <div style="display:flex; gap:10px; align-items:flex-start;">
                <span style="color:#7ad79f; font-size:14px; flex-shrink:0; margin-top:1px;">🎯</span>
                <span style="font-family:'DM Sans',sans-serif; font-size:0.8rem; color:#ecf4ee; line-height:1.4;">Skills &amp; keyword gaps <span style="color:#9fb6a8;">— exactly what's missing</span></span>
              </div>
              <div style="display:flex; gap:10px; align-items:flex-start;">
                <span style="color:#7ad79f; font-size:14px; flex-shrink:0; margin-top:1px;">🤖</span>
                <span style="font-family:'DM Sans',sans-serif; font-size:0.8rem; color:#ecf4ee; line-height:1.4;">ATS optimisation <span style="color:#9fb6a8;">— get past the robot screeners</span></span>
              </div>
              <div style="display:flex; gap:10px; align-items:flex-start;">
                <span style="color:#7ad79f; font-size:14px; flex-shrink:0; margin-top:1px;">✍️</span>
                <span style="font-family:'DM Sans',sans-serif; font-size:0.8rem; color:#ecf4ee; line-height:1.4;">Tailored cover letter <span style="color:#9fb6a8;">— ready to send in seconds</span></span>
              </div>
              <div style="display:flex; gap:10px; align-items:flex-start;">
                <span style="color:#7ad79f; font-size:14px; flex-shrink:0; margin-top:1px;">✨</span>
                <span style="font-family:'DM Sans',sans-serif; font-size:0.8rem; color:#ecf4ee; line-height:1.4;">Resume auto-updater <span style="color:#9fb6a8;">— stronger words, same facts</span></span>
              </div>
            </div>
          </div>

          <!-- Job tracker count -->
          <div style="background:#0c2019;border:1px solid rgba(159,182,168,0.12);border-radius:12px;padding:16px;">
            <p style="font-family:'Space Mono',monospace; font-size:9.5px; letter-spacing:0.16em; text-transform:uppercase; color:#6e8a7b; margin:0 0 0.5rem;">Applications tracked</p>
            <p style="font-family:'Bricolage Grotesque',system-ui,sans-serif; font-size:2rem; font-weight:800; color:#7ad79f; margin:0 0 0.2rem; line-height:1;">
              {count}
            </p>
            <p style="font-family:'Space Mono',monospace; font-size:0.68rem; color:#6e8a7b; margin:0; line-height:1.4;">
              {'Track your first application below' if count == 0 else 'role(s) in your tracker →'}
            </p>
          </div>

        </div>
        """, unsafe_allow_html=True)

        if st.button("↻ Refresh credits", key="refresh_credits_btn", use_container_width=True):
            st.session_state.pop('user_profile', None)
            st.rerun()

        _sb_email = auth_email
        _sb_links = []
        for _pk, _pi in PACK_INFO.items():
            _ref = f"{_sb_email}|{_pk}"[:50]
            _params = urllib.parse.urlencode({
                "paymentamount": _pi["amount"],
                "paymentdescription": f"ResumeSync {_pi['name']} ({_pi['credits']} analyses)",
                "paymentref": _ref,
                "email": _sb_email,
            })
            _url = PAY_ADVANCED_URL + "?" + _params
            _sb_links.append(f"""
            <a href="{_url}" target="_blank" style="
                display:block;padding:0.45rem 0.6rem;margin-bottom:0.4rem;
                background:linear-gradient(135deg,#7ad79f,#4fae7a);
                color:#000000;font-family:'DM Sans',sans-serif;font-weight:700;
                font-size:12px;text-align:center;border-radius:6px;
                text-decoration:none;">
                {_pi['credits']} analyses &nbsp;·&nbsp; {_pi['price']}
            </a>""")
        st.markdown(
            '<div style="border-top:1px solid rgba(159,182,168,0.12);padding-top:1rem;margin:0 0 0.6rem;">'
            '<p style="font-family:\'Space Mono\',monospace;font-size:9px;letter-spacing:0.16em;'
            'color:#7ad79f;text-transform:uppercase;margin:0 0 0.7rem;">Top up credits</p>'
            + "".join(_sb_links)
            + '</div>',
            unsafe_allow_html=True
        )

        _has_analysis = 'analysis_result' in st.session_state
        if not st.session_state.get('_confirm_leave_tracker'):
            if st.button("📋 View Applications", key="sidebar_view_tracker", use_container_width=True):
                if _has_analysis:
                    st.session_state['_confirm_leave_tracker'] = True
                    st.rerun()
                else:
                    st.session_state.page = 'tracker'
                    st.rerun()
        else:
            st.markdown(
                '<div style="background:rgba(224,122,95,0.10);border:1px solid rgba(224,122,95,0.3);'
                'border-radius:8px;padding:0.75rem 0.9rem;margin-bottom:0.5rem;">'
                '<p style="color:#e07a5f;font-family:\'DM Sans\',sans-serif;font-size:0.82rem;'
                'font-weight:600;margin:0 0 0.3rem;">⚠️ Leave this analysis?</p>'
                '<p style="color:#9fb6a8;font-family:\'DM Sans\',sans-serif;font-size:0.78rem;margin:0;">'
                'Download your report first — it won\'t be here when you come back.</p>'
                '</div>',
                unsafe_allow_html=True
            )
            _res = st.session_state.get('analysis_result', {})
            _rc  = st.session_state.get('resume_text', '')
            _jc  = st.session_state.get('job_content', '')
            _jurl = st.session_state.get('job_url', '')
            _rfn  = st.session_state.get('resume_filename', 'resume')
            if _res:
                st.download_button(
                    label="💾 Download Report",
                    data=create_analysis_docx(_res.get('analysis', ''), _jurl, _rfn, _jc),
                    file_name=f"ResumeAnalysis_{_rfn.rsplit('.', 1)[0][:20]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="sidebar_dl_before_leave",
                    use_container_width=True
                )
            leave_col, stay_col = st.columns(2)
            with leave_col:
                if st.button("Go anyway", key="confirm_leave_tracker", use_container_width=True):
                    st.session_state.pop('_confirm_leave_tracker', None)
                    st.session_state.page = 'tracker'
                    st.rerun()
            with stay_col:
                if st.button("Stay", key="cancel_leave_tracker", use_container_width=True):
                    st.session_state.pop('_confirm_leave_tracker', None)
                    st.rerun()

        st.markdown("""
        <a href='https://madelivniekerk.github.io/Resumesync/' target='_self'
           style='display:block;text-align:center;font-family:DM Sans,sans-serif;font-size:13px;
                  color:#6e8a7b;text-decoration:none;padding:6px 0;margin-bottom:6px;'>
          ← Back to home
        </a>""", unsafe_allow_html=True)

        if tracker_data:
            st.download_button(
                label="📥 Download Tracker",
                data=generate_tracker_excel(tracker_data),
                file_name="job_applications.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="sidebar_download_tracker",
                use_container_width=True
            )

        st.markdown("<div style='height:0.5rem;'></div>", unsafe_allow_html=True)
        if st.button("Sign out", key="sign_out_btn", use_container_width=True):
            for _k in ['auth_user_id','auth_email','auth_token','auth_refresh','user_profile','otp_sent_email']:
                st.session_state.pop(_k, None)
            st.rerun()

        pass  # chat is handled via st.chat_input() at the bottom of main()

    # ── Workspace compact CSS ─────────────────────────────────────────────────
    st.markdown("""
<style>
/* Kill oversized markdown headings everywhere in the workspace */
h1,h2,h3,h4{line-height:1.2!important;}
[data-testid="stMarkdownContainer"] h1,
.stMarkdown h1{font-size:0.95rem!important;font-weight:700!important;margin:0.5rem 0 0.15rem!important;color:#ecf4ee!important;}
[data-testid="stMarkdownContainer"] h2,
.stMarkdown h2{font-size:0.82rem!important;font-weight:600!important;margin:0.4rem 0 0.1rem!important;color:#ecf4ee!important;}
[data-testid="stMarkdownContainer"] h3,
.stMarkdown h3{font-size:0.76rem!important;font-weight:600!important;margin:0.3rem 0 0.1rem!important;color:#9fb6a8!important;}
[data-testid="stMarkdownContainer"] h4,
.stMarkdown h4{font-size:0.72rem!important;font-weight:600!important;margin:0.2rem 0 0.05rem!important;color:#9fb6a8!important;}
/* Tighten paragraph margins */
[data-testid="stMarkdownContainer"] p,.stMarkdown p{margin-block-end:0.2rem!important;font-size:0.82rem!important;}
/* Analysis results body — smaller text, scoped so it doesn't affect the rest of the app */
.st-key-analysis_results_body [data-testid="stMarkdownContainer"] p,
.st-key-analysis_results_body [data-testid="stMarkdownContainer"] li{
    font-size:0.74rem!important;line-height:1.5!important;
}
/* Thinner dividers */
hr{margin:0.4rem 0!important;border-color:rgba(159,182,168,0.12)!important;}
/* Checkbox labels */
[data-testid="stCheckbox"] label p{font-size:0.78rem!important;line-height:1.25!important;}
/* Text area */
[data-testid="stTextArea"] textarea{font-size:0.78rem!important;line-height:1.3!important;}
/* Radio labels */
[data-testid="stRadio"] label p{font-size:0.78rem!important;}
/* Form submit button smaller and unstyled */
[data-testid="stFormSubmitButton"] button{font-size:0.78rem!important;padding:0.25rem 0.6rem!important;}
/* Extra bottom padding so content isn't hidden behind the fixed chat bar */
section.main .block-container{padding-bottom:5rem!important;}
/* Fixed bottom container */
[data-testid="stBottom"]{background:#ffffff!important;border-top:1px solid rgba(0,0,0,0.10)!important;padding:0.25rem 1rem!important;}
[data-testid="stBottom"] > *{background:#ffffff!important;}
[data-testid="stBottom"] .block-container{padding:0.2rem 1rem!important;}
[data-testid="stBottomBlockContainer"]{padding:0.2rem 1rem!important;}
[data-testid="stBottom"] [data-testid="stVerticalBlock"]{gap:0!important;}
/* Chat input field */
[data-testid="stChatInput"]{background:#ffffff!important;border:1px solid rgba(0,0,0,0.18)!important;border-radius:8px!important;box-shadow:none!important;}
[data-testid="stChatInput"]:focus-within{border-color:rgba(0,0,0,0.35)!important;box-shadow:none!important;}
[data-testid="stChatInputContainer"]{padding:0.15rem 0.3rem!important;min-height:0!important;}
[data-testid="stChatInput"] textarea{background:#ffffff!important;color:#111111!important;font-family:'DM Sans',sans-serif!important;font-size:0.82rem!important;min-height:28px!important;max-height:28px!important;line-height:1.3!important;padding:0.3rem 0.5rem!important;resize:none!important;}
[data-testid="stChatInput"] textarea::placeholder{color:#888888!important;}
/* Send button */
[data-testid="stChatInput"] button{background:transparent!important;color:#444444!important;border:none!important;}
[data-testid="stChatInput"] button:hover{color:#111111!important;background:rgba(0,0,0,0.06)!important;}
/* Remove any outline/shadow */
[data-testid="stChatInput"] *{outline:none!important;box-shadow:none!important;}
/* Chat message bubbles */
[data-testid="stChatMessage"]{background:transparent!important;padding:0.3rem 0!important;}
/* Expander — default light theme clashes with the dark workspace */
[data-testid="stExpander"]{background:transparent!important;border:1px solid rgba(159,182,168,0.25)!important;border-radius:8px!important;}
[data-testid="stExpander"] summary{background:#0d1f16!important;border-radius:8px!important;}
[data-testid="stExpander"] summary:hover{background:#132a1f!important;}
[data-testid="stExpander"] summary p,
[data-testid="stExpander"] summary span{color:#ecf4ee!important;}
[data-testid="stExpander"] summary svg{fill:#ecf4ee!important;}
[data-testid="stExpander"] [data-testid="stExpanderDetails"]{background:#0d1f16!important;border-radius:0 0 8px 8px!important;}
</style>""", unsafe_allow_html=True)

    # ── Compact welcome topbar (matches handoff App design — lean workspace, not landing hero) ──
    st.markdown("""
    <div style="display:flex;align-items:center;justify-content:space-between;gap:20px;
                padding:20px 0 22px;">
      <div style="display:flex;align-items:center;gap:14px;">
        <div style="width:46px;height:46px;border-radius:11px;
                    background:linear-gradient(150deg,#1d3a31,#11251d);
                    border:1px solid rgba(122,215,159,0.22);
                    display:grid;place-items:center;
                    font-family:'Bricolage Grotesque',system-ui,sans-serif;
                    font-weight:700;font-size:18px;color:#7ad79f;flex-shrink:0;">M</div>
        <div>
          <div style="font-family:'Bricolage Grotesque',system-ui,sans-serif;font-weight:700;
                      font-size:22px;letter-spacing:-0.02em;color:#ecf4ee;line-height:1.1;">
            Let's find your next match.
          </div>
          <div style="font-size:13px;color:#9fb6a8;margin-top:3px;font-family:'DM Sans',sans-serif;">
            Upload your resume and a job posting to get your compatibility score.
          </div>
        </div>
      </div>
      <div style="display:flex;align-items:center;gap:7px;
                  font-family:'Space Mono',monospace;font-size:10px;
                  letter-spacing:0.12em;text-transform:uppercase;color:#6e8a7b;white-space:nowrap;">
        <span style="width:7px;height:7px;border-radius:50%;background:#7ad79f;
                     box-shadow:0 0 0 3px rgba(122,215,159,0.10);display:inline-block;"></span>
        <span style="width:7px;height:7px;border-radius:50%;background:rgba(159,182,168,0.22);display:inline-block;"></span>
        <span style="width:7px;height:7px;border-radius:50%;background:rgba(159,182,168,0.22);display:inline-block;"></span>
        <span style="margin-left:4px;">Step 1 of 3</span>
      </div>
    </div>
    """, unsafe_allow_html=True)

    client = get_client()
    if not client:
        st.error("⚠️ ANTHROPIC_API_KEY not found. On Streamlit Cloud: go to App Settings → Secrets and add your key in TOML format: ANTHROPIC_API_KEY = \"sk-ant-...\"")
        return

    # ── Workspace card ────────────────────────────────────────────────────────
    st.markdown("""
    <div style="background:#0c2019;border:1px solid rgba(159,182,168,0.12);border-radius:20px;
                padding:28px 28px 4px;box-shadow:0 24px 60px rgba(0,0,0,0.45),0 4px 14px rgba(0,0,0,0.35);
                margin-bottom:4px;">
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])
    with col1:
        st.markdown("""
        <div style="display:flex;align-items:center;gap:9px;
                    font-family:'Space Mono',monospace;font-size:11px;letter-spacing:0.14em;
                    text-transform:uppercase;color:#7ad79f;margin-bottom:10px;">
          <span style="width:19px;height:19px;border-radius:50%;background:rgba(122,215,159,0.16);
                       display:grid;place-items:center;font-size:10px;color:#7ad79f;flex-shrink:0;
                       text-align:center;line-height:19px;">1</span>
          Upload your resume
        </div>
        """, unsafe_allow_html=True)
        resume_file = st.file_uploader(
            "Resume", type=['pdf', 'docx', 'doc', 'txt'], label_visibility="collapsed"
        )
        if resume_file:
            # Cache bytes immediately so they survive all subsequent reruns
            st.session_state['_resume_bytes'] = resume_file.getvalue()
            st.session_state['_resume_name'] = resume_file.name
            st.markdown(f'<p style="font-size:0.75rem;color:#7ad79f;margin:4px 0 0;font-family:DM Sans,sans-serif;">✅ {resume_file.name}</p>', unsafe_allow_html=True)

    with col2:
        st.markdown("""
        <div style="display:flex;align-items:center;gap:9px;
                    font-family:'Space Mono',monospace;font-size:11px;letter-spacing:0.14em;
                    text-transform:uppercase;color:#7ad79f;margin-bottom:10px;">
          <span style="width:19px;height:19px;border-radius:50%;background:rgba(122,215,159,0.16);
                       display:grid;place-items:center;font-size:10px;color:#7ad79f;flex-shrink:0;
                       text-align:center;line-height:19px;">2</span>
          Add the job posting
        </div>
        """, unsafe_allow_html=True)
        input_method = st.radio(
            "Input method", ["Paste URL · auto-scrape", "Paste manually"],
            horizontal=True, label_visibility="collapsed"
        )
        job_url = None
        manual_job_text = None
        if input_method == "Paste URL · auto-scrape":
            job_url = st.text_input("Job URL", placeholder="https://www.seek.com.au/job/...", label_visibility="collapsed")
            if job_url:
                st.markdown(f'<p style="font-size:0.75rem;color:#7ad79f;margin:4px 0 0;font-family:DM Sans,sans-serif;">✅ {job_url[:55]}</p>', unsafe_allow_html=True)
        else:
            manual_job_text = st.text_area("Job description", placeholder="Paste the full job description here...", height=100, label_visibility="collapsed")
            if manual_job_text:
                st.markdown(f'<p style="font-size:0.75rem;color:#7ad79f;margin:4px 0 0;font-family:DM Sans,sans-serif;">✅ {len(manual_job_text)} characters</p>', unsafe_allow_html=True)
            manual_apply_url = st.text_input(
                "Application URL *(optional — detects your ATS)*",
                placeholder="https://boards.greenhouse.io/company/jobs/...",
                key="manual_apply_url"
            )

    st.markdown("""
    <div style="border-top:1px solid rgba(159,182,168,0.07);margin:12px 0 4px;"></div>
    """, unsafe_allow_html=True)

    has_resume = bool(st.session_state.get('_resume_bytes'))
    has_job = bool(job_url or manual_job_text)

    col_hint, col_btn = st.columns([2, 1])
    with col_hint:
        ready = has_resume and has_job
        dot_color = "#7ad79f" if ready else "#6e8a7b"
        hint_text = "Ready to analyze" if ready else "Add both to analyze"
        st.markdown(f"""
        <div style="display:flex;align-items:center;gap:8px;font-size:12px;color:{dot_color};
                    font-family:'DM Sans',sans-serif;padding-top:10px;">
          <span style="width:6px;height:6px;border-radius:50%;background:{dot_color};
                       display:inline-block;flex-shrink:0;"></span>
          {hint_text}
        </div>
        """, unsafe_allow_html=True)
    with col_btn:
        if st.button(
            "Analyze compatibility →",
            type="primary",
            use_container_width=True,
        ):
            if not has_resume:
                st.error("❌ Please upload a resume file first.")
            elif not has_job:
                st.error("❌ Please add a job posting (URL or paste).")
            else:
                st.session_state['_do_analysis'] = True
                st.session_state['_job_url_snap'] = job_url
                st.session_state['_manual_job_snap'] = manual_job_text

    if st.session_state.get('_do_analysis'):
        st.session_state['_do_analysis'] = False
        resume_bytes = st.session_state.get('_resume_bytes')
        resume_name  = st.session_state.get('_resume_name', 'resume')
        job_url      = st.session_state.get('_job_url_snap') or job_url
        manual_job_text = st.session_state.get('_manual_job_snap') or manual_job_text
        if not resume_bytes:
            st.error("❌ Please upload a resume file first.")
            return
        if not (job_url or manual_job_text):
            st.error("❌ Please add a job posting.")
            return

        # ── Usage limit check ─────────────────────────────────────────────────
        _allowed, _reason = can_run_analysis(profile)
        if not _allowed:
            st.error(f"🔒 {_reason}")
            _ue = st.session_state.get('auth_email', '')
            _pack_links = []
            for _pk, _pi in PACK_INFO.items():
                _ref = f"{_ue}|{_pk}"[:50]
                _params = urllib.parse.urlencode({
                    "paymentamount": _pi["amount"],
                    "paymentdescription": f"ResumeSync {_pi['name']} ({_pi['credits']} analyses)",
                    "paymentref": _ref,
                    "email": _ue,
                })
                _url = PAY_ADVANCED_URL + "?" + _params
                _pack_links.append(f"""
                <a href="{_url}" target="_blank" style="
                    display:block;padding:0.6rem 1rem;margin-bottom:0.5rem;
                    background:linear-gradient(135deg,#7ad79f,#4fae7a);
                    color:#000000;font-family:'DM Sans',sans-serif;font-weight:700;
                    font-size:14px;text-align:center;border-radius:8px;
                    text-decoration:none;box-shadow:0 4px 12px rgba(122,215,159,0.25);">
                    {_pi['name']} — {_pi['price']} &nbsp;·&nbsp; {_pi['credits']} analyses
                </a>""")
            st.markdown(
                '<p style="font-family:\'DM Sans\',sans-serif;font-size:13px;'
                'color:#9fb6a8;margin:0.75rem 0 0.5rem;">Top up — one-off, never expires:</p>'
                + "".join(_pack_links),
                unsafe_allow_html=True
            )
            return

        with st.spinner("🔍 Analyzing your resume against the job posting..."):

            with st.status("Extracting text from resume...", expanded=True) as status:
                resume_file_obj = io.BytesIO(resume_bytes)
                resume_file_obj.name = resume_name
                resume_text = extract_resume_text(resume_file_obj)
                if "Error" in resume_text:
                    st.error(f"❌ {resume_text}")
                    return
                if len(resume_text.strip()) < 100:
                    st.error("❌ Could not read enough text from your resume. If it's a PDF, it may be image-based (scanned) — try saving as DOCX instead. If it's a DOCX, re-save it from Word and re-upload.")
                    return
                st.write(f"✅ Extracted {len(resume_text)} characters from resume")
                status.update(label="Resume extracted!", state="complete")

            if job_url:
                with st.status("Fetching job posting from URL...", expanded=True) as status:
                    job_data = scrape_job_url(job_url)
                    if not job_data['success']:
                        st.error(f"❌ Could not fetch job posting: {job_data['error']}")
                        st.info("💡 Switch to 'Paste job description manually'.")
                        return
                    job_content = job_data['content']
                    st.session_state['detected_ats'] = detect_ats(job_url) or detect_ats_from_text(job_content)
                    st.write(f"✅ Fetched {len(job_content)} characters from job posting")
                    status.update(label="Job posting fetched!", state="complete")
            else:
                job_content = manual_job_text
                # Try to detect ATS from optional application URL or embedded URLs in the text
                _manual_apply_url = st.session_state.get('manual_apply_url', '')
                _detected_ats = detect_ats(_manual_apply_url) or detect_ats_from_text(manual_job_text)
                st.session_state['detected_ats'] = _detected_ats
                job_url = _manual_apply_url or "Manual Input"
                st.success(f"✅ Using manually pasted job description ({len(job_content)} characters)")

            with st.status("Analysing with Claude AI...", expanded=True) as status:
                result = analyze_resume_vs_job(resume_text, job_content, job_url, client)
                if not result['success']:
                    st.error(f"❌ Analysis failed: {result['error']}")
                    return
                status.update(label="Analysis complete!", state="complete")

        # Clear previous comparison results before storing new ones
        for _k in ['cover_letter', 'proposed_updates', 'updated_resume_bytes',
                   'updated_resume_name', 'updated_match_pct', 'tracker_saved',
                   '_confirm_leave_tracker', 'upd_guidance', '_upd_guidance_saved',
                   'analysis_chat', 'ats_fixed_docx_bytes', 'ats_fixed_notes']:
            st.session_state.pop(_k, None)

        st.session_state['analysis_result'] = result
        st.session_state['resume_text'] = resume_text
        st.session_state['job_content'] = job_content
        st.session_state['job_url'] = job_url
        st.session_state['resume_filename'] = resume_name
        st.session_state['resume_file_bytes'] = resume_bytes
        st.session_state['resume_is_docx'] = resume_name.lower().endswith(('.docx', '.doc'))
        st.session_state['tracker_saved'] = False

        # Only charge one analysis per unique resume+job combination
        _analysis_key = hashlib.md5((resume_text + job_content).encode()).hexdigest()
        if st.session_state.get('_last_analysis_key') != _analysis_key:
            decrement_credits(auth_user_id)
            st.session_state['_last_analysis_key'] = _analysis_key

        st.session_state['_scroll_to_results'] = True

    # Results
    if 'analysis_result' in st.session_state:
        result = st.session_state['analysis_result']
        resume_text = st.session_state['resume_text']
        job_content = st.session_state['job_content']
        job_url = st.session_state['job_url']
        resume_filename = st.session_state['resume_filename']

        st.markdown('<div id="analysis-results-anchor"></div>', unsafe_allow_html=True)

        if st.session_state.pop('_scroll_to_results', False):
            _components.html(
                '<script>'
                'window.parent.document.getElementById("analysis-results-anchor")'
                '.scrollIntoView({behavior:"smooth",block:"start"});'
                '</script>',
                height=0,
            )

        st.divider()

        col_new = st.columns([3, 1])[1]
        with col_new:
            if st.button("🔄 New Analysis", key="new_analysis"):
                for key in ['analysis_result', 'resume_text', 'job_content', 'job_url',
                            'resume_filename', 'resume_file_bytes', 'resume_is_docx',
                            'cover_letter', 'tracker_saved', 'proposed_updates',
                            'updated_resume_bytes', 'updated_resume_name', 'updated_match_pct',
                            'upd_guidance', '_upd_guidance_saved',
                            'trimmed_resume_text', 'trimmed_resume_cuts', 'analysis_chat',
                            'trim_docx_bytes', 'trim_pairs', 'trim_applied_count',
                            'ats_fixed_docx_bytes', 'ats_fixed_notes']:
                    st.session_state.pop(key, None)
                st.rerun()

        # ── Analysis full-width ───────────────────────────────────────────────
        fields = parse_analysis_fields(result['analysis'])

        st.markdown(
            '<div style="display:flex;align-items:center;gap:8px;margin:0.6rem 0 0.4rem;">'
            '<span style="font-family:\'Space Mono\',monospace;font-size:9.5px;letter-spacing:0.18em;'
            'text-transform:uppercase;color:#7ad79f;">📋 Analysis Results</span>'
            '<span style="flex:1;height:1px;background:rgba(159,182,168,0.15);"></span>'
            '</div>',
            unsafe_allow_html=True
        )

        # ---- Prominent match score badge ----
        _score_raw = fields.get('match_pct', '').replace('%', '').strip()
        if _score_raw.isdigit():
            _score_val = int(_score_raw)
            if _score_val >= 80:
                _score_color, _score_bg, _score_tier = '#7ad79f', 'rgba(122,215,159,0.10)', 'Strong fit'
            elif _score_val >= 60:
                _score_color, _score_bg, _score_tier = '#e0a14a', 'rgba(224,161,74,0.10)', 'Good fit'
            else:
                _score_color, _score_bg, _score_tier = '#e07a5f', 'rgba(224,122,95,0.10)', 'Needs work'
            st.markdown(
                f'<div style="background:{_score_bg};border:1px solid {_score_color};border-radius:14px;'
                f'padding:1rem 1.2rem;margin:0 0 0.8rem;display:flex;align-items:baseline;gap:0.6rem;'
                f'justify-content:center;">'
                f'<span style="font-family:\'Bricolage Grotesque\',serif;font-size:3rem;font-weight:800;'
                f'color:{_score_color};line-height:1;">{_score_val}%</span>'
                f'<span style="font-family:\'Space Mono\',monospace;font-size:0.85rem;letter-spacing:0.05em;'
                f'text-transform:uppercase;color:{_score_color};">{_score_tier}</span>'
                f'</div>',
                unsafe_allow_html=True
            )

        with st.container(key="analysis_results_body"):
            render_analysis(result['analysis'])

        # Build shared filename parts used across all downloads
        _fn_date = datetime.now().strftime('%Y-%m-%d')
        _fn_person = re.sub(r'[^\w\-]', '_', resume_filename.rsplit('.', 1)[0])[:25].strip('_')
        _fn_role = re.sub(r'[^\w\-]', '_', fields.get('job_title', 'Role').replace(' ', '_'))[:25].strip('_')

        # ---- ATS Detection panel ----
        _ats = st.session_state.get('detected_ats')
        if _ats:
            _tips_html = ''.join(
                f'<li style="color:#ecf4ee;font-size:0.78rem;font-family:\'DM Sans\',sans-serif;'
                f'line-height:1.5;margin-bottom:0.1rem;">{t}</li>'
                for t in _ats['tips']
            )
            st.markdown(
                f'<div style="background:rgba(255,255,255,0.02);border-left:3px solid {_ats["color"]};'
                f'border-radius:8px;padding:0.5rem 0.9rem;margin:0.3rem 0 0.4rem;">'
                f'<span style="font-family:\'Space Mono\',monospace;font-size:9px;letter-spacing:0.14em;'
                f'text-transform:uppercase;color:{_ats["color"]};">'
                f'{_ats["icon"]} {_ats["name"]} — {_ats["summary"]}</span>'
                f'<ul style="margin:0.3rem 0 0;padding-left:1.1rem;">{_tips_html}</ul>'
                f'</div>',
                unsafe_allow_html=True
            )

        # ---- Quick ATS pre-flight checks ----
        _ats_warnings = []
        _resume_name_lower = resume_filename.lower()
        if _resume_name_lower.endswith('.pdf'):
            _ats_warnings.append(
                '📄 <strong>File format:</strong> You uploaded a PDF. '
                'Many ATS systems parse Word documents more reliably — '
                'submit as <strong>.docx</strong> when the job application allows it.'
            )
        _word_count = len(resume_text.split())
        _est_pages  = estimate_pages(_word_count)
        if _word_count < 250:
            _ats_warnings.append(
                f'📏 <strong>Resume length:</strong> Only ~{_word_count:,} words extracted — your resume may look thin, '
                'or the file may be image-based (try uploading as .docx). '
                'Expand bullet points with context and measurable outcomes.'
            )
        if _ats_warnings:
            st.markdown(
                '<div style="background:rgba(224,161,74,0.08);border-left:3px solid #e0a14a;'
                'border-radius:8px;padding:0.5rem 0.9rem;margin:0.3rem 0 0.5rem;">'
                '<span style="font-family:\'Space Mono\',monospace;font-size:9px;letter-spacing:0.14em;'
                'text-transform:uppercase;color:#e0a14a;">⚠ ATS &nbsp;</span>'
                + ''.join(
                    f'<span style="color:#ecf4ee;font-size:0.80rem;font-family:\'DM Sans\',sans-serif;">{w}&nbsp; </span>'
                    for w in _ats_warnings
                )
                + '</div>',
                unsafe_allow_html=True
            )

        # ---- ATS Parse Preview — what the robot actually sees ----
        if _resume_name_lower.endswith('.docx'):
            try:
                _ats_struct = analyze_docx_ats_structure(st.session_state.get('resume_file_bytes'))
            except Exception:
                _ats_struct = None

            if _ats_struct:
                _struct_warnings = []
                if _ats_struct['has_tables']:
                    _struct_warnings.append(
                        f"🔲 <strong>{_ats_struct['table_count']} table(s) detected</strong> — strict ATS platforms "
                        "(Taleo, Workday) often drop table content entirely or scramble its reading order."
                    )
                if _ats_struct['has_multicolumn']:
                    _struct_warnings.append(
                        "▥ <strong>Multi-column layout detected</strong> — many ATS read left-to-right across the "
                        "whole line, interleaving your columns into garbled text."
                    )
                if _ats_struct['has_textboxes']:
                    _struct_warnings.append(
                        "🔳 <strong>Text box(es) detected</strong> — most ATS cannot read text inside text boxes at all; "
                        "that content is effectively invisible to the parser."
                    )
                if _ats_struct['has_contact_in_header_footer']:
                    _struct_warnings.append(
                        "📵 <strong>Contact info found in the header/footer</strong> — most ATS ignore headers and "
                        "footers completely, meaning your email or phone number may never reach the parsed profile."
                    )
                if not _ats_struct['has_skills_section']:
                    _struct_warnings.append(
                        "🏷️ <strong>No dedicated Skills section found</strong> — ATS platforms scan a Skills/Core "
                        "Competencies heading separately from your work history; add one if you don't have it."
                    )

                st.markdown(
                    '<div style="display:flex;align-items:center;gap:8px;margin:0.6rem 0 0.4rem;">'
                    '<span style="font-family:\'Space Mono\',monospace;font-size:9.5px;letter-spacing:0.18em;'
                    'text-transform:uppercase;color:#7ad79f;">🔬 ATS Parse Preview</span>'
                    '<span style="flex:1;height:1px;background:rgba(159,182,168,0.15);"></span>'
                    '</div>',
                    unsafe_allow_html=True
                )

                if _struct_warnings:
                    st.markdown(
                        '<div style="background:rgba(224,122,95,0.08);border-left:3px solid #e07a5f;'
                        'border-radius:8px;padding:0.5rem 0.9rem;margin:0 0 0.4rem;">'
                        + ''.join(
                            f'<p style="color:#ecf4ee;font-size:0.80rem;font-family:\'DM Sans\',sans-serif;'
                            f'margin:0.15rem 0;">{w}</p>'
                            for w in _struct_warnings
                        )
                        + '</div>',
                        unsafe_allow_html=True
                    )
                    if _ats_struct['at_risk_pct'] > 0:
                        st.markdown(
                            f'<p style="color:#9fb6a8;font-size:0.78rem;font-family:\'DM Sans\',sans-serif;'
                            f'margin:0.3rem 0 0.5rem;">~<strong style="color:#e07a5f;">{_ats_struct["at_risk_pct"]}%</strong> '
                            f'of your extracted content lives in tables/headers/footers — the parts most at risk of '
                            f'being dropped by a strict ATS.</p>',
                            unsafe_allow_html=True
                        )

                    _fixable = _ats_struct['has_tables'] or _ats_struct['has_contact_in_header_footer']
                    if _fixable:
                        if _ats_struct['has_multicolumn'] or _ats_struct['has_textboxes']:
                            st.markdown(
                                '<p style="color:#9fb6a8;font-size:0.76rem;font-family:\'DM Sans\',sans-serif;'
                                'margin:0 0 0.4rem;">Multi-column layout and text boxes can\'t be safely auto-fixed — '
                                'they need re-formatting in Word. The table/header issues below can be fixed automatically.</p>',
                                unsafe_allow_html=True
                            )
                        col_fix = st.columns([1, 2, 1])[1]
                        with col_fix:
                            _fix_btn = st.button(
                                "🔧 Fix ATS-risk content",
                                key="fix_ats_risk_btn",
                                use_container_width=True,
                                type="primary",
                            )
                        if _fix_btn:
                            _fixed_bytes, _fix_notes = build_ats_safe_docx(
                                st.session_state.get('resume_file_bytes'), _ats_struct
                            )
                            st.session_state['ats_fixed_docx_bytes'] = _fixed_bytes
                            st.session_state['ats_fixed_notes'] = _fix_notes

                        if st.session_state.get('ats_fixed_docx_bytes'):
                            st.markdown(
                                '<div style="background:rgba(122,215,159,0.06);border-left:3px solid #7ad79f;'
                                'border-radius:8px;padding:0.5rem 0.9rem;margin:0.4rem 0;">'
                                '<span style="font-family:\'Space Mono\',monospace;font-size:9px;letter-spacing:0.14em;'
                                'text-transform:uppercase;color:#7ad79f;">✓ Fixed — nothing removed, only added</span>'
                                + ''.join(
                                    f'<p style="color:#ecf4ee;font-size:0.78rem;font-family:\'DM Sans\',sans-serif;'
                                    f'margin:0.3rem 0 0;">{html.escape(n)}</p>'
                                    for n in st.session_state.get('ats_fixed_notes', [])
                                )
                                + '</div>',
                                unsafe_allow_html=True
                            )
                            st.download_button(
                                label="💾 Download ATS-Safe Version (.docx)",
                                data=st.session_state['ats_fixed_docx_bytes'],
                                file_name=f"ATSSafe_{_fn_person}_{_fn_date}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_ats_safe",
                                use_container_width=True
                            )
                else:
                    st.markdown(
                        '<div style="background:rgba(122,215,159,0.06);border-left:3px solid #7ad79f;'
                        'border-radius:8px;padding:0.5rem 0.9rem;margin:0 0 0.4rem;">'
                        '<p style="color:#ecf4ee;font-size:0.80rem;font-family:\'DM Sans\',sans-serif;margin:0;">'
                        '✅ Clean structure — no tables, columns, text boxes, or header-only contact info detected. '
                        'This resume should parse consistently across ATS platforms.</p>'
                        '</div>',
                        unsafe_allow_html=True
                    )

                with st.expander("👁 See exactly what a strict ATS extracts from your file"):
                    st.caption(
                        "This is your resume with formatting stripped — paragraphs only, no tables, no "
                        "headers/footers. It's the closest approximation of what Taleo/Workday-style parsers read."
                    )
                    st.markdown(
                        '<pre style="background:#0d1f16;color:#ecf4ee;border:1px solid rgba(159,182,168,0.25);'
                        'border-radius:8px;padding:0.75rem 0.9rem;max-height:260px;overflow-y:auto;'
                        'white-space:pre-wrap;word-break:break-word;font-family:\'DM Sans\',sans-serif;'
                        'font-size:0.74rem;line-height:1.5;margin:0 0 0.6rem;">'
                        + html.escape(_ats_struct['strict_text'] or "(no plain-paragraph text found)")
                        + '</pre>',
                        unsafe_allow_html=True
                    )
                    if _ats_struct['table_text'] or _ats_struct['header_footer_text']:
                        st.caption("Content found only in tables/headers/footers (at risk of being dropped):")
                        _dropped = "\n".join(filter(None, [_ats_struct['table_text'], _ats_struct['header_footer_text']]))
                        st.markdown(
                            '<pre style="background:#0d1f16;color:#e0a14a;border:1px solid rgba(224,161,74,0.35);'
                            'border-radius:8px;padding:0.75rem 0.9rem;max-height:140px;overflow-y:auto;'
                            'white-space:pre-wrap;word-break:break-word;font-family:\'DM Sans\',sans-serif;'
                            'font-size:0.74rem;line-height:1.5;margin:0;">'
                            + html.escape(_dropped)
                            + '</pre>',
                            unsafe_allow_html=True
                        )
        elif _resume_name_lower.endswith('.pdf'):
            st.caption(
                "🔬 ATS Parse Preview is available for .docx uploads — structural checks (tables, columns, "
                "text boxes) need the document's underlying layout, which isn't inspectable in a PDF."
            )

        # ── Resume Trim — always visible after analysis ───────────────────────
        st.markdown(
            '<div style="background:rgba(224,161,74,0.06);border-left:3px solid rgba(224,161,74,0.50);'
            'border-radius:8px;padding:0.5rem 0.9rem;margin:0.3rem 0 0.4rem;'
            'display:flex;align-items:baseline;gap:0.6rem;flex-wrap:wrap;">'
            '<span style="font-family:\'Space Mono\',monospace;font-size:9px;letter-spacing:0.14em;'
            'text-transform:uppercase;color:#e0a14a;white-space:nowrap;">✂ Trim</span>'
            f'<span style="color:#9fb6a8;font-size:0.80rem;font-family:\'DM Sans\',sans-serif;">'
            f'{_word_count:,} words (~{_est_pages} pages) — AI removes low-impact content, condenses old roles. Facts stay intact.</span>'
            '</div>',
            unsafe_allow_html=True
        )

        _trim_target = st.radio(
            "Target length",
            ["2 pages (recommended)", "1 page (very tight)"],
            horizontal=True,
            key="trim_target"
        )
        _trim_pages = 2 if "2 pages" in _trim_target else 1

        col_trim = st.columns([1, 2, 1])[1]
        with col_trim:
            trim_btn = st.button("✂ Trim My Resume", key="trim_resume_btn", use_container_width=True)

        if trim_btn:
            with st.status("Trimming your resume...", expanded=True) as trim_status:
                trim_status.write(f"Identifying cuts to reach {_trim_pages} pages...")
                trim_result = trim_resume(resume_text, client, target_pages=_trim_pages)
                if trim_result['success']:
                    pairs = trim_result['pairs']
                    trim_status.write(f"Applying {len(pairs)} cut(s) to your original document...")
                    _base_bytes = st.session_state.get('resume_file_bytes')
                    if _base_bytes:
                        _trimmed_bytes, _applied = apply_updates_to_docx(_base_bytes, pairs, resume_filename)
                        st.session_state['trim_docx_bytes']   = _trimmed_bytes
                        st.session_state['trim_pairs']        = pairs
                        st.session_state['trim_applied_count'] = _applied
                        # Clear any previous proposed changes — they were based on the untrimmed doc
                        for _k in ['proposed_updates', 'updated_resume_bytes', 'updated_resume_name',
                                   'updated_match_pct', 'upd_guidance', '_upd_guidance_saved']:
                            st.session_state.pop(_k, None)
                    trim_status.update(label=f"✓ {len(pairs)} item(s) removed — proceed to Update My Resume below", state="complete")
                else:
                    st.error(f"❌ Trim failed: {trim_result['error']}")

        if st.session_state.get('trim_docx_bytes'):
            _pairs        = st.session_state.get('trim_pairs', [])
            _applied      = st.session_state.get('trim_applied_count', 0)
            st.markdown(
                f'<div style="background:rgba(122,215,159,0.06);border-left:3px solid #7ad79f;'
                f'border-radius:8px;padding:0.5rem 0.9rem;margin:0.3rem 0 0.5rem;">'
                f'<span style="font-family:\'Space Mono\',monospace;font-size:9px;letter-spacing:0.14em;'
                f'text-transform:uppercase;color:#7ad79f;">✓ Trimmed — {_applied} item(s) removed</span>'
                f'<span style="color:#9fb6a8;font-size:0.78rem;font-family:\'DM Sans\',sans-serif;'
                f'display:block;margin-top:0.25rem;">Scroll down to <strong style="color:#ecf4ee;">✨ Update My Resume</strong> '
                f'to improve what remains, then download your final document.</span>'
                f'</div>',
                unsafe_allow_html=True
            )
            if _pairs:
                with st.expander(f"What was cut ({len(_pairs)} items)", expanded=False):
                    for p in _pairs:
                        st.markdown(f'- {p.get("description","")}: `{p.get("find","")[:80]}`')

            if st.button("↩ Undo trim", key="clear_trim"):
                for _k in ['trim_docx_bytes', 'trim_pairs', 'trim_applied_count']:
                    st.session_state.pop(_k, None)
                st.rerun()

        st.divider()

        # ============= RESUME UPDATER =============
        st.markdown(
            '<div style="display:flex;align-items:center;gap:8px;margin:0.8rem 0 0.4rem;">'
            '<span style="font-family:\'Space Mono\',monospace;font-size:9.5px;letter-spacing:0.18em;'
            'text-transform:uppercase;color:#7ad79f;">✨ Update My Resume</span>'
            '<span style="flex:1;height:1px;background:rgba(159,182,168,0.15);"></span>'
            '</div>',
            unsafe_allow_html=True
        )

        if st.session_state.get('resume_is_docx'):
            st.markdown(
                '<div style="background:rgba(122,215,159,0.06);padding:0.55rem 0.9rem;border-radius:10px;'
                'border-left:3px solid #7ad79f;margin-bottom:0.6rem;">'
                '<p style="color:#9fb6a8;font-size:0.82rem;margin:0;font-family:\'DM Sans\',sans-serif;line-height:1.5;">'
                'Stronger verbs, clearer outcomes — <strong style="color:#ecf4ee;">never adds skills you don\'t have.</strong> '
                'Review and approve every change before it\'s applied.</p>'
                '</div>',
                unsafe_allow_html=True
            )

            st.markdown(
                '<div style="background:rgba(111,177,224,0.08);padding:0.55rem 0.9rem;border-radius:10px;'
                'border-left:3px solid #6fb1e0;margin-bottom:0.6rem;">'
                '<p style="color:#6fb1e0;font-family:\'Space Mono\',monospace;font-size:9px;'
                'letter-spacing:0.14em;text-transform:uppercase;margin:0 0 0.3rem;">Tip — quantified impact</p>'
                '<p style="color:#ecf4ee;font-size:0.81rem;margin:0;font-family:\'DM Sans\',sans-serif;line-height:1.6;">'
                '📊 <b>%</b> &nbsp;·&nbsp; ⏱ <b>time saved</b> &nbsp;·&nbsp; 💰 <b>revenue / cost</b> &nbsp;·&nbsp; ⚡ <b>efficiency</b> &nbsp;·&nbsp; 📈 <b>scale</b><br>'
                '<span style="color:#6e8a7b;font-size:0.82rem;">'
                'e.g. "reduced processing time by 35%" &nbsp;·&nbsp; "increased accuracy by 20%" &nbsp;·&nbsp; "supported 5M+ records daily"'
                '</span></p>'
                '</div>',
                unsafe_allow_html=True
            )

            _guidance_val = st.session_state.get('upd_guidance', '').strip()
            if _guidance_val:
                c1, c2 = st.columns([10, 1])
                with c1:
                    st.markdown(
                        f'<div style="background:rgba(111,177,224,0.08);border-left:3px solid #6fb1e0;'
                        f'padding:0.4rem 0.9rem;border-radius:6px;margin-bottom:0.6rem;font-family:\'DM Sans\',sans-serif;">'
                        f'<span style="color:#6fb1e0;font-size:0.75rem;font-family:\'Space Mono\',monospace;'
                        f'text-transform:uppercase;letter-spacing:0.1em;">Guidance: </span>'
                        f'<span style="color:#ecf4ee;font-size:0.82rem;">{_guidance_val}</span>'
                        f'</div>',
                        unsafe_allow_html=True
                    )
                with c2:
                    if st.button("✕", key="clear_guidance", help="Clear guidance"):
                        st.session_state.pop('upd_guidance', None)
                        st.rerun()

            st.markdown('<div id="resume-updater-anchor"></div>', unsafe_allow_html=True)
            col_upd = st.columns([1, 2, 1])[1]
            with col_upd:
                improve_btn = st.button("✨ Propose Resume Changes", type="primary", key="update_resume_btn", use_container_width=True)

            if improve_btn:
                upd_guidance = st.session_state.get('upd_guidance', '').strip()
                st.session_state['_upd_guidance_saved'] = upd_guidance
                all_updates = []
                guidance_finds = set()
                boost_implied_count = 0

                with st.status("Improving your resume...", expanded=True) as upd_status:
                    # Step 0 — user guidance (locked, highest priority)
                    if upd_guidance:
                        upd_status.write("Applying your guidance first...")
                        guidance_result = generate_guidance_updates(resume_text, upd_guidance, client)
                        if guidance_result['success'] and guidance_result['updates']:
                            all_updates.extend(guidance_result['updates'])
                            guidance_finds = {u['find'] for u in guidance_result['updates']}
                            upd_status.write(f"✅ {len(guidance_result['updates'])} guidance change(s) locked in")
                        elif not guidance_result['success']:
                            upd_status.write(f"⚠️ Could not apply guidance: {guidance_result['error']}")

                    # Step 1 — general expression improvements (skip any bullet already touched by guidance)
                    upd_status.write("Analysing general improvements...")
                    upd_result = generate_resume_updates(resume_text, result['analysis'], client)
                    if upd_result['success']:
                        new_general = [u for u in upd_result['updates'] if u['find'] not in guidance_finds]
                        all_updates.extend(new_general)
                        upd_status.write(f"✅ {len(new_general)} expression improvement(s) found")
                    else:
                        upd_status.write(f"⚠️ Could not generate general changes: {upd_result['error']}")

                    # Step 2 — implied → exact terminology boost
                    upd_status.write("Converting implied keyword matches to exact terminology...")
                    boost_result = generate_implied_to_exact_updates(resume_text, result['analysis'], client)
                    if boost_result['success'] and boost_result['updates']:
                        boost_implied_count = boost_result.get('implied_count', len(boost_result['updates']))
                        boost_finds = {u['find'] for u in boost_result['updates']}
                        # Boost takes priority over general but not over user guidance
                        all_updates = [u for u in all_updates if u.get('type') == 'user_guidance' or u['find'] not in boost_finds]
                        all_updates.extend(boost_result['updates'])
                        upd_status.write(f"✅ {len(boost_result['updates'])} implied→exact terminology fix(es) added")
                    elif boost_result['success']:
                        upd_status.write("✅ No implied matches to convert — all keywords already exact")

                    # Step 3 — removal suggestions (unnecessary content)
                    upd_status.write("Checking for unnecessary content to remove...")
                    removal_result = generate_removal_suggestions(resume_text, client)
                    if removal_result['success'] and removal_result['updates']:
                        existing_finds = {u['find'] for u in all_updates}
                        new_removals = [u for u in removal_result['updates'] if u['find'] not in existing_finds]
                        all_updates.extend(new_removals)
                        upd_status.write(f"✅ {len(new_removals)} unnecessary line(s) flagged for removal")
                    elif removal_result['success']:
                        upd_status.write("✅ No unnecessary content found")

                    upd_status.update(label=f"{len(all_updates)} total change(s) ready for review", state="complete")

                if all_updates:
                    st.session_state['proposed_updates'] = all_updates
                    st.session_state['update_source'] = 'boost' if boost_implied_count else 'regular'
                    st.session_state['_boost_implied_count'] = boost_implied_count
                    st.session_state.pop('updated_resume_bytes', None)
                else:
                    st.error("❌ Could not generate any changes.")

                _components.html(
                    '<script>'
                    'window.parent.document.getElementById("resume-updater-anchor")'
                    '.scrollIntoView({behavior:"smooth",block:"start"});'
                    '</script>',
                    height=0,
                )

            # ── Review step ────────────────────────────────────────────────
            if 'proposed_updates' in st.session_state and st.session_state['proposed_updates']:
                proposed = st.session_state['proposed_updates']

                _saved_guidance = st.session_state.get('_upd_guidance_saved', '').strip()
                _current_guidance = st.session_state.get('upd_guidance', '').strip()
                if _current_guidance and _current_guidance != _saved_guidance:
                    st.markdown(
                        '<div style="background:rgba(224,161,74,0.10);border-left:3px solid #e0a14a;'
                        'padding:0.6rem 1rem;border-radius:8px;margin-bottom:0.8rem;font-family:\'DM Sans\',sans-serif;">'
                        '<span style="color:#e0a14a;font-size:0.75rem;font-family:\'Space Mono\',monospace;'
                        'text-transform:uppercase;letter-spacing:0.1em;">⚠ Guidance updated</span>'
                        '<span style="color:#9fb6a8;font-size:0.85rem;display:block;margin-top:0.3rem;">'
                        'Your additional guidance changed after these proposals were generated. '
                        'Click <strong style="color:#ecf4ee;">✨ Propose Resume Changes</strong> again to include it.'
                        '</span></div>',
                        unsafe_allow_html=True
                    )
                elif _saved_guidance:
                    st.markdown(
                        f'<div style="background:rgba(111,177,224,0.08);border-left:3px solid #6fb1e0;'
                        f'padding:0.5rem 0.9rem;border-radius:6px;margin-bottom:0.8rem;font-family:\'DM Sans\',sans-serif;">'
                        f'<span style="color:#6fb1e0;font-size:0.75rem;font-family:\'Space Mono\',monospace;'
                        f'text-transform:uppercase;letter-spacing:0.1em;">✓ Guidance included: </span>'
                        f'<span style="color:#9fb6a8;font-size:0.82rem;">{_saved_guidance}</span>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                st.markdown(
                    '<p style="font-family:\'Space Mono\',monospace;font-size:9px;letter-spacing:0.14em;'
                    'text-transform:uppercase;color:#7ad79f;margin:0.6rem 0 0.4rem;">'
                    'Review changes — untick any you want to skip:</p>',
                    unsafe_allow_html=True
                )

                selected = []
                for i, change in enumerate(proposed):
                    find           = change.get('find', '')
                    replace        = change.get('replace', '')
                    description    = change.get('description', 'Improve phrasing')
                    change_type    = change.get('type', '')
                    is_removal     = change_type == 'remove' or replace == ''
                    is_user_guided = change_type == 'user_guidance'

                    label_prefix = "⭐ YOUR GUIDANCE — " if is_user_guided else f"Change {i+1}: "
                    with st.container():
                        checked = st.checkbox(f"**{label_prefix}**{description}", value=True, key=f"chk_{i}")

                        if is_removal:
                            st.markdown(
                                f'<div style="background:rgba(224,80,70,0.08);border-left:3px solid #e05046;'
                                f'padding:0.7rem 0.9rem;border-radius:10px;font-size:0.82rem;'
                                f'font-family:\'DM Sans\',sans-serif;color:#9fb6a8;margin-bottom:0.8rem;">'
                                f'<strong style="color:#e05046;font-size:0.72rem;font-family:\'Space Mono\',monospace;'
                                f'letter-spacing:0.1em;text-transform:uppercase;">✂ Remove this line</strong><br>'
                                f'<span style="text-decoration:line-through;opacity:0.65;">{find}</span></div>',
                                unsafe_allow_html=True
                            )
                            edited_replace = ''
                        else:
                            before_border = '#e0a14a' if is_user_guided else '#e07a5f'
                            before_bg     = 'rgba(224,161,74,0.08)' if is_user_guided else 'rgba(224,122,95,0.08)'
                            col_b, col_a = st.columns(2)
                            with col_b:
                                st.markdown(
                                    f'<div style="background:{before_bg};border-left:3px solid {before_border};'
                                    f'padding:0.6rem 0.8rem;border-radius:10px;font-size:0.82rem;'
                                    f'font-family:\'DM Sans\',sans-serif;color:#9fb6a8;">'
                                    f'<strong style="color:{before_border};font-size:0.75rem;font-family:\'Space Mono\',monospace;'
                                    f'letter-spacing:0.1em;text-transform:uppercase;">{"⭐ Your guidance — before" if is_user_guided else "Before"}</strong><br>{find}</div>',
                                    unsafe_allow_html=True
                                )
                            with col_a:
                                st.markdown(
                                    f'<p style="color:#7ad79f;font-size:0.75rem;font-family:\'Space Mono\',monospace;'
                                    f'letter-spacing:0.1em;text-transform:uppercase;margin:0 0 0.3rem;">After — edit if needed</p>',
                                    unsafe_allow_html=True
                                )
                                edited_replace = st.text_area(
                                    f"After {i+1}",
                                    value=replace,
                                    height=90,
                                    key=f"edit_replace_{i}",
                                    label_visibility="collapsed"
                                )
                            st.markdown("<div style='margin-bottom:0.8rem'></div>", unsafe_allow_html=True)

                        if checked:
                            selected.append({**change, 'replace': edited_replace})

                n_selected = len(selected)
                col_apply = st.columns([1, 2, 1])[1]
                with col_apply:
                    apply_btn = st.button(
                        "✅ Apply " + str(n_selected) + " Selected Change" + ("s" if n_selected != 1 else ""),
                        type="primary",
                        key="apply_selected_btn",
                        disabled=n_selected == 0
                    )

                if apply_btn and selected:
                    # Use trimmed DOCX as base if the user trimmed first; else use original
                    _base_for_update = st.session_state.get('trim_docx_bytes') or st.session_state['resume_file_bytes']
                    updated_bytes, applied = apply_updates_to_docx(
                        _base_for_update,
                        selected,
                        resume_filename
                    )
                    new_filename = f"{_fn_person}_Updated_{_fn_role}_{_fn_date}.docx"
                    st.session_state['updated_resume_bytes'] = updated_bytes
                    st.session_state['updated_resume_name'] = new_filename
                    st.success(f"✅ {applied} change(s) applied to your resume.")
                    # Score improvement — method depends on change type
                    is_boost = st.session_state.get('update_source') == 'boost'
                    with st.spinner("Calculating updated compatibility score..."):
                        if is_boost:
                            # Analytical: promote `applied` implied items to exact using
                            # the original keyword tables — consistent with the analysis methodology
                            _orig_pct = fields.get('match_pct', '')
                            if not _orig_pct:
                                _m = re.search(r'COMPATIBILITY SCORE[^0-9]*(\d+)%', result.get('analysis', ''), re.IGNORECASE)
                                if _m:
                                    _orig_pct = _m.group(1) + '%'
                            _orig_num = int(re.search(r'\d+', _orig_pct).group()) if re.search(r'\d+', _orig_pct) else None
                            _tables = parse_keyword_tables(result.get('analysis', ''))
                            _before = compute_score_analytically(_tables, n_promote=0)
                            _after  = compute_score_analytically(_tables, n_promote=applied)
                            if _before is not None and _after is not None and _orig_num is not None:
                                _delta = _after - _before
                                new_score = min(100, _orig_num + _delta)
                            else:
                                new_score = _after  # fallback: raw analytical score
                        else:
                            # LLM re-score for general expression improvements
                            _upd_doc = Document(io.BytesIO(updated_bytes))
                            _upd_paras = list(_upd_doc.paragraphs)
                            for _t in _upd_doc.tables:
                                for _r in _t.rows:
                                    for _c in _r.cells:
                                        _upd_paras.extend(_c.paragraphs)
                            updated_text = "\n".join(p.text for p in _upd_paras if p.text.strip())
                            _orig_pct = fields.get('match_pct', '')
                            _orig_num = int(re.search(r'\d+', _orig_pct).group()) if re.search(r'\d+', _orig_pct) else None
                            new_score = re_score_resume(updated_text, job_content, client,
                                                        orig_score=_orig_num, changes=selected)
                        if new_score is not None:
                            st.session_state['updated_match_pct'] = f"{new_score}%"

            if 'updated_resume_bytes' in st.session_state:
                new_filename = st.session_state.get('updated_resume_name', f"{_fn_person}_Updated_{_fn_role}_{_fn_date}.docx")
                # Show score lift if re-score is available
                updated_pct = st.session_state.get('updated_match_pct')
                if updated_pct:
                    orig_pct = fields.get('match_pct', '')
                    # Also try scanning the raw analysis text directly as a fallback
                    if not orig_pct:
                        _m = re.search(r'COMPATIBILITY SCORE[^0-9]*(\d+)%', result.get('analysis', ''), re.IGNORECASE)
                        if _m:
                            orig_pct = _m.group(1) + '%'
                    orig_num = int(re.search(r'\d+', orig_pct).group()) if re.search(r'\d+', orig_pct) else None
                    new_num  = int(re.search(r'\d+', updated_pct).group()) if re.search(r'\d+', updated_pct) else None

                    if orig_num is not None and new_num is not None and orig_num != new_num:
                        # Full before → after with delta badge
                        delta = new_num - orig_num
                        delta_str = f"+{delta}" if delta > 0 else str(delta)
                        delta_color = "#7ad79f" if delta > 0 else "#ef4444"
                        st.markdown(
                            f'<div style="text-align:center;margin:8px 0 4px;">'
                            f'<span style="font-family:\'Space Mono\',monospace;font-size:12px;color:#9fb6a8;">Compatibility: </span>'
                            f'<span style="font-family:\'Bricolage Grotesque\',sans-serif;font-weight:800;font-size:17px;color:#9fb6a8;text-decoration:line-through;">{orig_pct}</span>'
                            f'<span style="font-family:\'Bricolage Grotesque\',sans-serif;font-size:17px;color:#6e8a7b;margin:0 6px;">→</span>'
                            f'<span style="font-family:\'Bricolage Grotesque\',sans-serif;font-weight:800;font-size:17px;color:#ecf4ee;">{updated_pct}</span>'
                            f'<span style="font-family:\'Space Mono\',monospace;font-size:11px;color:{delta_color};'
                            f'background:{"rgba(122,215,159,0.12)" if delta > 0 else "rgba(239,68,68,0.10)"};'
                            f'padding:2px 7px;border-radius:5px;margin-left:8px;">{delta_str} pts</span>'
                            f'</div>',
                            unsafe_allow_html=True
                        )
                    else:
                        # Show updated score on its own (no delta or delta = 0)
                        st.markdown(
                            f'<div style="text-align:center;margin:8px 0 4px;">'
                            f'<span style="font-family:\'Space Mono\',monospace;font-size:12px;color:#9fb6a8;">Updated compatibility: </span>'
                            f'<span style="font-family:\'Bricolage Grotesque\',sans-serif;font-weight:800;font-size:17px;color:#ecf4ee;">{updated_pct}</span>'
                            f'</div>',
                            unsafe_allow_html=True
                        )
                col_dl_upd = st.columns([1, 2, 1])[1]
                with col_dl_upd:
                    st.download_button(
                        label="📥 Download Updated Resume",
                        data=st.session_state['updated_resume_bytes'],
                        file_name=new_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_updated_resume"
                    )
        else:
            st.markdown(
                '<div style="background:rgba(122,215,159,0.04);border:1px solid rgba(159,182,168,0.12);'
                'border-radius:10px;padding:0.75rem 1rem;">'
                '<p style="color:#6e8a7b;font-size:0.85rem;margin:0;font-family:\'DM Sans\',sans-serif;">'
                '💡 Upload a Word document (.docx) to enable automatic resume updating.</p>'
                '</div>',
                unsafe_allow_html=True
            )

        st.divider()

        # ============= COVER LETTER =============
        st.markdown('<h2 style="color:#ecf4ee; font-size:1.8rem; font-weight:700; text-align:center; margin:2rem 0; font-family:Bricolage Grotesque,serif; letter-spacing:-0.02em;">✍️ Generate Cover Letter</h2>', unsafe_allow_html=True)
        st.markdown(
            '<div style="background:rgba(109,193,138,0.06); padding:1.2rem 1.5rem; border-radius:8px; border-left:4px solid #7ad79f; margin-bottom:2rem;">'
            '<p style="color:#9fb6a8; font-size:0.95rem; margin:0; font-family:DM Sans,sans-serif;">🎯 Create a personalised cover letter tailored to this job based on your match analysis</p>'
            '</div>',
            unsafe_allow_html=True
        )

        cl_col1, cl_col2, cl_col3 = st.columns(3)
        with cl_col1:
            cl_tone = st.selectbox(
                "Tone",
                ["Professional", "Conversational", "Enthusiastic", "Formal", "Confident & Direct"],
                key="cl_tone"
            )
        with cl_col2:
            cl_length = st.selectbox(
                "Length",
                ["Brief (≈200 words)", "Standard (300–350 words)", "Detailed (≈450 words)"],
                index=1,
                key="cl_length"
            )
        with cl_col3:
            cl_incorporate = st.checkbox(
                "Incorporate improvement recommendations",
                value=True,
                key="cl_incorporate",
                help="Weaves 1–2 specific recommendations from the analysis into the letter"
            )

        cl_guidance = st.text_area(
            "Additional guidance *(optional)*",
            placeholder="e.g. Mention my leadership of the 2023 digital transformation project. Emphasise my Python skills. Keep it humble but confident.",
            height=90,
            key="cl_guidance"
        )

        col_gen = st.columns([1, 2, 1])[1]
        with col_gen:
            generate_cl_button = st.button("🚀 Generate Cover Letter", type="primary", key="generate_cover_letter", use_container_width=True)

        if generate_cl_button:
            with st.spinner("✍️ Writing your personalised cover letter..."):
                with st.status("Generating cover letter...", expanded=True) as status:
                    cl_result = generate_cover_letter(
                        resume_text, job_content, job_url, result['analysis'], client,
                        tone=cl_tone, length=cl_length, incorporate_recs=cl_incorporate,
                        user_guidance=cl_guidance
                    )
                    if not cl_result['success']:
                        st.error(f"❌ Cover letter generation failed: {cl_result['error']}")
                    else:
                        st.write("✅ Cover letter generated!")
                        status.update(label="Cover letter ready!", state="complete")
                        st.session_state['cover_letter'] = cl_result['cover_letter']

        if 'cover_letter' in st.session_state:
            st.markdown("---")
            st.markdown(
                '<p style="font-family:\'DM Sans\',sans-serif;font-size:0.8rem;color:#9fb6a8;'
                'margin:0 0 0.3rem;">Edit directly below — your changes are reflected in the download.</p>',
                unsafe_allow_html=True
            )
            cl_edited = st.text_area(
                "Cover letter",
                value=st.session_state['cover_letter'],
                height=420,
                key="cl_edit_area",
                label_visibility="collapsed"
            )

            # ── Regenerate section ─────────────────────────────────────────
            st.markdown(
                '<p style="color:#9fb6a8; font-size:0.85rem; font-family:DM Sans,sans-serif; '
                'margin:0.5rem 0 0.3rem;">Want changes? Describe what to adjust and regenerate.</p>',
                unsafe_allow_html=True
            )
            cl_changes = st.text_area(
                "Proposed changes",
                placeholder="e.g. Make the opening more confident. Mention my Tableau experience earlier. Shorten the second paragraph. Use a more conversational tone in the closing.",
                height=100,
                key="cl_changes",
                label_visibility="collapsed"
            )

            regen_col, dl_col = st.columns([1, 1])
            with regen_col:
                regen_btn = st.button("🔄 Regenerate Cover Letter", key="regen_cover_letter", use_container_width=True)
            with dl_col:
                st.download_button(
                    label="💾 Download Cover Letter (.docx)",
                    data=create_cover_letter_docx(cl_edited),
                    file_name=f"CoverLetter_{_fn_person}_{_fn_role}_{_fn_date}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_cl",
                    use_container_width=True
                )

            if regen_btn:
                with st.spinner("✍️ Rewriting your cover letter..."):
                    with st.status("Applying your changes...", expanded=True) as regen_status:
                        cl_result = generate_cover_letter(
                            resume_text, job_content, job_url, result['analysis'], client,
                            tone=cl_tone, length=cl_length, incorporate_recs=cl_incorporate,
                            prior_letter=cl_edited,
                            change_instructions=cl_changes if cl_changes.strip() else None,
                            user_guidance=cl_guidance
                        )
                        if not cl_result['success']:
                            st.error(f"❌ Regeneration failed: {cl_result['error']}")
                        else:
                            st.write("✅ Cover letter updated!")
                            regen_status.update(label="Done!", state="complete")
                            st.session_state['cover_letter'] = cl_result['cover_letter']
                            st.session_state['cl_edit_area'] = cl_result['cover_letter']
                            st.rerun()

        st.divider()

        # ============= JOB TRACKER =============
        st.markdown('<h2 style="color:#ecf4ee; font-size:1.8rem; font-weight:700; text-align:center; margin:2rem 0; font-family:Bricolage Grotesque,serif; letter-spacing:-0.02em;">📊 Save to Job Tracker</h2>', unsafe_allow_html=True)

        t1, t2, t3 = st.columns(3)
        with t1:
            job_title_input = st.text_input("Job Title", value=fields['job_title'], key="tracker_title")
        with t2:
            company_input = st.text_input("Company", value=fields['company'], key="tracker_company")
        with t3:
            location_input = st.text_input("Location", value=fields['location'], key="tracker_location")

        col_save = st.columns([1, 2, 1])[1]
        with col_save:
            if st.session_state.get('tracker_saved'):
                st.success("✅ Saved to tracker!")
                if st.button("📋 View Applications →", key="goto_tracker_after_save", use_container_width=True):
                    st.session_state.page = 'tracker'
                    st.rerun()
            else:
                if st.button("💾 Save to Job Tracker", key="save_tracker", use_container_width=True):
                    try:
                        # Use edited cover letter text if available (from cl_edit_area widget)
                        cover_letter_text = (
                            st.session_state.get('cl_edit_area')
                            or st.session_state.get('cover_letter', '')
                        )
                        save_to_tracker(
                            job_title=job_title_input,
                            company=company_input,
                            location=location_input,
                            resume_filename=resume_filename,
                            match_pct=st.session_state.get('updated_match_pct') or fields['match_pct'],
                            job_url=job_url,
                            cover_letter=cover_letter_text,
                            cover_letter_path='',
                            notes=extract_recommendations_summary(result['analysis']),
                            updated_resume_file=st.session_state.get('updated_resume_name', ''),
                            user_id=st.session_state.get('auth_user_id'),
                        )
                        load_tracker_data.clear()
                        st.session_state['tracker_saved'] = True
                        st.rerun()
                    except Exception as _save_err:
                        st.error(f"❌ Could not save to tracker: {_save_err}")

        tracker_data = load_tracker_data()
        if tracker_data:
            with st.expander(f"📋 View Job Tracker ({len(tracker_data)} roles)", expanded=False):
                st.dataframe(tracker_data, use_container_width=True)
                st.download_button(
                    label="📥 Download Tracker (Excel)",
                    data=generate_tracker_excel(tracker_data),
                    file_name="job_applications.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_tracker"
                )

        st.divider()

        # Download full report
        col_download = st.columns([1, 2, 1])[1]
        with col_download:
            st.download_button(
                label="💾 Download Analysis Report (.docx)",
                data=create_analysis_docx(result['analysis'], job_url, resume_filename, job_content),
                file_name=f"ResumeAnalysis_{_fn_person}_{_fn_role}_{_fn_date}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_analysis",
                use_container_width=True
            )

    # Footer
    st.divider()
    st.markdown(
        '<div style="text-align:center;padding:2rem 0 1rem;font-family:\'DM Sans\',sans-serif;">'
        '<p style="color:#7ad79f; font-size:0.85rem; margin-bottom:0.4rem; font-family:Space Mono,monospace; letter-spacing:0.1em;">Powered by Claude AI</p>'
        '<p style="color:#9fb6a8; font-size:0.8rem; font-family:Space Mono,monospace; letter-spacing:0.08em;">© 2026 ResumeSync by VisualizePro · AI-Powered Career Intelligence</p>'
        '</div>',
        unsafe_allow_html=True
    )

    # ── Fixed bottom guidance bar — shown when analysis is active and resume is .docx ──
    if st.session_state.get('analysis_result') and st.session_state.get('resume_is_docx'):
        _guidance_prompt = st.chat_input(
            "Add guidance for resume update… e.g. 'reduced costs by 30%', 'managed 8 people', 'make it more senior'"
        )
        if _guidance_prompt:
            _existing = st.session_state.get('upd_guidance', '').strip()
            st.session_state['upd_guidance'] = (_existing + '; ' + _guidance_prompt) if _existing else _guidance_prompt
            st.toast(f"Added: \"{_guidance_prompt[:60]}{'…' if len(_guidance_prompt) > 60 else ''}\"", icon="✅")


if _IMPORT_ERROR:
    st.error(f"**Startup import error:** {_IMPORT_ERROR}")
    st.code(_tb.format_exc(), language="python")
    st.info("Check Streamlit Cloud → App Settings → Secrets for ANTHROPIC_API_KEY, SUPABASE_URL, SUPABASE_KEY.")
else:
    try:
        main()
    except Exception as _e:
        st.error(f"**App error:** {_e}")
        st.code(_tb.format_exc(), language="python")
        st.info("Check Streamlit Cloud → App Settings → Secrets for ANTHROPIC_API_KEY, SUPABASE_URL, SUPABASE_KEY.")
